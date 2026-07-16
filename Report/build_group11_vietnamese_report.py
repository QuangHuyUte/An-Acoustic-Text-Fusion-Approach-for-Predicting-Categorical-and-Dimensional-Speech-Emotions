from pathlib import Path
import hashlib
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "Group11_Final_Report_Vietnamese_Draft.md"
OUT_PATH = ROOT / "Group11_Final_Report.docx"
ASSET_DIR = ROOT / "group11_report_assets"
FEATURE_FIG_DIR = (
    ROOT.parent
    / "06_w2v_based_models"
    / "02_IEMOCAP Feature Extraction Emotion2Vec Acoustic"
    / "output"
    / "feature_figures"
)
TEXT_BRANCH_FIG_DIR = (
    ROOT.parent
    / "06_w2v_based_models"
    / "03C_Transcript_Pretrained_Text_MultiTask_5_10Fold"
    / "output_03c_transcript_text_branch"
    / "figures"
)
FORMULA_DIR = ROOT / "group11_report_assets" / "rendered_formulas"


FIGURES_AFTER_HEADING = {
    "3.2. Mô tả dataset IEMOCAP": [
        (
            ASSET_DIR / "fig_dataset_distribution.png",
            "Figure 3.1. Distribution of IEMOCAP four-class emotion labels after filtering.",
        )
    ],
    "3.6. Giao thức chia dữ liệu": [
        (
            ASSET_DIR / "fig_split_overview_group11.png",
            "Figure 3.2. Speaker/session-independent evaluation protocols used in Stage 2.",
        )
    ],
    "4.2. Feature extraction ở giai đoạn giữa kỳ": [
        (
            ASSET_DIR / "fig_midterm_domain_dataset_evidence.png",
            "Figure 4.1. Midterm feature evidence across RAVDESS, CREMA-D, TESS and SAVEE.",
        )
    ],
    "4.2.1. Waveform, RMS energy và intensity pattern": [
        (
            ASSET_DIR / "fig_midterm_waveform_rms_6emotions.png",
            "Figure 4.2. Midterm waveform and RMS energy comparison across six emotion classes.",
        )
    ],
    "4.2.2. MFCC, delta MFCC và delta-delta MFCC": [
        (
            ASSET_DIR / "fig_midterm_mfcc_delta_6emotions.png",
            "Figure 4.3. MFCC, delta MFCC and delta-delta MFCC comparison across six emotion classes.",
        )
    ],
    "4.2.3. Log-Mel spectrogram, delta log-Mel và spectral contrast": [
        (
            ASSET_DIR / "fig_midterm_logmel_contrast_6emotions.png",
            "Figure 4.4. Log-Mel, delta Log-Mel and spectral contrast comparison across six emotion classes.",
        ),
        (
            ASSET_DIR / "fig_midterm_logmel_delta_sample.png",
            "Figure 4.5. Log-Mel channel comparison for one sample.",
        ),
    ],
    "4.2.4. ZCR, pitch/F0, chroma, tonnetz và spectral descriptors": [
        (
            ASSET_DIR / "fig_midterm_spectral_shape_6emotions.png",
            "Figure 4.6. Spectral shape feature evidence across six emotion classes.",
        ),
        (
            ASSET_DIR / "fig_midterm_zcr_6emotions.png",
            "Figure 4.7. ZCR feature evidence across six emotion classes.",
        ),
    ],
    "4.5. Minh họa đặc trưng trên IEMOCAP": [
        (
            FEATURE_FIG_DIR / "one_sample_per_emotion_feature_panels.png",
            "Figure 4.8. IEMOCAP one-sample-per-emotion feature panels.",
        ),
        (
            FEATURE_FIG_DIR / "same_speaker_multi_emotion_feature_panels.png",
            "Figure 4.9. IEMOCAP same-speaker multi-emotion feature panels.",
        ),
        (
            FEATURE_FIG_DIR / "selected_feature_means_by_emotion_heatmap.png",
            "Figure 4.10. Selected acoustic feature means by emotion class.",
        ),
        (
            FEATURE_FIG_DIR / "selected_feature_vad_correlation_heatmap.png",
            "Figure 4.11. Correlation between selected features and VAD dimensions.",
        ),
        (
            FEATURE_FIG_DIR / "xstats_pca_by_emotion.png",
            "Figure 4.12. PCA visualization of X_stats by emotion class.",
        ),
    ],
    "5.2. Giai đoạn giữa kỳ: mô hình 06D acoustic-only multi-branch SER": [
        (
            ASSET_DIR / "fig_06d_midterm_architecture_redrawn.png",
            "Figure 5.1. Redrawn 06D midterm acoustic-only multi-branch SER architecture.",
        )
    ],
    "5.3. Giai đoạn cuối kỳ: mở rộng 06D lên IEMOCAP": [
        (
            ASSET_DIR / "fig_03b_03d_architecture.png",
            "Figure 5.2. Final-stage acoustic-text bridge and expert-level fusion architecture with acoustic, text, bridge, and output heads.",
        )
    ],
    "6.3. Kết quả các mô hình đề xuất": [
        (
            ASSET_DIR / "fig_model_result_comparison.png",
            "Figure 6.1. Summary comparison of WA, UAR, and Macro-F1.",
        ),
        (
            ASSET_DIR / "fig_ccc_result_comparison.png",
            "Figure 6.2. Summary comparison of CCC_mean for dimensional emotion regression.",
        ),
    ],
    "6.3.2. Kết quả mô hình Script-Only Tuned RoBERTa Text Encoder": [
        (
            TEXT_BRANCH_FIG_DIR / "transcript_dataset_overview.png",
            "Figure 6.3. Transcript branch dataset overview.",
        ),
        (
            TEXT_BRANCH_FIG_DIR / "training_curves_by_fold.png",
            "Figure 6.4. Training curves of the script-only tuned text branch by fold.",
        ),
        (
            TEXT_BRANCH_FIG_DIR / "paper_style_metric_summary.png",
            "Figure 6.5. Metric summary of the script-only tuned text branch.",
        ),
    ],
    "7.2. Quy trình suy luận": [
        (
            ASSET_DIR / "fig_group11_system_pipeline.png",
            "Figure 7.1. Web demo inference route from audio/script input to emotion and VAD outputs.",
        )
    ],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text.strip())
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_inline_markdown(paragraph, text, size=11, bold_default=False):
    text = text.strip()
    if not text:
        return
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.bold = bold_default
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(size)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(max(size - 1, 8))
            run.font.color.rgb = RGBColor(30, 74, 120)
        else:
            label = token[1 : token.index("]")]
            url = token[token.index("(") + 1 : -1]
            run = paragraph.add_run(f"{label} ({url})")
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0, 84, 147)
            run.underline = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold_default


def add_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_figure(doc, image_path, caption, width=6.25):
    if not image_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Missing figure: {image_path}]")
        run.italic = True
        run.font.color.rgb = RGBColor(180, 0, 0)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(31, 78, 121)),
        ("Heading 3", 11, RGBColor(31, 78, 121)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_cover_page(doc):
    lines = [
        "HO CHI MINH CITY UNIVERSITY OF TECHNOLOGY AND EDUCATION",
        "FACULTY OF HIGH QUALITY TRAINING",
        "",
        "FINAL REPORT",
        "SPEECH PROGRAMMING",
        "",
        "Emotion2Vec-Guided Acoustic-Text Cross-Attention",
        "Multi-Task Speech Emotion Recognition",
        "",
        "Group 11",
        "Nguyễn Tài Huy - Bùi Quang Huy - Nguyễn Minh Cường",
        "",
        "Ho Chi Minh City, 2026",
    ]
    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = "Calibri"
        if idx in [3, 4]:
            run.font.size = Pt(22 if idx == 3 else 18)
            run.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
        elif idx in [6, 7]:
            run.font.size = Pt(15)
            run.bold = True
        elif idx in [0, 1]:
            run.font.size = Pt(12)
            run.bold = True
        else:
            run.font.size = Pt(12)
    doc.add_page_break()


def split_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line):
    content = line.strip().strip("|").replace(" ", "")
    return bool(content) and set(content) <= set("-:")


def add_table(doc, table_lines):
    rows = [split_table_row(line) for line in table_lines if line.strip()]
    if len(rows) >= 2 and is_table_separator(table_lines[1]):
        header = rows[0]
        data = rows[2:]
    else:
        header = rows[0]
        data = rows[1:]
    ncols = max(len(header), *(len(r) for r in data)) if data else len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i in range(ncols):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "5B9BD5")
        set_cell_text(cell, header[i] if i < len(header) else "", True, (255, 255, 255), 8)
    for row in data:
        cells = table.add_row().cells
        for i in range(ncols):
            set_cell_text(cells[i], row[i] if i < len(row) else "", False, None, 8)
    doc.add_paragraph()


def normalize_formula_text(text):
    formula = " ".join(line.strip() for line in text.splitlines() if line.strip())
    formula = formula.strip()
    formula = formula.replace("\\mathbb{1}", "\\mathbf{1}")
    return formula


def latex_to_pretty_text(formula):
    formula = normalize_formula_text(formula)
    exact = {
        r"\mathcal{L} = \lambda_{cls}\mathcal{L}_{CE} + \lambda_{vad}\mathcal{L}_{VAD}":
            "ℒ = λ_cls · ℒ_CE + λ_vad · ℒ_VAD",
        r"Attention(Q,K,V)=softmax\left(\frac{QK^T}{\sqrt{d_k}}\right)V":
            "Attention(Q, K, V) = softmax( QKᵀ / √dₖ ) V",
        r"Attention(Q,K,V)=softmax(\frac{QK^T}{\sqrt{d_k}})V":
            "Attention(Q, K, V) = softmax( QKᵀ / √dₖ ) V",
        r"head_i = Attention(QW_i^Q, KW_i^K, VW_i^V)":
            "headᵢ = Attention(QWᵢᵠ, KWᵢᵏ, VWᵢᵛ)",
        r"head_i=Attention(QW_i^Q,KW_i^K,VW_i^V)":
            "headᵢ = Attention(QWᵢᵠ, KWᵢᵏ, VWᵢᵛ)",
        r"MHA(Q,K,V)=Concat(head_1,...,head_h)W^O":
            "MHA(Q, K, V) = Concat(head₁, …, headₕ)Wᴼ",
        r"m_a \sim Bernoulli(p_a), \quad m_t \sim Bernoulli(p_t)":
            "mₐ ∼ Bernoulli(pₐ),    mₜ ∼ Bernoulli(pₜ)",
        r"z_{fused}=Fusion(m_a z_{acoustic}, m_t z_{text})":
            "z_fused = Fusion(mₐ z_acoustic, mₜ z_text)",
        r"\hat{y}_{emo}=softmax(W_c h + b_c)":
            "ŷₑₘₒ = softmax(W_c h + b_c)",
        r"\hat{y}_{vad}=W_r h + b_r":
            "ŷᵥₐd = W_r h + b_r",
        r"L = \lambda_{emo}L_{CE} + \lambda_{vad}L_{VAD}":
            "L = λ_emo · L_CE + λ_vad · L_VAD",
        r"WA = \frac{\sum_i \mathbf{1}(y_i=\hat{y}_i)}{N}":
            "WA = (Σᵢ 𝟙(yᵢ = ŷᵢ)) / N",
        r"UAR = \frac{1}{C}\sum_{c=1}^{C}\frac{TP_c}{TP_c+FN_c}":
            "UAR = (1/C) Σ_{c=1}^{C} TP_c / (TP_c + FN_c)",
        r"MacroF1 = \frac{1}{C}\sum_{c=1}^{C}F1_c":
            "Macro-F1 = (1/C) Σ_{c=1}^{C} F1_c",
        r"CCC = \frac{2\rho\sigma_x\sigma_y}{\sigma_x^2+\sigma_y^2+(\mu_x-\mu_y)^2}":
            "CCC = 2ρσₓσᵧ / (σₓ² + σᵧ² + (μₓ − μᵧ)²)",
        r"MAE = \frac{1}{N}\sum_i |y_i-\hat{y}_i|":
            "MAE = (1/N) Σᵢ |yᵢ − ŷᵢ|",
    }
    if formula in exact:
        return exact[formula]
    pretty = formula
    replacements = {
        "\\lambda": "λ",
        "\\rho": "ρ",
        "\\sigma": "σ",
        "\\mu": "μ",
        "\\sum": "Σ",
        "\\sqrt": "√",
        "\\sim": "∼",
        "\\quad": "   ",
        "\\hat{y}": "ŷ",
        "\\mathbf{1}": "𝟙",
        "\\mathcal{L}": "ℒ",
        "\\left": "",
        "\\right": "",
        "\\frac": "",
    }
    for old, new in replacements.items():
        pretty = pretty.replace(old, new)
    pretty = pretty.replace("{", "").replace("}", "")
    pretty = pretty.replace("^T", "ᵀ").replace("^2", "²")
    pretty = pretty.replace("...", "…")
    return pretty


def render_formula_image_with_pillow(formula, out):
    from PIL import Image, ImageDraw, ImageFont

    pretty = latex_to_pretty_text(formula)
    font_candidates = [
        Path(r"C:\Windows\Fonts\cambria.ttc"),
        Path(r"C:\Windows\Fonts\segoeui.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    font_path = next((fp for fp in font_candidates if fp.exists()), None)
    font = ImageFont.truetype(str(font_path), 38) if font_path else ImageFont.load_default()
    dummy = Image.new("RGBA", (10, 10), (255, 255, 255, 0))
    draw = ImageDraw.Draw(dummy)
    bbox = draw.textbbox((0, 0), pretty, font=font)
    w = max(bbox[2] - bbox[0] + 60, 480)
    h = max(bbox[3] - bbox[1] + 34, 82)
    img = Image.new("RGBA", (w, h), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    x = (w - (bbox[2] - bbox[0])) // 2
    y = (h - (bbox[3] - bbox[1])) // 2 - bbox[1]
    draw.rounded_rectangle((6, 6, w - 6, h - 6), radius=12, fill=(245, 248, 252, 255), outline=(209, 222, 236, 255), width=1)
    draw.text((x, y), pretty, font=font, fill=(23, 50, 77, 255))
    img.save(out)
    return out


def render_formula_image(formula):
    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    formula = normalize_formula_text(formula)
    digest = hashlib.md5(("formula-style-v2:" + formula).encode("utf-8")).hexdigest()[:12]
    out = FORMULA_DIR / f"formula_{digest}.png"
    if out.exists():
        return out
    try:
        import matplotlib.pyplot as plt

        fig = plt.figure(figsize=(0.01, 0.01), dpi=220)
        fig.patch.set_alpha(0)
        text = fig.text(0, 0, f"${formula}$", fontsize=17, color="#17324d")
        fig.canvas.draw()
        bbox = text.get_window_extent()
        width = max(bbox.width / fig.dpi + 0.22, 1.0)
        height = max(bbox.height / fig.dpi + 0.18, 0.35)
        plt.close(fig)

        fig = plt.figure(figsize=(width, height), dpi=220)
        fig.patch.set_alpha(0)
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")
        ax.text(0.5, 0.5, f"${formula}$", fontsize=17, color="#17324d", ha="center", va="center")
        fig.savefig(out, transparent=True, bbox_inches="tight", pad_inches=0.04)
        plt.close(fig)
        return out
    except Exception:
        try:
            return render_formula_image_with_pillow(formula, out)
        except Exception:
            return None


def add_formula_image(doc, formula):
    formula = normalize_formula_text(formula)
    image_path = render_formula_image(formula)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path and image_path.exists():
        width_inches = 4.2
        try:
            from PIL import Image

            with Image.open(image_path) as img:
                width_inches = min(max(img.width / 220, 1.6), 5.8)
        except Exception:
            pass
        p.add_run().add_picture(str(image_path), width=Inches(width_inches))
        return
    run = p.add_run(formula)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(25, 55, 85)


def add_formula_block(doc, lines):
    formula = "\n".join(line for line in lines if line.strip())
    if formula.strip():
        add_formula_image(doc, formula)


def add_code_block(doc, lines):
    for line in lines:
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_shading(p, "EEF3F8")
        run = p.add_run(line)
        run.font.name = "Cambria Math"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(25, 55, 85)


def set_para_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Group 11 Final Report - Speech Emotion Recognition"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if header.runs:
            header.runs[0].font.size = Pt(8)
            header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        footer = section.footer.paragraphs[0]
        footer.text = "Ho Chi Minh City University of Technology and Education"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer.runs:
            footer.runs[0].font.size = Pt(8)
            footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def insert_heading_figures(doc, heading_text):
    for image_path, caption in FIGURES_AFTER_HEADING.get(heading_text, []):
        add_figure(doc, image_path, caption)


def build_report():
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    style_document(doc)
    add_cover_page(doc)

    start_idx = 0
    body_start_markers = {"# ACKNOWLEDGEMENT", "# LỜI CẢM ƠN"}
    for i, line in enumerate(lines):
        if line.strip() in body_start_markers:
            start_idx = i
            break

    in_code = False
    code_kind = ""
    code_lines = []
    in_math = False
    math_lines = []
    table_lines = []
    pending_bullets = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_table(doc, table_lines)
            table_lines = []

    def flush_bullets():
        nonlocal pending_bullets
        for item in pending_bullets:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, item, 10.5)
        pending_bullets = []

    for raw in lines[start_idx:]:
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                if code_kind == "math":
                    add_formula_block(doc, code_lines)
                else:
                    add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
                code_kind = ""
            else:
                flush_table()
                flush_bullets()
                in_code = True
                code_kind = line.strip()[3:].strip().lower()
            continue
        if in_code:
            code_lines.append(line)
            continue

        if line.strip() == "\\[":
            flush_table()
            flush_bullets()
            in_math = True
            math_lines = []
            continue
        if line.strip() == "\\]":
            if in_math:
                add_formula_block(doc, math_lines)
                math_lines = []
                in_math = False
            continue
        if in_math:
            math_lines.append(line)
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            flush_bullets()
            table_lines.append(line)
            continue
        else:
            flush_table()

        stripped = line.strip()
        if not stripped:
            flush_bullets()
            continue

        if stripped.startswith("- "):
            pending_bullets.append(stripped[2:])
            continue
        flush_bullets()

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            if (title.startswith("CHAPTER ") or title.startswith("CHƯƠNG ")) and len(doc.paragraphs) > 5:
                doc.add_page_break()
            p = doc.add_heading(title, level=1)
            insert_heading_figures(doc, title)
            continue
        if stripped.startswith("## "):
            title = stripped[3:].strip()
            doc.add_heading(title, level=2)
            insert_heading_figures(doc, title)
            continue
        if stripped.startswith("### "):
            title = stripped[4:].strip()
            doc.add_heading(title, level=3)
            insert_heading_figures(doc, title)
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline_markdown(p, stripped, 11)

    flush_table()
    flush_bullets()
    add_header_footer(doc)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build_report()
