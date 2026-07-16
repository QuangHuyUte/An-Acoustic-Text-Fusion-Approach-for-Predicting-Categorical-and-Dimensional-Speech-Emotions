from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\UTE\Speech Programming\Speech Project")
ROOT06 = ROOT / "06_w2v_based_models"
ASSETS = ROOT / "Report" / "group11_report_assets"
OUT_DOC = ROOT / "Report" / "Group11_Final_Report.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(15, 35, 64)
MUTED = RGBColor(85, 95, 110)
HEADER_FILL = "E8EEF5"


REFERENCES = [
    ("Busso et al.", "IEMOCAP: Interactive emotional dyadic motion capture database", "Language Resources and Evaluation, 2008", "https://doi.org/10.1007/s10579-008-9076-6"),
    ("Schuller et al.", "The INTERSPEECH 2009 Emotion Challenge", "INTERSPEECH, 2009", "https://www.isca-archive.org/interspeech_2009/schuller09_interspeech.html"),
    ("Vaswani et al.", "Attention Is All You Need", "NeurIPS, 2017", "https://arxiv.org/abs/1706.03762"),
    ("Devlin et al.", "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding", "NAACL, 2019", "https://arxiv.org/abs/1810.04805"),
    ("Liu et al.", "RoBERTa: A Robustly Optimized BERT Pretraining Approach", "arXiv, 2019", "https://arxiv.org/abs/1907.11692"),
    ("Baevski et al.", "wav2vec 2.0: A Framework for Self-Supervised Learning of Speech Representations", "NeurIPS, 2020", "https://arxiv.org/abs/2006.11477"),
    ("Chen et al.", "WavLM: Large-Scale Self-Supervised Pre-training for Full Stack Speech Processing", "IEEE JSTSP, 2022", "https://arxiv.org/abs/2110.13900"),
    ("Ma et al.", "emotion2vec: Self-Supervised Pre-Training for Speech Emotion Representation", "arXiv, 2023", "https://arxiv.org/abs/2312.15185"),
    ("Zhao et al.", "CA-MSER: Speech Emotion Recognition via Multi-level Acoustic Information and Co-Attention", "arXiv, 2022", "https://arxiv.org/abs/2203.15326"),
    ("Ye et al.", "TIM-Net: A Time Delay Neural Network Based Method for Speech Emotion Recognition", "arXiv, 2022", "https://arxiv.org/abs/2211.08233"),
    ("Chen et al.", "SpeechFormer: A Hierarchical Efficient Framework Incorporating the Characteristics of Speech", "arXiv, 2022", "https://arxiv.org/abs/2203.03812"),
    ("Lian et al.", "DST: Deformable Speech Transformer for Emotion Recognition", "arXiv, 2023", "https://arxiv.org/abs/2302.13729"),
    ("Li et al.", "Multi-Task Multi-Modal Learning for Categorical and Dimensional Emotion Recognition", "arXiv, 2024", "https://arxiv.org/abs/2401.00536"),
    ("Latif et al.", "An Attention-Augmented End-to-End Multi-Task Learning Model for Emotion Prediction from Speech", "arXiv, 2019", "https://arxiv.org/abs/1903.12424"),
    ("Siriwardhana et al.", "Jointly Fine-Tuning Wav2Vec 2.0 for Speech Emotion Recognition", "arXiv, 2021", "https://arxiv.org/abs/2110.06309"),
    ("Pepino et al.", "Emotion Recognition from Speech Using wav2vec 2.0 Embeddings", "INTERSPEECH, 2021", "https://arxiv.org/abs/2104.03502"),
    ("Lin", "A Concordance Correlation Coefficient to Evaluate Reproducibility", "Biometrics, 1989", "https://doi.org/10.2307/2532051"),
    ("OpenAI Whisper", "Robust Speech Recognition via Large-Scale Weak Supervision", "ICML, 2023", "https://arxiv.org/abs/2212.04356"),
]


figures = []
tables = []


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=8.5, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9E2EF"):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn("w:" + edge))
        if elem is None:
            elem = OxmlElement("w:" + edge)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def add_caption(doc, kind, text):
    n = len(figures) + 1 if kind == "Figure" else len(tables) + 1
    (figures if kind == "Figure" else tables).append(text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{kind} {n}. {text}")
    set_run_font(r, size=9, italic=True, color=MUTED)


def add_table(doc, headers, rows, caption=None, font_size=8.0):
    if caption:
        add_caption(doc, "Table", caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], HEADER_FILL)
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=font_size, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, align=align)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_para(doc, text="", bold=False, italic=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), color=BLUE if level < 3 else DARK, bold=True)


def add_image(doc, path, caption, width=6.3):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, "Figure", caption)


def load_data():
    dist = pd.read_csv(ROOT06 / "01_IEMOCAP Dataset Analysis and Speaker-Independent Splits/output/reports/4class_emotion_distribution.csv")
    s5 = pd.read_csv(ROOT06 / "01_IEMOCAP Dataset Analysis and Speaker-Independent Splits/output/reports/5fold_session_split_sizes.csv")
    overview = pd.read_csv(ROOT06 / "01_IEMOCAP Dataset Analysis and Speaker-Independent Splits/output/reports/dataset_overview.csv")
    cache = pd.read_csv(ROOT06 / "02_IEMOCAP Feature Extraction Emotion2Vec Acoustic/output/feature_reports/feature_cache_summary.csv").iloc[0]
    fref = pd.read_csv(ROOT06 / "02_IEMOCAP Feature Extraction Emotion2Vec Acoustic/output/feature_reports/feature_reference_table.csv")
    res = pd.read_csv(ASSETS / "compiled_result_summary.csv")
    res03b = pd.read_csv(ROOT06 / "03B_CoAttention_Emotion2Vec_Acoustic_MultiTask_5_10Fold/03b-coattention-emotion2vec-acoustic-multitask-5fold/reports/03B_acoustic_text_bridge_rmm_results_by_fold.csv")
    return dist, s5, overview, cache, fref, res, res03b


def build():
    dist, s5, overview, cache, fref, res, res03b = load_data()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(12)

    header = sec.header.paragraphs[0]
    header.text = "Group 11 | Speech Emotion Recognition and VAD Regression"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.text = "Speech Processing Final Report - July 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    cover_lines = [
        ("HCMC UNIVERSITY OF TECHNOLOGY AND ENGINEERING", 13, True, 2),
        ("FACULTY OF INFORMATION TECHNOLOGY", 13, True, 36),
        ("FINAL REPORT", 20, True, 8),
        ("Course name: Speech Processing", 12, False, 20),
        ("SPEECH EMOTION RECOGNITION AND DIMENSIONAL AFFECT REGRESSION USING ACOUSTIC-TEXT BRIDGE FUSION", 17, True, 18),
        ("Lecturer name: MSc. Phu Khac Anh", 12, False, 8),
        ("List of members: Group 11", 12, False, 8),
        ("Nguyen Tai Huy - Bui Quang Huy - Nguyen Minh Cuong", 12, False, 88),
        ("Ho Chi Minh City, 7/2026", 12, False, 0),
    ]
    for text, size, bold, after in cover_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=DARK)
    doc.add_page_break()

    add_heading(doc, "ACKNOWLEDGEMENT")
    add_para(doc, "We would like to express our sincere gratitude to MSc. Phu Khac Anh for his guidance throughout the Speech Processing course. The project gave our team an opportunity to move from a midterm acoustic emotion recognition prototype toward a more complete research-style system with dataset analysis, feature extraction, multi-task learning, model comparison, and a live web demonstration.")
    add_para(doc, "We also thank the Faculty of Information Technology at HCMC University of Technology and Engineering for providing the learning environment and evaluation criteria that encouraged systematic experimentation rather than only model implementation.")

    add_heading(doc, "ABSTRACT")
    add_para(doc, "This report presents Group 11's final project on Speech Emotion Recognition (SER) and dimensional affect regression. The midterm stage focused on an acoustic 06D model for emotion classification. The final stage extends the project to the IEMOCAP corpus, where each utterance provides categorical emotion labels and dimensional ratings for valence, arousal, and dominance.")
    add_para(doc, "The implemented pipeline includes dataset construction, feature visualization, multi-task heads, bridge/cross-attention fusion, random modality masking, and a local web demo that accepts audio upload or microphone recording. Current results show that expert fusion between the 03B acoustic-text branch and the 03C transcript branch reaches 72.90% WA / 73.56% UAR with 0.652 CCC mean on the 5-fold session-independent protocol, and 74.02% WA / 74.64% UAR with 0.661 CCC mean on the 10-fold speaker-independent protocol.")

    add_heading(doc, "FIGURE LIST")
    for i, name in enumerate([
        "IEMOCAP 4-class subset distribution",
        "Speaker/session independent split overview",
        "One sample per emotion feature panels",
        "Same speaker, multiple emotions feature panels",
        "Proposed acoustic-text bridge fusion architecture",
        "Group 11 final system and web demo pipeline",
        "Current emotion metric comparison",
        "Current CCC comparison",
    ], 1):
        add_para(doc, f"Figure {i}. {name}")
    add_heading(doc, "TABLE LIST")
    for i, name in enumerate([
        "Evaluation criteria coverage",
        "Main related models and their role",
        "IEMOCAP dataset overview",
        "4-class emotion distribution",
        "VAD statistics by emotion",
        "Feature groups extracted in Notebook 02",
        "Experiment matrix for final ablation",
        "Current result summary",
        "03B fold-level result",
        "Comparison with literature baselines",
        "Notebook and artifact map",
    ], 1):
        add_para(doc, f"Table {i}. {name}")
    doc.add_page_break()

    add_heading(doc, "CHAPTER 1: INTRODUCTION")
    add_heading(doc, "1.1. Problem statement and motivation", 2)
    add_para(doc, "Speech emotion recognition is a central speech processing task because spoken language carries both semantic content and paralinguistic affective information. Emotion is expressed through pitch, energy, rhythm, pauses, spectral shape, and linguistic choices. This project therefore studies both categorical emotion recognition and dimensional affect prediction.")
    add_para(doc, "The final direction deliberately moved away from the earlier presentation-feedback formulation and focused on a research-computable SER problem. This choice is appropriate for the course timeline because IEMOCAP provides audio, transcript information, speaker/session metadata, categorical emotion labels, and valence-arousal-dominance ratings [1].")
    add_heading(doc, "1.2. Research subject and scope", 2)
    add_para(doc, "The research subject is an end-to-end speech emotion processing pipeline for IEMOCAP utterances. The main emotion classes are neutral, happy, sad, and angry, where excited utterances are merged into happy following the common 4-class IEMOCAP setup. The dimensional targets are valence, arousal/activation, and dominance.")
    add_para(doc, "The system scope covers audio preprocessing, speaker-independent data splits, multi-branch acoustic feature extraction, transcript preparation, acoustic-text fusion, two-head multi-task prediction, experimental comparison, and a local web demonstration. It does not claim cross-corpus robustness, multilingual SER, or production-level low-latency streaming deployment.")
    add_heading(doc, "1.3. Project objectives aligned with evaluation criteria", 2)
    add_table(doc, ["Criterion", "How this report addresses it"], [
        ("1.1 Complete pipeline", "Audio input, preprocessing, feature extraction, model inference, reporting, and web demo are implemented."),
        ("1.2 Selected model/method", "The selected method is a multi-task acoustic-text bridge fusion model inspired by 06D, emotion2vec, co-attention, and multi-task affective learning."),
        ("1.3 Improvements", "The project improves the midterm 06D direction by adding IEMOCAP strict splits, stronger feature cache, transcript branch, bridge fusion, RMM, and expert fusion."),
        ("1.4 Impact analysis", "Experiments compare 03B, 03C, and 03D across WA, UAR, Macro-F1, CCC, MAE, and fold stability."),
        ("1.5 Demo", "The web demo accepts audio upload/recording, runs preprocessing/model inference, shows feature evidence, exports reports, and compares processed queues."),
    ], "Evaluation criteria coverage")
    add_heading(doc, "1.4. Main contributions", 2)
    add_bullets(doc, [
        "Constructed a clean IEMOCAP 4-class + VAD metadata set with 5,531 utterances and speaker/session-independent folds.",
        "Implemented Notebook 02 to extract full 06D-style features: X_temporal, X_spectral, X_stats, X_e2v, and text-ready metadata.",
        "Implemented 03B acoustic-text bridge fusion, 03C transcript text tuning, and 03D expert fusion between acoustic and text experts.",
        "Reported categorical emotion metrics and dimensional regression metrics, avoiding the mistake of using accuracy for continuous VAD targets.",
        "Built a local web demo that visualizes waveform, log-Mel, feature trends, transcript, model route, and queue-to-queue comparison.",
    ])

    add_heading(doc, "CHAPTER 2: RELATED WORKS")
    add_heading(doc, "2.1. IEMOCAP and SER evaluation protocols", 2)
    add_para(doc, "IEMOCAP is a widely used benchmark for speech emotion research. It contains five sessions with ten speakers, dyadic conversations, audio, transcripts, categorical emotion labels, and dimensional affect ratings [1]. Because utterances from the same speaker can leak speaker-specific patterns, this project prioritizes leave-one-session-out and leave-one-speaker-out evaluation.")
    add_heading(doc, "2.2. Acoustic features and self-supervised speech representations", 2)
    add_para(doc, "Classical SER systems commonly use low-level descriptors such as MFCCs, energy, zero-crossing rate, pitch/F0, and spectral descriptors. Modern systems increasingly use self-supervised speech encoders such as wav2vec 2.0, WavLM, HuBERT, and emotion2vec [6]-[8].")
    add_heading(doc, "2.3. Co-attention, transformers, and feature fusion", 2)
    add_para(doc, "Co-attention models such as CA-MSER combine MFCC, spectrogram, and wav2vec2 branches through attention-based fusion [9]. Transformer-style attention is also central to DST and SpeechFormer [11], [12]. This motivates the 03B design: acoustic branch tokens are refined by self-attention, while text and acoustic representations meet through bridge cross-attention.")
    add_heading(doc, "2.4. Multi-task categorical and dimensional affect learning", 2)
    add_para(doc, "Categorical emotion and dimensional affect are related but not identical. Multi-task papers argue that shared representations can improve affect modeling when classification and regression heads are trained together [13], [14]. CCC is used for dimensional targets because it penalizes mismatch in correlation, mean, and variance [17].")
    add_table(doc, ["Model", "Input", "Architecture", "Role in this project"], [
        ("emotion2vec", "Raw speech -> pretrained speech emotion embedding", "Self-supervised emotion representation + downstream classifier", "Feature extractor/backbone reference [8]"),
        ("CA-MSER", "MFCC + spectrogram + wav2vec2", "BiLSTM/CNN/wav2vec2 with co-attention", "Closest reference for 06D-style co-attention [9]"),
        ("TIM-Net", "MFCC", "Temporal-aware multi-scale dilated CNN", "Lightweight acoustic baseline [10]"),
        ("SpeechFormer", "spectrogram/log-Mel/wav2vec", "Hierarchical speech transformer", "Reference for speech-aware temporal modeling [11]"),
        ("DST", "WavLM features", "Deformable Speech Transformer", "Strong transformer baseline [12]"),
        ("Multi-task multimodal 2024", "acoustic + linguistic", "classifier + regressor heads, multimodal fusion", "Motivates categorical + VAD multi-task setting [13]"),
    ], "Main related models and their role", font_size=7.2)

    add_heading(doc, "CHAPTER 3: DATASET AND PREPROCESSING")
    add_heading(doc, "3.1. IEMOCAP dataset overview", 2)
    ov = {r["metric"]: r["value"] for _, r in overview.iterrows()}
    add_table(doc, ["Item", "Value"], [
        ("Total utterances in parsed full metadata", f"{int(float(ov.get('total_utterances', 0))):,}"),
        ("Sessions", int(float(ov.get("sessions", 0)))),
        ("Speakers", int(float(ov.get("speakers", 0)))),
        ("Conversations", int(float(ov.get("conversations", 0)))),
        ("WAV found rate", f"{float(ov.get('wav_found_rate', 0)) * 100:.1f}%"),
        ("Transcript found rate", f"{float(ov.get('transcript_found_rate', 0)) * 100:.1f}%"),
        ("Mean duration", f"{float(ov.get('mean_duration_sec', 0)):.2f} seconds"),
        ("Mean word count", f"{float(ov.get('mean_word_count', 0)):.2f} words"),
    ], "IEMOCAP dataset overview")
    add_image(doc, ASSETS / "fig_dataset_distribution.png", "IEMOCAP 4-class subset distribution used in the final experiments.", 6.0)
    add_table(doc, ["Emotion", "Count", "Percent"], [(r.emotion_4class, int(r["count"]), f"{r['percent']:.2f}%") for _, r in dist.iterrows()], "4-class emotion distribution")
    add_heading(doc, "3.2. Valence, arousal, and dominance labels", 2)
    add_para(doc, "IEMOCAP dimensional labels are human ratings. Valence reflects positive versus negative affect, arousal/activation reflects emotional energy, and dominance reflects perceived control or assertiveness. They should be treated as continuous regression targets, not as class labels unless explicitly discretized.")
    add_table(doc, ["Emotion", "Valence mean", "Arousal mean", "Dominance mean"], [
        ("neutral", "2.971", "2.726", "2.831"),
        ("angry", "1.906", "3.636", "3.950"),
        ("sad", "2.253", "2.563", "2.828"),
        ("happy", "3.947", "3.411", "3.228"),
    ], "VAD statistics by emotion")
    add_heading(doc, "3.3. Speaker/session-independent split protocols", 2)
    add_para(doc, "Two strict split protocols are used. The 5-fold protocol holds out one full session as test and uses another session as validation. The 10-fold protocol holds out one speaker as test and another speaker as validation.")
    add_image(doc, ASSETS / "fig_split_overview_group11.png", "Train/validation/test sizes for the 5-fold and 10-fold split protocols.", 6.3)
    add_table(doc, ["Fold", "Train", "Val", "Test"], [(r["fold"], r["train"], r["val"], r["test"]) for _, r in s5.iterrows()], "5-fold session-independent split sizes", font_size=7.5)
    add_heading(doc, "3.4. Audio preprocessing", 2)
    add_bullets(doc, [
        "Load each utterance waveform and resolve it against the full IEMOCAP release directory.",
        "Convert audio to mono and use a consistent sample rate for downstream model extraction.",
        "Segment or pad/truncate consistently when a model branch expects fixed temporal dimensions.",
        "Use training-fold scalers/statistics for model input; do not fit normalization on validation, test, or live demo audio.",
        "Prepare transcript text with the same fold identifiers, emotion labels, and VAD targets used by the acoustic data.",
    ])

    add_heading(doc, "CHAPTER 4: FEATURE EXTRACTION AND PROPOSED METHOD")
    add_heading(doc, "4.1. Midterm 06D baseline", 2)
    add_para(doc, "The midterm stage developed a 06D acoustic SER prototype. It combined hand-crafted acoustic features and neural branches, including temporal acoustic descriptors, log-Mel/spectrogram representations, statistical functionals, and attention/fusion components. The strongest random-split result was around 80%, but the stricter speaker-independent setting was around 69%.")
    add_para(doc, "The final project keeps the useful idea from 06D - multi-branch acoustic evidence - but moves to IEMOCAP, strict folds, text integration, and dimensional VAD regression.")
    add_heading(doc, "4.2. Notebook 02 feature extraction", 2)
    add_para(doc, f"Notebook 02 produces a complete feature cache with {int(cache['n_samples'])} samples. The cache contains X_temporal {cache['X_temporal_shape']}, X_spectral {cache['X_spectral_shape']}, X_stats {cache['X_stats_shape']}, and X_e2v {cache['X_e2v_shape']}. emotion2vec extraction succeeded for all samples in the current cache.")
    add_table(doc, ["Feature group", "Shape", "Contains", "Used by", "Reason"], [(r["feature_group"], r["shape"], r["contains"], r["used_by"], r["why"]) for _, r in fref.iterrows()], "Feature groups extracted in Notebook 02", font_size=6.5)
    add_image(doc, ROOT06 / "02_IEMOCAP Feature Extraction Emotion2Vec Acoustic/output/feature_figures/one_sample_per_emotion_feature_panels.png", "Feature visualization for one sample per emotion from Notebook 02.", 6.3)
    add_image(doc, ROOT06 / "02_IEMOCAP Feature Extraction Emotion2Vec Acoustic/output/feature_figures/same_speaker_multi_emotion_feature_panels.png", "Same-speaker multi-emotion feature visualization from Notebook 02.", 6.3)
    add_heading(doc, "4.3. 03B acoustic-text bridge fusion model", 2)
    add_para(doc, "The 03B model builds an acoustic-token set from four branches: temporal, spectral, statistical, and emotion2vec. The text branch projects transcript representations into the same fusion space. Bridge cross-attention lets acoustic tokens query text evidence and text tokens query acoustic evidence. A balanced fusion gate then produces a shared representation for two heads: emotion classification and VAD regression.")
    add_para(doc, "The main formulas are: Attention(Q,K,V)=softmax(QK^T/sqrt(d_k))V; MultiHead(Q,K,V)=Concat(head_1,...,head_h)W^O. Random Modality Masking (RMM) randomly masks one modality during training so the fused model does not collapse into depending on only one branch.")
    add_image(doc, ASSETS / "fig_03b_03d_architecture.png", "Proposed 03B acoustic-text bridge fusion and two-head prediction architecture.", 6.3)
    add_heading(doc, "4.4. 03C transcript branch and 03D expert fusion", 2)
    add_para(doc, "The 03C branch studies how much information is available from transcript alone. It uses a pretrained text encoder/tokenizer and task-specific heads for emotion and VAD. The 03D model combines trained acoustic and text experts by loading fold-specific outputs and training a small fusion head.")
    add_heading(doc, "4.5. Web demo integration", 2)
    add_para(doc, "The web demo implements the same inference logic as the test pipeline: audio upload or recording, waveform standardization, optional ASR transcript generation, feature evidence display, model inference, report export, and queue comparison.")
    add_image(doc, ASSETS / "fig_group11_system_pipeline.png", "Final system and web demo processing route.", 6.3)

    add_heading(doc, "CHAPTER 5: IMPLEMENTATION")
    add_heading(doc, "5.1. Notebook and artifact organization", 2)
    add_table(doc, ["Notebook / module", "Purpose", "Outputs"], [
        ("01_IEMOCAP_Dataset_Analysis_and_Splits", "Parse full IEMOCAP, build metadata, statistics, 5-fold and 10-fold splits", "metadata, split CSV/JSON, dataset figures"),
        ("02_IEMOCAP_Feature_Extraction_Emotion2Vec_Acoustic", "Extract 06D acoustic features, emotion2vec features, text-ready metadata", "feature cache, feature figures, feature reference tables"),
        ("03B_CoAttention_Emotion2Vec_Acoustic_MultiTask", "Train acoustic-text bridge/RMM model", "fold metrics, history, plots, checkpoints when enabled"),
        ("03C_Transcript_Pretrained_Text_MultiTask", "Train transcript-only branch", "5-fold/10-fold metrics and predictions"),
        ("03D_03B_03C_Expert_Fusion", "Fuse acoustic expert and tuned text expert outputs", "5-fold/10-fold fusion results and predictions"),
        ("07_Web_Demo", "Interactive audio upload/recording demo", "feature evidence, queue comparison, report export"),
    ], "Notebook and artifact map", font_size=7.0)
    add_heading(doc, "5.2. Live demo processing design", 2)
    add_numbered(doc, [
        "User uploads an audio file or records speech from the browser microphone.",
        "The system decodes the signal, standardizes it, and places it into a processing queue.",
        "The backend or browser-side pipeline extracts waveform evidence, log-Mel evidence, frame-level feature trends, and transcript tokens.",
        "The model returns emotion probabilities and valence/arousal/dominance values.",
        "The UI shows feature evidence, segment diagnostics, transcript/timestamp, and allows JSON/HTML report export.",
        "After two queue items are processed, the user can compare their waveform, log-Mel, and feature statistics.",
    ])

    add_heading(doc, "CHAPTER 6: EXPERIMENTAL EVALUATION")
    add_heading(doc, "6.1. Metrics", 2)
    add_para(doc, "For emotion classification, the report uses Weighted Accuracy (WA), Unweighted Average Recall (UAR/UA), Macro-F1, and Weighted-F1 when available. For VAD regression, the report uses CCC, MAE, and RMSE. CCC is defined as 2*rho*sigma_x*sigma_y / (sigma_x^2 + sigma_y^2 + (mu_x - mu_y)^2), where x is the ground-truth rating and y is the prediction [17].")
    add_heading(doc, "6.2. Planned ablation matrix", 2)
    add_table(doc, ["Experiment", "Description", "Protocol", "Metrics", "Current status"], [
        ("acoustic_only", "03B acoustic branches only", "5-fold + 10-fold", "WA, UAR, Macro-F1, CCC, MAE", "Pending run"),
        ("acoustic_text_concat", "Acoustic token + text embedding concatenation", "5-fold + 10-fold", "WA, UAR, Macro-F1, CCC, MAE", "Pending run"),
        ("acoustic_text_bridge", "Bridge cross-attention without RMM", "5-fold + 10-fold", "WA, UAR, Macro-F1, CCC, MAE", "Pending run"),
        ("acoustic_text_bridge_rmm", "Bridge cross-attention + Random Modality Masking", "5-fold + 10-fold", "WA, UAR, Macro-F1, CCC, MAE", "5-fold available; 10-fold partially available via 03D internal acoustic expert"),
        ("text_tune", "Pretrained transcript branch", "5-fold + 10-fold", "WA, UAR, Macro-F1, CCC, MAE", "Available"),
        ("03D expert fusion", "Fusion of 03B acoustic expert and 03C text expert", "5-fold + 10-fold", "WA, UAR, Macro-F1, CCC, MAE", "Available"),
    ], "Experiment matrix for final ablation", font_size=6.7)
    add_heading(doc, "6.3. Current results", 2)
    add_table(doc, ["Model / protocol", "WA (%)", "UAR (%)", "Macro-F1 (%)", "CCC mean"], [(r["Label"].replace("\n", " / "), f"{r['WA']:.2f}", f"{r['UAR']:.2f}", f"{r['Macro_F1']:.2f}", f"{r['CCC_mean']:.3f}") for _, r in res.iterrows()], "Current result summary", font_size=7.4)
    add_image(doc, ASSETS / "fig_model_result_comparison.png", "Current WA, UAR, and Macro-F1 comparison across implemented branches.", 6.3)
    add_image(doc, ASSETS / "fig_ccc_result_comparison.png", "Current CCC mean comparison across implemented branches.", 6.3)
    add_heading(doc, "6.4. 03B fold-level analysis", 2)
    add_table(doc, ["Fold", "Best epoch", "WA", "UAR", "Macro-F1", "CCC mean"], [(r["fold"].replace("fold_", "F"), r["best_epoch"], f"{r['WA']*100:.2f}", f"{r['UAR']*100:.2f}", f"{r['Macro_F1']*100:.2f}", f"{r['CCC_mean']:.3f}") for _, r in res03b.iterrows()], "03B acoustic-text bridge-RMM 5-fold results by fold", font_size=7.0)
    add_para(doc, "03B bridge-RMM has a mean 5-fold result of 69.72% WA, 70.35% UAR, 69.58% Macro-F1, and 0.639 CCC mean. Fold 2 and Fold 5 are stronger for emotion classification, while Fold 4 has the highest CCC mean. This suggests that acoustic-text fusion is useful but still sensitive to session differences.")
    add_heading(doc, "6.5. Literature comparison", 2)
    add_table(doc, ["Model", "Input", "Split", "Reported result", "Comment"], [
        ("emotion2vec linear", "audio", "LOSO / LOSpeaker", "WA around 71.79-72.94; starred variants higher", "Closest pretrained speech emotion reference [8]"),
        ("FT-w2v2 P-TAPT", "raw waveform", "LOSO", "UA 74.3", "Strong SSL fine-tuning baseline [15]"),
        ("CA-MSER", "MFCC + spectrogram + wav2vec2", "LOSO / LOSpeaker", "LOSO WA/UA 69.80/71.05; LOSpeaker WA/UA 71.64/72.70", "Closest co-attention reference [9]"),
        ("DST", "WavLM", "LOSO", "WA 71.8, UA 73.6", "Transformer reference [12]"),
        ("03D expert fusion (ours)", "06D acoustic + transcript expert", "5-fold / 10-fold", "5-fold WA/UAR 72.90/73.56; 10-fold WA/UAR 74.02/74.64", "Best current implemented result"),
    ], "Comparison with literature baselines", font_size=6.6)
    add_heading(doc, "6.6. Discussion of results", 2)
    add_para(doc, "The current result trend supports the decision to keep transcript information as an auxiliary signal. The transcript-only 03C branch is already competitive, reaching 67.81% UAR on 5-fold and 69.47% UAR on 10-fold. However, 03D expert fusion is stronger than either branch alone, suggesting that the acoustic and text experts capture complementary cues.")
    add_para(doc, "Dominance remains the weakest VAD dimension, which is also common in affective speech regression because dominance is less directly observable from short utterances than arousal or valence. The planned ablation table is necessary because fusion can help or hurt depending on whether the acoustic branch is strong enough.")

    add_heading(doc, "CHAPTER 7: USER INTERFACE AND DEMONSTRATION")
    add_para(doc, "The demo is implemented in the 07_Web_Demo folder. It accepts audio upload or microphone recording and follows the same academic rule as the test pipeline: preprocess the waveform, run feature extraction, apply saved model artifacts, and return predictions without fitting new scalers on the live sample.")
    add_heading(doc, "7.1. Demo outputs", 2)
    add_bullets(doc, [
        "Emotion probability bars and final emotion label.",
        "Valence, arousal, and dominance radar/metric values.",
        "Waveform with transcript/timestamp markers.",
        "Log-Mel spectrogram and frame-level feature trend visualization.",
        "Feature evidence table showing X_temporal, X_spectral, X_stats, X_e2v, and transcript information.",
        "HTML report export that saves audio, feature images, transcript, prediction, and raw JSON.",
        "Queue comparison that compares waveform, log-Mel, emotion, confidence, VAD, RMS, ZCR, pitch, pause ratio, and centroid.",
    ])
    add_heading(doc, "7.2. Limitations and future work", 2)
    add_bullets(doc, [
        "The model is trained on acted IEMOCAP emotional speech, so real-world spontaneous speech may require domain adaptation.",
        "The current report contains placeholder slots for ablation modes that still need to be executed: acoustic_only, acoustic_text_concat, and acoustic_text_bridge without RMM.",
        "The web demo currently compares queue items inside the running session. Importing old exported reports for later comparison can be added as a follow-up.",
        "Dominance prediction is weaker than valence/arousal and may need better prosody modeling or additional supervision.",
        "End-to-end fine-tuning of speech and text backbones may improve accuracy but increases compute cost and fold-leakage risk if not carefully isolated per fold.",
    ])

    add_heading(doc, "CHAPTER 8: CONCLUSION")
    add_para(doc, "This project turns the midterm 06D acoustic SER prototype into a fuller research-style speech processing system. The final pipeline uses IEMOCAP, strict speaker/session-independent splits, multi-branch acoustic features, transcript modeling, bridge fusion, and two-head multi-task outputs for emotion and VAD. The strongest current implemented result is the 03D expert fusion model, reaching 72.90% WA / 73.56% UAR on 5-fold and 74.02% WA / 74.64% UAR on 10-fold, with CCC mean around 0.65-0.66.")

    add_heading(doc, "REFERENCES")
    for i, (authors, title, venue, url) in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f'[{i}] {authors}. "{title}." {venue}. {url}')
        set_run_font(r, size=9)

    doc.save(OUT_DOC)
    print(OUT_DOC)


if __name__ == "__main__":
    build()
