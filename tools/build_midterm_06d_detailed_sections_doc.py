from pathlib import Path
import csv
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Midterm_Param"
ASSET_DIR = OUT_DIR / "06D_midterm_assets"
REPORT_DIR = ROOT / "06_Advanced_Model" / "outputs" / "06D_Emotion2Vec_CoAttention_outputs_package" / "reports"

DOCX_PATH = OUT_DIR / "Midterm_Param_Sections_1.5_3.1_3.2_3.3_06D_Detailed.docx"
ZIP_PATH = ASSET_DIR / "Midterm_Param_06D_Detailed_Doc_Package.zip"

ARCH_IMG = ASSET_DIR / "06D_architecture_landscape_final.png"
FEATURE_IMG = ASSET_DIR / "savee_feature_example.png"
CURVES_IMG = ASSET_DIR / "06D_training_curves.png"
LEADERBOARD_IMG = ASSET_DIR / "06D_macro_f1_leaderboard.png"
CM_RANDOM_IMG = ASSET_DIR / "06D_confusion_combined_random_stacking_full.png"
CM_STRICT_IMG = ASSET_DIR / "06D_confusion_combined_strict_no_tess_stacking_full.png"


BLUE = RGBColor(31, 78, 121)
TEAL = RGBColor(0, 128, 128)
DARK = RGBColor(16, 42, 67)
MUTED = RGBColor(82, 96, 109)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=DARK, size=font_size)
        set_cell_shading(hdr[i], "E8EEF5")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = BLUE if level <= 2 else DARK
        r.font.name = "Calibri"
    return p


def add_p(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = DARK
        r1.font.size = Pt(11)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.size = Pt(10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.size = Pt(10.5)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(38, 50, 56)
    return p


def add_image(doc, path, caption, width=6.4):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def read_csv(path):
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def pct(x):
    try:
        return f"{float(x) * 100:.2f}%"
    except Exception:
        return str(x)


def build_doc():
    metrics = read_csv(REPORT_DIR / "06D_emotion2vec_coattention_metrics.csv")
    per_dataset = read_csv(REPORT_DIR / "06D_emotion2vec_coattention_per_dataset.csv")
    refs = read_csv(REPORT_DIR / "06D_reference_model_comparison.csv")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Phân tích chi tiết kiến trúc đề xuất 06D Emotion2Vec Co-Attention SER")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DARK
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Các mục Midterm Param: 1.5, 3.1, 3.2, 3.3")
    sr.font.size = Pt(11)
    sr.font.color.rgb = MUTED

    add_p(
        doc,
        "Tài liệu này viết lại phần phân tích mô hình theo hướng 06D hiện tại. Nội dung bám theo roadmap của project: giai đoạn hiện tại tập trung vào Speech Emotion Recognition, còn hướng phát triển dài hạn là hệ thống feedback giọng nói cho thuyết trình. Vì vậy, phần mô hình không chỉ mô tả accuracy mà còn giải thích vì sao các đặc trưng, branch và cơ chế fusion có liên quan tới việc phân tích cảm xúc, độ tự tin và xu hướng cảm xúc trong giọng nói."
    )

    add_image(doc, ARCH_IMG, "Hình 1. Kiến trúc 06D Emotion2Vec Co-Attention SER đã được chuẩn hóa lại để dùng trong báo cáo.", width=7.0)

    add_heading(doc, "1.5. Select the proposed solution for the project and justify the selection", 1)
    add_p(
        doc,
        "Giải pháp được chọn cho giai đoạn hiện tại là mô hình 06D Emotion2Vec Co-Attention SER, thuộc nhóm multi-representation Speech Emotion Recognition. Mô hình nhận audio tiếng nói 16 kHz làm đầu vào và dự đoán xác suất cho 6 cảm xúc chung: neutral, happy, sad, angry, fear và disgust."
    )
    add_p(
        doc,
        "Lý do chọn hướng multi-representation là vì cảm xúc trong giọng nói không chỉ nằm ở một loại tín hiệu. Một người nói có thể thể hiện cảm xúc qua năng lượng giọng, màu giọng, tốc độ thay đổi phổ âm, vùng năng lượng trên spectrogram, nhịp nói, hoặc các cue cảm xúc trừu tượng hơn đã được mô hình pretrained học trước. Nếu chỉ dùng một nhóm đặc trưng như MFCC hoặc chỉ dùng một mô hình CNN đơn lẻ thì hệ thống có thể bỏ sót nhiều cue quan trọng."
    )
    add_p(
        doc,
        "Do đó, 06D được thiết kế thành bốn nhánh biểu diễn. Mỗi nhánh nhìn cùng một audio ở một góc khác nhau, sau đó mô hình mới kết hợp chúng ở tầng fusion. Đây là điểm khác với baseline truyền thống ở phần 04, nơi mô hình chủ yếu dùng MFCC/statistical features với SVM, Random Forest hoặc Logistic Regression."
    )

    add_table(
        doc,
        ["Thành phần", "Input chính", "Vai trò trong mô hình", "Lý do chọn"],
        [
            ["Branch A - Temporal acoustic", "MFCC + delta + delta-delta + RMS/ZCR/spectral sequence", "Học biến đổi cảm xúc theo thời gian bằng 1D-CNN, BiLSTM/GRU và attention.", "Cảm xúc thường nằm trong dynamics của giọng nói, không chỉ trong một frame tĩnh."],
            ["Branch B - Spectrogram", "Log-Mel + delta log-Mel + delta-delta log-Mel", "Học pattern thời gian - tần số bằng 2D-CNN, residual CNN và SE attention.", "Log-Mel giữ rõ cấu trúc time-frequency, phù hợp với CNN."],
            ["Branch C - Emotion pretrained", "Raw waveform 16 kHz", "Dùng frozen emotion2vec để lấy embedding cảm xúc đã học trước, sau đó train adapter nhỏ.", "Tận dụng pretrained representation mà không cần GPU mạnh để fine-tune toàn bộ mô hình."],
            ["Branch D - Statistical", "Handcrafted statistical vector", "Dùng StandardScaler, Stats MLP và RBF-SVM để tạo p_stats.", "Feature engineering cổ điển vẫn có giá trị trong SER và bổ sung thông tin toàn clip."],
            ["Fusion", "z_temporal, z_spectral, z_e2v, p_stats", "Co-attention tạo p_deep, sau đó stacking kết hợp p_deep và p_stats.", "Giúp mô hình kết hợp nhiều nguồn thông tin thay vì phụ thuộc một branch."],
        ],
        widths=[1.55, 1.65, 2.2, 2.0],
        font_size=8,
    )

    add_p(
        doc,
        "Giải pháp này cũng phù hợp với điều kiện tài nguyên của đồ án sinh viên. Thay vì train một pretrained model lớn end-to-end, project dùng emotion2vec ở chế độ frozen. Điều này giúp giảm thời gian train, giảm yêu cầu GPU, và vẫn giữ được giá trị học thuật vì mô hình có cơ chế kết hợp nhiều biểu diễn, co-attention và stacking."
    )
    add_p(
        doc,
        "Về mặt định hướng ứng dụng, mô hình 06D là nền tảng cho hệ thống feedback giọng nói. Đầu ra không chỉ có nhãn cảm xúc, mà còn có vector xác suất p_final[6] và confidence. Về sau, hệ thống có thể mở rộng thành phân tích xu hướng cảm xúc theo từng đoạn nói, phát hiện đoạn nói thiếu tự tin, căng thẳng hoặc quá đơn điệu, từ đó đưa ra phản hồi cho người thuyết trình."
    )

    add_heading(doc, "3.1. Present the architecture or workflow of the proposed solution", 1)
    add_p(
        doc,
        "Workflow của 06D bắt đầu từ audio đầu vào 16 kHz. Audio được resample, chuyển mono nếu cần, chuẩn hóa biên độ và crop/pad về độ dài cố định. Một điểm quan trọng là dữ liệu phải được chia train/validation/test trước khi augmentation hoặc chọn tham số, nhằm tránh data leakage. Sau bước tiền xử lý, cùng một audio được chuyển thành bốn biểu diễn khác nhau."
    )
    add_numbered(
        doc,
        [
            "Input audio: một utterance hoặc segment giọng nói ở 16 kHz.",
            "Preprocess: resample, mono, normalize, crop/pad theo duration của notebook.",
            "Split protocol: combined_random, combined_strict_no_tess và các thí nghiệm single-dataset nếu cần.",
            "Feature router: tạo temporal sequence, log-Mel representation, emotion2vec embedding và statistical vector.",
            "Branch modeling: mỗi branch học một loại cue cảm xúc khác nhau.",
            "Fusion: co-attention kết hợp các embedding sâu; stacking kết hợp p_deep với p_stats.",
            "Output: vector xác suất 6 lớp, nhãn dự đoán và confidence."
        ],
    )

    add_heading(doc, "3.1.1. Các đặc trưng tiếng nói được dùng trong mô hình", 2)
    add_p(
        doc,
        "Các đặc trưng dưới đây được dùng vì mỗi nhóm thể hiện một mặt khác nhau của cảm xúc. Bảng này được viết theo đúng logic của notebook 06D: temporal branch dùng đặc trưng theo frame, spectrogram branch dùng biểu diễn time-frequency, emotion2vec branch dùng waveform, còn statistical branch dùng vector thống kê cố định."
    )
    add_table(
        doc,
        ["Đặc trưng", "Là gì?", "Nó thể hiện điều gì?", "Vì sao dùng trong branch?"],
        [
            ["Waveform", "Dạng sóng âm theo thời gian, biểu diễn biên độ tín hiệu tại từng thời điểm.", "Cho biết cường độ, khoảng nghỉ, nhịp nói và biến động biên độ thô.", "Dùng ở bước tiền xử lý và Branch C vì emotion2vec nhận waveform/raw speech làm đầu vào."],
            ["MFCC", "Mel-Frequency Cepstral Coefficients, mô tả envelope phổ âm trên thang Mel gần với cảm nhận tai người.", "Thể hiện timbre/màu giọng và hình dạng phổ âm.", "Dùng trong Branch A vì tạo chuỗi frame theo thời gian; dùng trong Branch D để lấy thống kê toàn clip."],
            ["Delta MFCC", "Đạo hàm bậc 1 của MFCC theo thời gian.", "Thể hiện tốc độ thay đổi màu giọng giữa các frame.", "Bổ sung dynamics cho Branch A, giúp mô hình thấy cảm xúc thay đổi ra sao."],
            ["Delta-delta MFCC", "Đạo hàm bậc 2 của MFCC theo thời gian.", "Thể hiện gia tốc thay đổi phổ âm.", "Giúp Branch A bắt các biến động nhanh trong giọng nói."],
            ["Log-Mel spectrogram", "Bản đồ năng lượng theo thời gian và dải tần Mel, sau đó lấy log để nén dải giá trị.", "Cho thấy vùng năng lượng cao/thấp, formant, nhịp phát âm và texture phổ.", "Dùng trong Branch B vì CNN xử lý tốt biểu diễn dạng ảnh time-frequency."],
            ["RMS energy", "Root Mean Square energy, đo độ lớn/năng lượng của tín hiệu theo frame.", "Angry/happy thường có năng lượng cao hơn sad/neutral trong nhiều trường hợp.", "Dùng trong Branch A như chuỗi năng lượng và Branch D như thống kê toàn clip."],
            ["ZCR", "Zero-Crossing Rate, số lần tín hiệu đổi dấu trong một frame.", "Liên quan tới độ sắc, nhiễu, hoặc thành phần tần số cao.", "Bổ sung low-level cue cho Branch A và Branch D."],
            ["Spectral centroid", "Tâm khối năng lượng phổ.", "Centroid cao gợi âm sáng/sắc; centroid thấp gợi âm trầm hơn.", "Bổ sung thông tin độ sáng phổ cho Branch A và Branch D."],
            ["Bandwidth / Rolloff", "Độ rộng phổ và ngưỡng tần số chứa phần lớn năng lượng.", "Cho biết năng lượng tập trung hẹp hay trải rộng trên nhiều dải tần.", "Dùng để mô tả chất lượng âm và sự phân bố phổ."],
            ["Spectral contrast", "Độ chênh giữa đỉnh và đáy năng lượng ở các dải tần.", "Phản ánh cấu trúc phổ và độ rõ của vùng tần số.", "Dùng trong feature engineering vì SER cần cue về chất giọng và cường độ phổ."],
            ["Chroma", "Năng lượng gom theo 12 lớp cao độ.", "Bổ sung thông tin pitch-class/harmonic texture.", "Dùng trong Branch D như đặc trưng thống kê phụ, không phải feature chính của speech emotion."],
            ["Statistical vector", "Các thống kê mean, std, min, max, percentile, entropy trên feature.", "Tóm tắt toàn bộ clip thành vector cố định.", "Cần cho SVM/RBF-SVM vì SVM cần input có độ dài cố định."],
        ],
        widths=[1.2, 2.0, 2.1, 2.2],
        font_size=7.5,
    )
    add_image(doc, FEATURE_IMG, "Hình 2. Minh họa waveform, spectrogram, MFCC và log-Mel từ mẫu audio trong SAVEE.", width=6.6)

    add_heading(doc, "3.1.2. Branch A - Temporal Acoustic Model", 2)
    add_p(
        doc,
        "Branch A xử lý chuỗi đặc trưng theo thời gian. Đây là nhánh quan trọng vì cảm xúc trong giọng nói không chỉ nằm ở một frame riêng lẻ, mà nằm ở cách năng lượng, màu giọng và phổ âm thay đổi trong suốt câu nói."
    )
    add_p(
        doc,
        "Input của Branch A là tensor có dạng [B, 132, T]. Trong đó B là batch size, 132 là số đặc trưng trên mỗi frame, và T là số frame theo thời gian. 132 đặc trưng gồm 40 MFCC, 40 delta MFCC, 40 delta-delta MFCC, RMS, ZCR, spectral centroid, spectral bandwidth, spectral rolloff và 7 spectral contrast."
    )
    add_code(
        doc,
        "temporal = concatenate([\n"
        "    mfcc, delta_mfcc, delta2_mfcc,\n"
        "    rms, zcr, centroid, bandwidth, rolloff, contrast\n"
        "], axis=feature_axis)  # [132, T]\n"
        "model_input = temporal  # Conv1D scans along time axis"
    )
    add_table(
        doc,
        ["Khối", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng?"],
        [
            ["1D-CNN", "Conv1D trượt kernel theo trục thời gian để phát hiện pattern cục bộ như năng lượng tăng nhanh, MFCC đổi mạnh hoặc giọng sắc hơn.", "[B, 132, T] -> temporal feature maps", "Nhẹ hơn 2D-CNN và phù hợp với chuỗi MFCC/low-level acoustic features."],
            ["BiLSTM/BiGRU", "Đọc chuỗi theo hai hướng: forward dùng quá khứ -> hiện tại, backward dùng tương lai -> hiện tại.", "feature maps -> contextual sequence", "Cảm xúc phụ thuộc ngữ cảnh nhiều frame, không chỉ một frame."],
            ["Attention pooling", "Tính trọng số quan trọng cho từng frame rồi lấy tổng có trọng số.", "sequence -> z_temporal", "Không phải frame nào cũng chứa cảm xúc rõ; attention tập trung vào đoạn giàu thông tin."],
            ["Classifier/embedding head", "LayerNorm, Dropout, Linear và activation phi tuyến để tạo embedding.", "context vector -> z_temporal 192-D", "Tạo biểu diễn cố định để đưa vào fusion."],
        ],
        widths=[1.2, 2.8, 1.5, 2.0],
        font_size=8,
    )

    add_heading(doc, "3.1.3. Branch B - Spectrogram CNN-SE Model", 2)
    add_p(
        doc,
        "Branch B khai thác log-Mel spectrogram. Khác với MFCC, log-Mel giữ cấu trúc thời gian - tần số rõ ràng: trục ngang là thời gian, trục dọc là Mel bin, màu/pixel là năng lượng. Điều này giúp CNN học được các vùng phổ giống như học pattern trên ảnh."
    )
    add_p(
        doc,
        "Input của Branch B có dạng [B, 3, 96, T]. Ba kênh là log-Mel, delta log-Mel và delta-delta log-Mel. Log-Mel cho biết năng lượng hiện tại; delta log-Mel cho biết năng lượng đang tăng/giảm nhanh thế nào; delta-delta log-Mel cho biết gia tốc thay đổi đó."
    )
    add_code(
        doc,
        "mel = melspectrogram(audio, n_mels=96)\n"
        "logmel = power_to_db(mel)\n"
        "d_logmel = delta(logmel)\n"
        "d2_logmel = delta(logmel, order=2)\n"
        "spectral = stack([logmel, d_logmel, d2_logmel], axis=channel_axis)  # [3, 96, T]"
    )
    add_table(
        doc,
        ["Khối", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng?"],
        [
            ["2D-CNN", "Dùng kernel 2D quét trên trục thời gian và Mel frequency để bắt pattern cục bộ.", "[B, 3, 96, T] -> feature maps", "Phù hợp với biểu diễn time-frequency dạng ảnh."],
            ["Residual CNN", "Thêm skip connection để học sâu hơn nhưng tránh mất tín hiệu gốc.", "feature maps -> deeper feature maps", "Ổn định quá trình train khi mô hình có nhiều lớp."],
            ["SE attention", "Squeeze-and-Excitation tính trọng số cho từng channel rồi nhân lại vào feature map.", "feature maps -> re-weighted maps", "Channel chứa cue cảm xúc mạnh sẽ được nhấn hơn."],
            ["Spectral embedding head", "Pooling và projection để tạo vector cố định.", "maps -> z_spectral 192-D", "Đưa biểu diễn phổ vào fusion với các branch khác."],
        ],
        widths=[1.2, 2.7, 1.6, 2.0],
        font_size=8,
    )

    add_heading(doc, "3.1.4. Branch C - Emotion Pretrained Branch", 2)
    add_p(
        doc,
        "Branch C dùng raw waveform 16 kHz và mô hình pretrained emotion2vec. Trong 06D, emotion2vec được dùng ở chế độ frozen: trọng số pretrained không bị cập nhật trong quá trình train chính. Chỉ adapter MLP phía sau được train để ánh xạ embedding về không gian 6 nhãn cảm xúc của project."
    )
    add_p(
        doc,
        "Lý do không fine-tune toàn bộ pretrained model là vì project bị giới hạn tài nguyên Kaggle/Colab và mục tiêu demo thực tế cần mô hình không quá phụ thuộc vào một tập dữ liệu nhỏ. Frozen embedding giúp tận dụng tri thức đã học trên dữ liệu lớn, đồng thời giảm nguy cơ overfit khi dataset SER hiện tại không quá lớn."
    )
    add_code(
        doc,
        "waveform -> emotion2vec_base (frozen)\n"
        "hidden_states -> mean_pooling -> z_raw\n"
        "z_raw -> AdapterMLP(Linear + GELU + Dropout) -> z_e2v / p_e2v"
    )
    add_table(
        doc,
        ["Khối", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng?"],
        [
            ["emotion2vec frozen encoder", "Biến waveform thành chuỗi hidden states giàu thông tin cảm xúc.", "waveform -> hidden states", "Tận dụng pretrained speech emotion representation."],
            ["Mean pooling", "Lấy trung bình hidden states theo thời gian để có vector toàn utterance.", "hidden states -> utterance vector", "Classifier cần vector cố định cho mỗi audio."],
            ["Adapter MLP", "Linear + activation + dropout để đưa embedding pretrained về không gian task.", "utterance vector -> z_e2v / p_e2v", "Train nhẹ, ít tài nguyên, giảm overfit so với fine-tune toàn bộ."],
        ],
        widths=[1.5, 2.7, 1.6, 1.9],
        font_size=8,
    )

    add_heading(doc, "3.1.5. Branch D - Statistical Acoustic RBF-SVM", 2)
    add_p(
        doc,
        "Branch D là nhánh machine learning cổ điển. Nó không giữ nguyên chuỗi frame mà tóm tắt toàn bộ clip thành vector thống kê cố định. Nhánh này được giữ lại vì nhiều hướng SER dựa trên feature engineering cho thấy các đặc trưng thủ công vẫn có giá trị, đặc biệt khi dữ liệu ít hoặc domain khác nhau."
    )
    add_p(
        doc,
        "Nếu đưa nguyên chuỗi feature vào SVM, độ dài vector sẽ phụ thuộc vào độ dài audio và dễ overfit. Vì vậy Branch D lấy các thống kê như mean, std, min, max, percentile và entropy trên từng nhóm feature. Kết quả là một vector có độ dài cố định cho mỗi mẫu audio."
    )
    add_code(
        doc,
        "stat_vec = concatenate([\n"
        "    stats(mfcc), stats(delta), stats(delta2),\n"
        "    stats(chroma_stft), stats(chroma_cqt), stats(chroma_cens),\n"
        "    stats(centroid), stats(bandwidth), stats(rolloff), stats(contrast),\n"
        "    stats(rms), stats(zcr), stats(energy), stats(entropy)\n"
        "])\n"
        "p_stats = StandardScaler() + RBF-SVM(probability=True, class_weight='balanced')"
    )
    add_table(
        doc,
        ["Khối", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng?"],
        [
            ["Statistical extraction", "Tính mean, std, min, max, percentile, entropy trên các feature frame-level.", "variable-length sequence -> fixed vector", "Tạo input cố định cho SVM và giảm phụ thuộc độ dài audio."],
            ["StandardScaler", "Chuẩn hóa mỗi feature về thang đo tương đương.", "X_stats -> scaled X_stats", "SVM dựa trên khoảng cách/kernel nên rất nhạy với scale."],
            ["RBF-SVM", "Dùng kernel phi tuyến để tìm biên quyết định trong không gian đặc trưng.", "scaled vector -> class probability", "Emotion không tách tuyến tính rõ trong không gian feature thủ công."],
            ["Probability output", "SVC(probability=True) trả xác suất từng lớp.", "decision score -> p_stats[6]", "Cần p_stats để stacking với p_deep."],
        ],
        widths=[1.4, 2.5, 1.7, 2.0],
        font_size=8,
    )

    add_heading(doc, "3.1.6. Fusion: Co-attention and validation-weighted stacking", 2)
    add_p(
        doc,
        "Sau khi Branch A, B, C tạo embedding, mô hình không cộng hoặc nối đơn giản ngay lập tức. 06D dùng co-attention để cho phép embedding pretrained từ emotion2vec hỗ trợ lựa chọn thông tin quan trọng từ temporal và spectrogram embedding. Nói cách khác, z_e2v đóng vai trò cue cảm xúc đã học trước, còn z_temporal và z_spectral cung cấp thông tin chi tiết từ feature engineering hiện tại."
    )
    add_code(
        doc,
        "z_temporal + z_spectral + z_e2v\n"
        "    -> co-attention\n"
        "    -> deep classifier\n"
        "    -> p_deep[6]\n\n"
        "p_final = validation_weighted_stacking(p_deep, p_stats, p_e2v_optional)"
    )
    add_p(
        doc,
        "Stacking được thực hiện trên xác suất, không phải trên nhãn cứng. Điều này giữ lại mức độ tự tin của từng branch. Ví dụ nếu deep model rất tự tin ở angry nhưng SVM nghiêng về neutral, stacking có thể học cách cân bằng dựa trên validation split. Test split chỉ được dùng để báo cáo cuối, không dùng để chọn trọng số."
    )

    add_heading(doc, "3.2. Implement the selected speech processing model, algorithm, or method", 1)
    add_p(
        doc,
        "Phần triển khai nằm trong notebook 06D Emotion2Vec Co-Attention Full SER. Dữ liệu đầu vào chính là thư mục ser_processed đã được tạo từ các notebook 01/02. Thư mục này chứa metadata, audio đã chuẩn hóa và các thông tin label/dataset/speaker cần thiết cho việc chia dữ liệu."
    )
    add_table(
        doc,
        ["Bước triển khai", "Nội dung", "Ý nghĩa kỹ thuật"],
        [
            ["1. Load metadata", "Đọc path audio, emotion label, dataset, speaker_id.", "Cần cho train/test split, speaker-aware split và per-dataset analysis."],
            ["2. Feature extraction", "Tạo X_temporal, X_spectral, X_e2v, X_stats.", "Chuẩn bị input cho bốn branch của 06D."],
            ["3. Split protocol", "combined_random và combined_strict_no_tess.", "So sánh hiệu năng dễ hơn và khả năng generalization khó hơn."],
            ["4. Train deep model", "Train Branch A/B/C + co-attention + classifier.", "Học p_deep từ multi-representation embedding."],
            ["5. Train stats SVM", "Train StandardScaler + RBF-SVM trên X_stats.", "Tạo p_stats từ feature engineering cổ điển."],
            ["6. Stacking", "Kết hợp p_deep, p_stats và optional emotion2vec probability.", "Tận dụng điểm mạnh của nhiều mô hình."],
            ["7. Evaluation", "Accuracy, macro-F1, weighted-F1, UAR, confusion matrix, per-dataset metrics.", "Đánh giá tổng thể và đánh giá theo từng domain dataset."],
        ],
        widths=[1.4, 2.7, 3.0],
        font_size=8,
    )
    add_p(
        doc,
        "Trong quá trình train, validation macro-F1 được dùng để chọn best epoch. Macro-F1 được ưu tiên vì bài toán có nhiều lớp cảm xúc và phân bố giữa các dataset không hoàn toàn cân bằng. Accuracy chỉ cho biết tỷ lệ đúng tổng thể; macro-F1 cho biết mô hình có học đều các lớp hay không."
    )
    add_p(
        doc,
        "Về mặt leakage control, augmentation hoặc các xử lý có tính thay đổi dữ liệu phải diễn ra sau khi split. Validation và test không được augment. Điều này đảm bảo test set vẫn là dữ liệu thật để đánh giá mô hình, tránh trường hợp mô hình nhìn thấy biến thể gần giống của test sample trong training."
    )
    add_p(
        doc,
        "Về mặt implementation, 06D chưa phải hệ thống feedback hoàn chỉnh. Nó là lõi nhận diện cảm xúc. Tuy nhiên, cấu trúc output đã phù hợp để mở rộng: p_final[6] có thể dùng để lấy confidence; nếu chạy theo từng đoạn audio, ta có thể tạo emotion timeline để feedback về xu hướng cảm xúc của bài nói."
    )

    add_heading(doc, "3.3. Provide preliminary experimental results or initial training results to demonstrate the feasibility of the proposed solution", 1)
    add_p(
        doc,
        "Kết quả sơ bộ được lấy từ output thật của notebook 06D. Notebook báo cáo hai protocol chính: combined_random và combined_strict_no_tess. combined_random là chia ngẫu nhiên trên tập gộp, thường dễ hơn vì speaker/domain có thể xuất hiện ở cả train và test. combined_strict_no_tess là protocol khó hơn, dùng để kiểm tra generalization khi speaker/domain ít bị trùng lặp hơn."
    )

    metric_rows = []
    for row in metrics:
        if row.get("split") == "test":
            metric_rows.append([
                row.get("protocol", ""),
                row.get("model", ""),
                row.get("n_samples", ""),
                pct(row.get("accuracy", "")),
                pct(row.get("macro_f1", "")),
                pct(row.get("uar", "")),
                row.get("best_epoch", "") or "-",
            ])
    add_table(
        doc,
        ["Protocol", "Model", "Test n", "Accuracy", "Macro-F1", "UAR", "Best epoch"],
        metric_rows,
        widths=[1.35, 1.65, 0.65, 0.85, 0.85, 0.85, 0.75],
        font_size=7.5,
    )

    add_p(
        doc,
        "Kết quả chính cho thấy stacking_full đạt khoảng 80.39% accuracy và 80.51% macro-F1 trên combined_random. Trên combined_strict_no_tess, stacking_full đạt khoảng 69.59% accuracy và 69.60% macro-F1. Khoảng cách này chứng minh mô hình học được tín hiệu cảm xúc, nhưng vẫn bị ảnh hưởng bởi speaker/domain shift."
    )
    add_image(doc, LEADERBOARD_IMG, "Hình 3. Macro-F1 leaderboard của các mô hình/nhánh trong notebook 06D.", width=6.5)
    add_image(doc, CURVES_IMG, "Hình 4. Training curves cho thấy train macro-F1 tăng cao hơn validation macro-F1, phản ánh nguy cơ overfit.", width=6.5)

    add_heading(doc, "3.3.1. Phân tích theo từng dataset trong test set", 2)
    per_rows = []
    for row in per_dataset:
        if row.get("model") == "stacking_full":
            per_rows.append([
                row.get("protocol", ""),
                row.get("dataset", ""),
                row.get("n_samples", ""),
                pct(row.get("accuracy", "")),
                pct(row.get("macro_f1", "")),
                pct(row.get("uar", "")),
            ])
    add_table(
        doc,
        ["Protocol", "Dataset", "n", "Accuracy", "Macro-F1", "UAR"],
        per_rows,
        widths=[1.55, 1.0, 0.6, 1.0, 1.0, 1.0],
        font_size=8,
    )
    add_p(
        doc,
        "Bảng per-dataset cho thấy TESS dễ hơn nhiều trong combined_random, trong khi CREMA-D thường khó hơn do số lượng speaker lớn, phong cách nói đa dạng và nhiễu domain cao hơn. Ở strict protocol, kết quả trên RAVDESS và SAVEE dao động do số mẫu test nhỏ hơn, còn CREMA-D phản ánh rõ hơn độ khó của generalization."
    )
    add_image(doc, CM_RANDOM_IMG, "Hình 5. Confusion matrix của stacking_full trên combined_random.", width=5.8)
    add_image(doc, CM_STRICT_IMG, "Hình 6. Confusion matrix của stacking_full trên combined_strict_no_tess.", width=5.8)

    add_heading(doc, "3.3.2. So sánh sơ bộ với các hướng nghiên cứu tham khảo", 2)
    ref_rows = []
    for row in refs:
        ref_rows.append([
            row.get("model", ""),
            row.get("protocol", ""),
            row.get("reported_accuracy_text", ""),
            row.get("main_idea", ""),
            row.get("link", ""),
        ])
    add_table(
        doc,
        ["Reference", "Protocol / split", "Reported result", "Main idea", "Link"],
        ref_rows,
        widths=[1.45, 1.45, 1.25, 2.1, 1.5],
        font_size=7,
    )
    add_p(
        doc,
        "Các bài báo tham khảo có thể đạt accuracy rất cao, nhưng cần đọc cùng protocol. Nhiều kết quả trên 90% đến từ single-dataset hoặc random split, nơi train/test có thể cùng điều kiện thu âm, cùng dataset style hoặc cùng speaker distribution. Vì vậy, không nên so trực tiếp kết quả đó với strict speaker/domain-aware protocol của project."
    )
    add_p(
        doc,
        "So với Ahmed et al., 06D có điểm giống là sử dụng ensemble/fusion và recurrent modeling, nhưng 06D dùng thêm emotion2vec pretrained branch và co-attention. So với hướng feature engineering mạnh như Chowdhury et al., 06D còn yếu hơn ở statistical branch: stats_rbf_svm trong kết quả hiện tại chưa mạnh, đặc biệt ở strict protocol. Đây là lý do các notebook sau như 06E thử richer statistics và augmentation sâu hơn."
    )

    add_heading(doc, "3.3.3. Nhận xét feasibility và hạn chế hiện tại", 2)
    add_bullets(
        doc,
        [
            "Mô hình khả thi vì stacking_full đạt trên 80% accuracy/macro-F1 ở combined_random, cao hơn từng nhánh đơn lẻ như stats_rbf_svm hoặc emotion2vec_logreg.",
            "Mô hình chưa giải quyết triệt để generalization vì strict accuracy/macro-F1 còn quanh 69.6%. Điều này cho thấy speaker/domain signal vẫn ảnh hưởng mạnh.",
            "Branch emotion2vec_mlp hoạt động khá tốt trên strict, cho thấy pretrained emotion representation có giá trị.",
            "Branch stats_rbf_svm yếu ở strict, nghĩa là feature engineering hiện tại chưa đủ mạnh hoặc RBF-SVM chưa được tối ưu tốt.",
            "Training curves cho thấy train macro-F1 tăng cao hơn validation macro-F1, phản ánh nguy cơ overfit vào speaker hoặc dataset style.",
        ],
    )
    add_p(
        doc,
        "Kết luận cho mục 3.3: 06D đủ khả thi để dùng làm proposed solution ở midterm vì đã có pipeline hoàn chỉnh từ audio -> feature extraction -> multi-branch modeling -> fusion -> evaluation. Đồng thời, kết quả cũng chỉ ra hướng future work rõ ràng: tăng độ mạnh của statistical features, cải thiện domain generalization, thử adapter/domain-aware architecture, và đánh giá thêm single-dataset experiments để so sánh công bằng hơn với các bài báo."
    )

    doc.add_page_break()
    add_heading(doc, "Phụ lục: mapping nhanh giữa kiến trúc và code/notebook", 1)
    add_table(
        doc,
        ["Thành phần trong sơ đồ", "Tên logic trong notebook", "Input", "Output"],
        [
            ["Branch A", "Temporal acoustic branch / CNN + BiLSTM/GRU + Attention", "X_temporal [B, 132, T]", "z_temporal, p_temporal nếu dùng head riêng"],
            ["Branch B", "Spectrogram CNN-SE branch", "X_spectral [B, 3, 96, T]", "z_spectral"],
            ["Branch C", "Frozen emotion2vec + adapter MLP", "raw waveform / X_e2v embedding", "z_e2v, optional p_e2v"],
            ["Branch D", "Rich/statistical features + RBF-SVM", "X_stats [B, D]", "p_stats [B, 6]"],
            ["Co-attention", "embedding fusion", "z_temporal, z_spectral, z_e2v", "z_fused / p_deep"],
            ["Stacking", "validation-weighted ensemble", "p_deep, p_stats, optional p_e2v", "p_final[6]"],
        ],
        widths=[1.4, 2.3, 1.8, 1.7],
        font_size=8,
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


def make_zip(docx_path):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in [
            docx_path,
            ARCH_IMG,
            FEATURE_IMG,
            CURVES_IMG,
            LEADERBOARD_IMG,
            CM_RANDOM_IMG,
            CM_STRICT_IMG,
        ]:
            if path.exists():
                zf.write(path, path.relative_to(ROOT))
    return ZIP_PATH


def main():
    docx = build_doc()
    zipped = make_zip(docx)
    print("DOCX", docx)
    print("ZIP", zipped)


if __name__ == "__main__":
    main()
