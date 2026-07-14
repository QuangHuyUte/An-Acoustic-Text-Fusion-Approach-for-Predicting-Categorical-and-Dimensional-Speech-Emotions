from pathlib import Path
import csv
import zipfile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Midterm_Param"
ASSET_DIR = OUT_DIR / "06D_midterm_assets"
EVIDENCE_DIR = ASSET_DIR / "feature_evidence_6_emotions"
REPORT_DIR = ROOT / "06_Advanced_Model" / "outputs" / "06D_Emotion2Vec_CoAttention_outputs_package" / "reports"
SINGLE_OUT_DIR = ROOT / "06_Advanced_Model" / "outputs" / "06D_Emotion2Vec_CoAttention_outputs_package" / "06D_Single_Dataset_Emotion2Vec_CoAttention_outputs_light_no_cache"
SINGLE_REPORT_DIR = SINGLE_OUT_DIR / "reports"
SINGLE_FIGURE_DIR = SINGLE_OUT_DIR / "figures"

DOCX_PATH = OUT_DIR / "Midterm_Param_06D_Full_Detailed_Feature_Block_Analysis.docx"
ALT_DOCX_PATH = OUT_DIR / "Midterm_Param_06D_Compact_Architecture_Walkthrough.docx"
ALT_DOCX_PATH_V2 = OUT_DIR / "Midterm_Param_06D_Compact_Architecture_Walkthrough_v2.docx"
ALT_DOCX_PATH_V3 = OUT_DIR / "Midterm_Param_06D_Compact_Architecture_Walkthrough_v3.docx"
ZIP_PATH = ASSET_DIR / "Midterm_Param_06D_Full_Detailed_Feature_Block_Analysis_Package.zip"

ARCH_IMG = ASSET_DIR / "06D_architecture_landscape_final.png"
EVIDENCE_WAVE_RMS = EVIDENCE_DIR / "feature_evidence_waveform_rms_6emotions.png"
EVIDENCE_MFCC = EVIDENCE_DIR / "feature_evidence_mfcc_delta_6emotions.png"
EVIDENCE_ZCR = EVIDENCE_DIR / "feature_evidence_zcr_6emotions.png"
EVIDENCE_SPECTRAL = EVIDENCE_DIR / "feature_evidence_spectral_shape_6emotions.png"
EVIDENCE_LOGMEL = EVIDENCE_DIR / "feature_evidence_logmel_contrast_6emotions.png"
BILSTM_IMG = ASSET_DIR / "bilstm_temporal_context_diagram.png"
DOMAIN_DATASET_IMG = ASSET_DIR / "domain_dataset_evidence.png"
CURVES_IMG = ASSET_DIR / "06D_training_curves.png"
LEADERBOARD_IMG = ASSET_DIR / "06D_macro_f1_leaderboard.png"
CM_RANDOM_IMG = ASSET_DIR / "06D_confusion_combined_random_stacking_full.png"
CM_STRICT_IMG = ASSET_DIR / "06D_confusion_combined_strict_no_tess_stacking_full.png"
CM_SINGLE_CREMA_IMG = SINGLE_FIGURE_DIR / "06D_confusion_single_CREMA-D_stacking_full.png"
CM_SINGLE_RAVDESS_IMG = SINGLE_FIGURE_DIR / "06D_confusion_single_RAVDESS_stacking_full.png"
CM_SINGLE_SAVEE_IMG = SINGLE_FIGURE_DIR / "06D_confusion_single_SAVEE_stacking_full.png"
CM_SINGLE_TESS_IMG = SINGLE_FIGURE_DIR / "06D_confusion_single_TESS_stacking_full.png"

INK = RGBColor(16, 42, 67)
BLUE = RGBColor(31, 78, 121)
MUTED = RGBColor(82, 96, 109)
GREEN = RGBColor(5, 150, 105)
SKIP_OLD_15_31_BLOCK = True
_SUPPRESS_OLD_15_31 = False


def suppress_old_15_31():
    return _SUPPRESS_OLD_15_31


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, font_size=8):
    if suppress_old_15_31():
        return None
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, INK, font_size)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    global _SUPPRESS_OLD_15_31
    if SKIP_OLD_15_31_BLOCK and text == "1.5 & 3.1. Proposed solution, research origin and architecture workflow":
        _SUPPRESS_OLD_15_31 = True
        return None
    if _SUPPRESS_OLD_15_31:
        if text == "3.2. Implement the selected speech processing model, algorithm, or method":
            _SUPPRESS_OLD_15_31 = False
        else:
            return None
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = BLUE if level <= 2 else INK
    return p


def add_p(doc, text):
    if suppress_old_15_31():
        return None
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    if suppress_old_15_31():
        return
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)


def add_code(doc, code):
    if suppress_old_15_31():
        return None
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8.3)
    r.font.color.rgb = RGBColor(38, 50, 56)
    return p


def add_image(doc, path, caption, width=6.6):
    if suppress_old_15_31():
        return
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def read_csv(path):
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def pct(x):
    try:
        return f"{float(x) * 100:.2f}%"
    except Exception:
        return str(x)


def add_feature_section(doc):
    add_heading(doc, "3.1.1. Phân tích đầy đủ các đặc trưng tiếng nói và ảnh dẫn chứng", 3)
    add_p(
        doc,
        "Bản này không chỉ liệt kê feature mà còn nối từng feature với vai trò trong kiến trúc 06D. Mỗi lập luận về feature đều có hình dẫn chứng từ 6 cảm xúc. Các hình được sinh từ audio thật trong ser_processed/audio_16k, ưu tiên cùng domain SAVEE để minh họa feature mà hạn chế nhiễu do khác dataset."
    )
    add_table(
        doc,
        ["Đặc trưng", "Là gì?", "Nó thể hiện điều gì?", "Dùng ở đâu trong 06D?", "Ảnh"],
        [
            ["Waveform", "Dạng sóng âm theo thời gian, biểu diễn biên độ tín hiệu tại từng thời điểm.", "Cường độ, khoảng nghỉ, nhịp nói và biến động biên độ thô.", "Branch C dùng waveform cho frozen emotion2vec. Waveform cũng là nguồn để trích xuất các feature khác.", "Hình 2"],
            ["MFCC", "Mel-Frequency Cepstral Coefficients, mô tả envelope phổ âm trên thang Mel.", "Timbre/màu giọng và hình dạng phổ âm.", "Branch A dùng chuỗi MFCC; Branch D lấy thống kê từ MFCC.", "Hình 3"],
            ["Delta MFCC", "Đạo hàm bậc 1 của MFCC theo thời gian.", "Tốc độ thay đổi màu giọng giữa các frame.", "Branch A dùng để học dynamics, không chỉ trạng thái tĩnh.", "Hình 3"],
            ["Delta-delta MFCC", "Đạo hàm bậc 2 của MFCC.", "Gia tốc thay đổi phổ âm, bắt biến động nhanh.", "Branch A dùng để bổ sung dynamics mạnh hơn.", "Hình 3"],
            ["Log-Mel", "Bản đồ năng lượng theo thời gian và dải tần Mel sau phép log.", "Vùng năng lượng cao/thấp, formant, texture phổ.", "Branch B dùng làm input ảnh time-frequency.", "Hình 6"],
            ["Delta log-Mel", "Đạo hàm bậc 1 của log-Mel theo thời gian.", "Dải tần nào đang tăng/giảm năng lượng nhanh.", "Branch B dùng làm kênh thứ 2.", "Hình 6"],
            ["Delta-delta log-Mel", "Đạo hàm bậc 2 của log-Mel.", "Gia tốc thay đổi năng lượng phổ.", "Branch B dùng làm kênh thứ 3.", "Hình 6"],
            ["RMS energy", "Năng lượng trung bình bình phương của tín hiệu trong frame.", "Độ lớn/loudness/cường độ giọng.", "Branch A dùng theo frame; Branch D lấy thống kê.", "Hình 2"],
            ["ZCR", "Zero-Crossing Rate, số lần tín hiệu đổi dấu trong một frame.", "Độ sắc, nhiễu, texture voiced/unvoiced hoặc high-frequency.", "Branch A dùng theo frame; Branch D lấy thống kê.", "Hình 4"],
            ["Centroid", "Tâm khối năng lượng phổ.", "Độ sáng của âm thanh.", "Branch A và D dùng như spectral shape cue.", "Hình 5"],
            ["Bandwidth", "Độ rộng phân bố năng lượng quanh centroid.", "Năng lượng tập trung hẹp hay trải rộng.", "Branch A và D.", "Hình 5"],
            ["Rolloff", "Ngưỡng tần số chứa phần lớn năng lượng phổ.", "Mức năng lượng lan lên vùng tần số cao.", "Branch A và D.", "Hình 5"],
            ["Spectral contrast", "Chênh lệch đỉnh - đáy năng lượng trong các dải tần.", "Độ rõ của cấu trúc phổ, peak-valley pattern.", "Branch A dùng 7 contrast channels; Branch D lấy thống kê.", "Hình 6"],
            ["Statistical vector", "Mean, std, min, max, median, IQR trên 132 temporal channels + 4 thống kê waveform.", "Tóm tắt toàn clip thành vector cố định.", "Stats branch MLP và RBF-SVM.", "Bảng 06D"],
        ],
        widths=[1.05, 1.65, 1.65, 1.75, 0.55],
        font_size=7.2,
    )

    add_image(doc, EVIDENCE_WAVE_RMS, "Hình 2. Waveform và RMS trên 6 emotion. RMS là bằng chứng trực quan cho cường độ/loudness theo thời gian.", width=6.8)
    add_p(
        doc,
        "Nhìn Hình 2, waveform cho thấy trực tiếp biên độ và khoảng nghỉ. RMS là đường năng lượng frame-level được tính từ waveform: khi waveform có biên độ mạnh, RMS thường tăng. Trong 06D, RMS không được dùng riêng lẻ để quyết định cảm xúc, mà được ghép vào temporal sequence để 1D-CNN/BiLSTM học cùng MFCC và spectral features. Lý do là cùng một mức năng lượng có thể mang ý nghĩa khác nhau tùy màu giọng và ngữ cảnh thời gian."
    )
    add_image(doc, EVIDENCE_MFCC, "Hình 3. MFCC, delta MFCC và delta-delta MFCC trên 6 emotion.", width=6.8)
    add_p(
        doc,
        "Hình 3 cho thấy MFCC là một bản đồ hệ số theo thời gian. MFCC mô tả spectral envelope/timbre; delta MFCC cho thấy tốc độ thay đổi của timbre; delta-delta MFCC cho thấy sự thay đổi đó đang tăng nhanh hay chậm lại. Vì cảm xúc không chỉ là màu giọng tại một thời điểm mà là cách màu giọng biến đổi trong cả câu nói, 06D ghép cả ba nhóm này vào Branch A."
    )
    add_image(doc, EVIDENCE_ZCR, "Hình 4. ZCR trên 6 emotion.", width=6.8)
    add_p(
        doc,
        "ZCR trong Hình 4 không phải là nhãn cảm xúc trực tiếp. Nó là low-level cue về texture tín hiệu: khi tín hiệu đổi dấu nhanh, ZCR tăng. Điều này có thể liên quan đến âm vô thanh, nhiễu, sắc giọng hoặc thành phần tần số cao. Do đó ZCR được dùng như một feature bổ sung cho Branch A và Branch D, nhưng luôn được học cùng các feature khác để tránh suy luận đơn giản quá mức."
    )
    add_image(doc, EVIDENCE_SPECTRAL, "Hình 5. Spectral centroid, bandwidth và rolloff trên 6 emotion.", width=6.8)
    add_p(
        doc,
        "Hình 5 trực quan hóa ba feature mô tả hình dạng phổ. Centroid cao thường gợi âm sáng/sắc hơn; bandwidth cao cho thấy năng lượng trải rộng hơn; rolloff cao cho thấy năng lượng lan tới vùng tần số cao hơn. Các feature này giải thích vì sao Branch A không chỉ dùng MFCC: nó cần thêm các cue phổ cụ thể để phân biệt giọng căng, sáng, mềm hoặc trầm."
    )
    add_image(doc, EVIDENCE_LOGMEL, "Hình 6. Log-Mel spectrogram và spectral contrast trên 6 emotion.", width=6.8)
    add_p(
        doc,
        "Hình 6 cho thấy log-Mel giữ cấu trúc thời gian - tần số rõ hơn MFCC: trục ngang là thời gian, trục dọc là dải Mel, màu là năng lượng. Vì vậy Branch B dùng log-Mel như ảnh đầu vào cho CNN. Spectral contrast bên phải cho thấy peak-valley structure ở các dải tần, được dùng như cue bổ sung trong Branch A/D."
    )


def add_reference_trace_section(doc):
    add_heading(doc, "1.5.1. Trace từ bài báo/reference sang kiến trúc đề xuất 06D", 3)
    add_p(
        doc,
        "Để tránh việc đề xuất kiến trúc theo cảm tính, bảng dưới đây chỉ rõ từng thành phần trong 06D được lấy cảm hứng từ đâu. Mỗi reference không được copy nguyên xi; nhóm chỉ lấy ý tưởng phù hợp với điều kiện dữ liệu, tài nguyên Kaggle/Colab và mục tiêu demo sau này. Vì vậy cột cuối giải thích cách 06D adapt lại ý tưởng đó."
    )
    add_table(
        doc,
        ["Thành phần trong 06D", "Reference / link", "Trong paper người ta dùng ở đâu?", "Người ta dùng để làm gì?", "06D adapt lại như thế nào?"],
        [
            [
                "Multi-representation design",
                "CA-MSER - Speech Emotion Recognition with Co-Attention based Multi-level Acoustic Information\nhttps://arxiv.org/abs/2203.15326",
                "Paper trích xuất multi-level acoustic information gồm MFCC, spectrogram và high-level acoustic information từ wav2vec2; sau đó xem các feature này như nhiều input/modalities để fusion.",
                "Giải quyết vấn đề một loại feature không đủ bao quát cảm xúc. MFCC giữ low-level cepstral cue, spectrogram giữ time-frequency cue, pretrained model giữ high-level acoustic cue.",
                "06D giữ tinh thần multi-level nhưng đổi high-level branch từ wav2vec2 sang emotion2vec vì emotion2vec chuyên hơn cho emotion representation. 06D dùng Branch A temporal, Branch B log-Mel, Branch C emotion2vec, Branch D stats."
            ],
            [
                "Co-attention fusion",
                "CA-MSER\nhttps://arxiv.org/abs/2203.15326",
                "Paper dùng co-attention module để các representation tương tác với nhau trước khi classification, thay vì chỉ concatenate đơn giản.",
                "Cho phép mô hình học quan hệ giữa nhiều nguồn acoustic information và chọn thông tin cảm xúc hữu ích hơn.",
                "06D dùng EmotionGuidedCoAttention: z_e2v làm query, z_temporal và z_spectral làm key/value. Như vậy pretrained emotion embedding đóng vai trò hướng dẫn temporal/spectral cues."
            ],
            [
                "Frozen pretrained emotion branch",
                "emotion2vec: Self-Supervised Pre-Training for Speech Emotion Representation\nhttps://arxiv.org/abs/2312.15185\nModelScope: https://www.modelscope.cn/models/iic/emotion2vec_base",
                "emotion2vec pretrain một universal speech emotion representation bằng self-supervised online distillation, kết hợp utterance-level và frame-level loss. Paper cho thấy chỉ train linear layers downstream vẫn hiệu quả trên SER.",
                "Tạo embedding cảm xúc mạnh hơn so với feature thủ công, đặc biệt khi dữ liệu downstream ít.",
                "06D dùng emotion2vec ở chế độ frozen, không fine-tune toàn bộ model. Sau đó dùng Adapter MLP để chiếu embedding về 192-D. Cách này nhẹ hơn và giảm overfit trên 4 dataset nhỏ."
            ],
            [
                "Temporal branch: local feature + recurrent context",
                "Ahmed et al. - An Ensemble 1D-CNN-LSTM-GRU Model with Data Augmentation for SER\nhttps://arxiv.org/abs/2112.05666",
                "Paper xây ba network: 1D-CNN, CNN-LSTM và CNN-GRU. Các model bắt đầu từ local feature acquiring blocks gồm dilated 1D-CNN, max pooling, batch normalization, sau đó thêm LSTM/GRU để học long-term dependencies.",
                "1D-CNN học local hidden features từ speech features; LSTM/GRU học phụ thuộc dài hạn trong tín hiệu giọng nói.",
                "06D không tái tạo nguyên ensemble của Ahmed. 06D lấy ý tưởng local temporal filters + recurrent context: Conv1D học local acoustic patterns, BiLSTM học context hai chiều, AttentionPooling chọn frame quan trọng."
            ],
            [
                "Weighted/stacking ensemble",
                "Ahmed et al.\nhttps://arxiv.org/abs/2112.05666",
                "Paper dùng weighted average ensemble để kết hợp kết quả từ 1D-CNN, CNN-LSTM và CNN-GRU.",
                "Giảm phụ thuộc vào một model đơn; model nào mạnh ở một dataset/feature có thể bù cho model khác.",
                "06D dùng validation-weighted stacking thay vì weighted average cố định. p_deep, p_stats và optional p_e2v được kết hợp dựa trên validation split; test không dùng để chọn trọng số."
            ],
            [
                "Handcrafted acoustic features",
                "Ahmed et al.\nhttps://arxiv.org/abs/2112.05666; Ullah et al. feature fusion 1D-CNN\nhttps://doi.org/10.1109/ICIT56493.2022.9989197",
                "Ahmed dùng MFCC, Log Mel-Scaled Spectrogram, ZCR, Chromagram và RMS. Ullah dùng nhóm feature như ZCR, energy, entropy of energy, RMS và MFCC rồi đưa vào 1D-CNN.",
                "Các low-level acoustic features giúp mô hình thấy loudness, zero-crossing texture, spectral envelope và energy distribution.",
                "06D dùng 132 temporal channels gồm MFCC/delta/delta2/RMS/ZCR/centroid/bandwidth/rolloff/contrast. Ngoài ra còn dùng statistical vector từ các temporal channels cho stats branch."
            ],
            [
                "Statistical feature branch",
                "Chowdhury et al. - A Novel Hybrid Deep Learning Technique for Speech Emotion Detection using Feature Engineering\nhttps://arxiv.org/abs/2507.07046",
                "Paper theo hướng feature engineering mạnh và dùng DCRF-BiLSTM trên nhiều dataset RAVDESS, TESS, SAVEE, EmoDB, CREMA-D. Trọng tâm là khai thác feature thủ công giàu thông tin trước khi model sequence học.",
                "Tận dụng acoustic descriptors ổn định để tăng hiệu năng khi dữ liệu SER nhỏ hoặc nhiều domain.",
                "06D chưa tái tạo DCRF-BiLSTM. 06D lấy ý tưởng feature engineering: từ 132 temporal channels tạo stats vector mean/std/min/max/median/IQR + waveform stats, đưa vào Stats MLP và RBF-SVM."
            ],
            [
                "Residual CNN in spectrogram branch",
                "ResNet - Deep Residual Learning for Image Recognition\nhttps://arxiv.org/abs/1512.03385",
                "ResNet dùng shortcut/residual connection để layer học residual function thay vì học mapping hoàn toàn mới.",
                "Giúp mạng sâu dễ tối ưu hơn, giảm tình trạng degradation khi tăng số layer.",
                "06D dùng ResidualSEBlock trong Branch B: Conv2D + BN + GELU + Conv2D + BN + SE2D + shortcut. Vì log-Mel là ảnh time-frequency, residual CNN phù hợp để học pattern phổ sâu hơn."
            ],
            [
                "SE/channel attention in spectrogram branch",
                "Squeeze-and-Excitation Networks\nhttps://arxiv.org/abs/1709.01507",
                "SENet dùng squeeze-excitation để recalibrate channel-wise feature responses bằng cách học trọng số cho từng channel.",
                "Channel nào hữu ích hơn sẽ được tăng trọng số; channel ít hữu ích bị giảm.",
                "06D dùng SE2D trong ResidualSEBlock. Với spectrogram feature maps, SE2D giúp nhấn mạnh các channel đang chứa cue cảm xúc rõ hơn."
            ],
            [
                "Multi-head attention mechanism",
                "Attention Is All You Need\nhttps://arxiv.org/abs/1706.03762",
                "Transformer dùng query, key, value và multi-head attention để học nhiều kiểu quan hệ giữa các token.",
                "Mỗi head có thể học một quan hệ khác nhau; attention giúp mô hình tập trung vào thông tin liên quan.",
                "06D dùng nn.MultiheadAttention trong EmotionGuidedCoAttention. Ở đây token không phải word mà là embedding acoustic: z_e2v query z_temporal/z_spectral."
            ],
            [
                "RBF-SVM probability branch",
                "SVM/kernel trick tradition; used widely as classical baseline in SER feature engineering.",
                "SVM nhận vector cố định và tìm decision boundary trong không gian kernel; RBF cho phép phân tách phi tuyến.",
                "Làm model cổ điển mạnh cho handcrafted statistics, nhất là khi feature vector có ý nghĩa rõ.",
                "06D dùng StandardScaler + RBF-SVM(probability=True) trên X_stats để tạo p_stats[6], sau đó stacking với p_deep. Nhánh này không thay deep model mà bổ sung góc nhìn thống kê."
            ],
        ],
        widths=[1.15, 1.55, 1.75, 1.55, 1.85],
        font_size=6.7,
    )
    add_p(
        doc,
        "Từ bảng trên, kiến trúc 06D không phải là ghép ngẫu nhiên nhiều kỹ thuật. Mỗi thành phần có một vai trò riêng: handcrafted/temporal features bắt tín hiệu âm học thấp; spectrogram CNN bắt cấu trúc time-frequency; emotion2vec cung cấp representation cảm xúc đã học trước; co-attention cho các representation tương tác; stats/SVM và stacking bổ sung hướng feature engineering cổ điển. Đây là lý do 06D phù hợp làm proposed architecture cho giai đoạn midterm."
    )
    add_heading(doc, "1.5.2. Vì sao không copy nguyên mô hình của paper?", 3)
    add_p(
        doc,
        "Một điểm cần nói rõ trong báo cáo là nhóm không copy nguyên CA-MSER, Ahmed ensemble hay DCRF-BiLSTM. Lý do là các paper dùng dataset/protocol/tài nguyên khác nhau. Nếu copy nguyên, mô hình có thể quá nặng, khó chạy Kaggle, hoặc không khớp với mục tiêu demo feedback sau này. Thay vào đó, 06D lấy các ý tưởng đã được chứng minh là hợp lý rồi adapt thành một pipeline gọn hơn: frozen pretrained embedding để nhẹ tài nguyên; co-attention nhẹ thay vì fusion lớn; temporal/spectral/statistical branches để bao phủ nhiều cue cảm xúc; stacking dùng validation để tránh test leakage."
    )


def add_implementation_shape_section(doc):
    add_heading(doc, "3.2.2. Giải thích đầy đủ các tensor shape trong code 06D", 2)
    add_p(
        doc,
        "Phần implementation cần đọc được cả với người chưa quen PyTorch. Vì vậy trước khi đi vào code, cần giải thích các ký hiệu shape. Trong notebook 06D, ký hiệu dạng [B, C, T] hoặc [B, D] là cách mô tả kích thước tensor khi train theo mini-batch."
    )
    add_table(
        doc,
        ["Ký hiệu", "Ý nghĩa", "Trong 06D là gì?", "Tại sao cần biết?"],
        [
            ["B", "Batch size, số audio sample được đưa vào model trong một lần forward.", "Ví dụ BATCH_SIZE=24 thì B=24 khi batch đầy.", "Ảnh hưởng RAM/GPU và shape của mọi tensor."],
            ["T", "Số frame thời gian sau khi chia audio thành các frame ngắn bằng hop_length.", "Với audio 3s, sr=16000, hop=160, T xấp xỉ 301 frames trước/pad theo MAX_FRAMES.", "Model sequence/CNN quét theo trục thời gian này."],
            ["132", "Số channel đặc trưng temporal trên mỗi frame.", "40 MFCC + 40 delta + 40 delta2 + RMS + ZCR + centroid + bandwidth + rolloff + 7 contrast = 132.", "Đây là input channel của Conv1D. Nếu ghi [B,126,T] là không đúng với 06D hiện tại."],
            ["3", "Số channel spectrogram.", "log-Mel, delta log-Mel, delta-delta log-Mel.", "Giống ảnh 3 kênh nhưng mỗi kênh là một loại biến đổi phổ âm."],
            ["96", "Số Mel bins.", "N_MELS=96.", "Là chiều tần số Mel của log-Mel spectrogram."],
            ["D_e2v", "Chiều embedding do emotion2vec trả về.", "Phụ thuộc backend/model emotion2vec, sau đó được scale và đưa vào MLPBranch.", "Không hard-code vì pretrained backend có thể trả vector khác chiều."],
            ["796", "Chiều statistical vector trong 06D.", "132 temporal channels x 6 stats + 4 waveform stats = 796.", "Input cho Stats MLP và RBF-SVM."],
            ["192", "Chiều embedding chung của các branch deep.", "z_temporal, z_spectral, z_e2v, z_context đều là 192-D.", "Giúp co-attention và fusion ghép các branch dễ dàng."],
            ["128", "Chiều z_stats sau Stats MLP.", "Stats MLP project X_stats thành z_stats [B,128].", "Stats branch nhỏ hơn để giảm overfit và giảm kích thước fusion."],
            ["896", "Chiều vector sau concat.", "192 + 192 + 192 + 192 + 128 = 896.", "Input cho Fusion MLP."],
            ["6", "Số lớp cảm xúc.", "neutral, happy, sad, angry, fear, disgust.", "Classifier cuối tạo logits/probability cho 6 nhãn."],
        ],
        widths=[0.75, 1.65, 2.6, 2.15],
        font_size=7.6,
    )
    add_p(
        doc,
        "Lưu ý quan trọng: trong notebook 06D hiện tại, temporal input là [B, 132, T], không phải [B, 126, T]. Con số 132 đến từ spectral_contrast của librosa trả về 7 bands khi dùng cấu hình mặc định n_bands=6. Công thức là 40 + 40 + 40 + 1 + 1 + 1 + 1 + 1 + 7 = 132."
    )

    add_heading(doc, "3.2.3. Code-level implementation flow", 2)
    add_p(
        doc,
        "Bảng sau mô tả đầy đủ hơn các khối code chính trong notebook 06D, input/output của từng bước và lý do bước đó tồn tại. Đây là phần cần dùng khi giảng viên hỏi: mô hình được implement như thế nào, dữ liệu đi qua code ra sao."
    )
    add_table(
        doc,
        ["Bước code", "Code/function chính", "Input", "Output", "Giải thích"],
        [
            ["1. Resolve data", "find_ser_processed(), metadata = pd.read_csv(...)", "Thư mục ser_processed", "metadata dataframe", "Tự tìm dataset trên Kaggle/local, đọc label, dataset, speaker_id, duration."],
            ["2. Load audio", "resolve_audio_path(row), load_audio_fixed(path)", "row trong metadata", "waveform 16 kHz độ dài cố định", "Đảm bảo mọi sample có cùng sample rate và độ dài để trích xuất feature ổn định."],
            ["3. Acoustic extraction", "extract_acoustic_features(y)", "waveform y", "temporal, spectral, stats", "Tạo 3 nhóm feature chính cho Branch A, B, D."],
            ["4. emotion2vec extraction", "Emotion2VecExtractor.extract_one(wav_path)", "file WAV 16 kHz", "X_e2v vector", "Tạo pretrained emotion embedding cho Branch C."],
            ["5. Cache", "build_or_load_feature_cache()", "metadata + audio", "X_temporal, X_spectral, X_stats, X_e2v", "Lưu feature để không phải trích xuất lại mỗi lần chạy."],
            ["6. Scaler", "compute_scalers(train_idx)", "train indices", "temporal/spectral mean-std, stats/e2v scaler", "Fit scaler chỉ trên train để tránh leakage."],
            ["7. Dataset", "SERDataset.__getitem__", "sample index", "dict tensor temporal/spectral/stats/e2v/label", "Chuẩn hóa feature, augment train-only, trả tensor cho DataLoader."],
            ["8. Deep model", "Emotion2VecCoAttentionSER.forward", "4 tensor input", "logits, fused, attention info", "Forward qua 4 branch, co-attention, fusion và classifier."],
            ["9. Training loop", "train_deep_model(), run_epoch()", "DataLoader", "best model by val macro-F1", "Train bằng loss phân loại, chọn checkpoint tốt nhất trên validation."],
            ["10. Classical models", "RBF-SVM, e2v_logreg, e2v_mlp", "X_stats hoặc X_e2v", "p_stats / p_e2v", "Tạo probability phụ để stacking."],
            ["11. Stacking", "validation predictions", "p_deep, p_stats, p_e2v", "p_final", "Kết hợp xác suất dựa trên validation, test chỉ dùng báo cáo."],
        ],
        widths=[1.0, 1.55, 1.25, 1.25, 2.3],
        font_size=7.1,
    )

    add_heading(doc, "3.2.4. Code trích xuất feature và ý nghĩa từng dòng chính", 2)
    add_code(
        doc,
        "def extract_acoustic_features(y):\n"
        "    mfcc = librosa.feature.mfcc(y=y, sr=TARGET_SR, n_mfcc=40, ...)\n"
        "    delta = librosa.feature.delta(mfcc)\n"
        "    delta2 = librosa.feature.delta(mfcc, order=2)\n"
        "    rms = librosa.feature.rms(y=y, ...)\n"
        "    zcr = librosa.feature.zero_crossing_rate(y, ...)\n"
        "    centroid = librosa.feature.spectral_centroid(y=y, ...)\n"
        "    bandwidth = librosa.feature.spectral_bandwidth(y=y, ...)\n"
        "    rolloff = librosa.feature.spectral_rolloff(y=y, ...)\n"
        "    contrast = librosa.feature.spectral_contrast(y=y, ...)  # 7 bands\n\n"
        "    temporal = np.vstack([mfcc, delta, delta2, rms, zcr, centroid, bandwidth, rolloff, contrast])\n"
        "    temporal = pad_or_trim_time(temporal).astype(np.float32)  # [132, T]\n\n"
        "    mel = librosa.feature.melspectrogram(y=y, sr=TARGET_SR, n_mels=96, ...)\n"
        "    logmel = librosa.power_to_db(mel, ref=np.max)\n"
        "    d_logmel = librosa.feature.delta(logmel)\n"
        "    d2_logmel = librosa.feature.delta(logmel, order=2)\n"
        "    spectral = np.stack([logmel, d_logmel, d2_logmel], axis=0)  # [3, 96, T]\n\n"
        "    # stats: mean/std/min/max/median/IQR over each temporal channel + waveform stats\n"
        "    stats = np.concatenate([mean, std, minv, maxv, med, iqr, duration_energy])  # [796]\n"
        "    return temporal, spectral, stats"
    )
    add_p(
        doc,
        "Đoạn code trên là phần nối trực tiếp giữa lý thuyết feature và mô hình. Branch A nhận temporal [132,T]; Branch B nhận spectral [3,96,T]; Branch D nhận stats [796]. Branch C không lấy từ hàm này mà lấy từ emotion2vec extractor trên waveform."
    )

    add_heading(doc, "3.2.5. Code Dataset, scaler và train-only augmentation", 2)
    add_code(
        doc,
        "def compute_scalers(train_idx):\n"
        "    temporal_mean/std = X_temporal[train_idx].mean/std(...)\n"
        "    spectral_mean/std = X_spectral[train_idx].mean/std(...)\n"
        "    stats_scaler = StandardScaler().fit(X_stats[train_idx])\n"
        "    e2v_scaler = StandardScaler().fit(X_e2v[train_idx])\n\n"
        "class SERDataset(Dataset):\n"
        "    def __getitem__(self, item):\n"
        "        i = self.indices[item]\n"
        "        temporal = normalize(X_temporal[i], temporal_mean/std)\n"
        "        spectral = normalize(X_spectral[i], spectral_mean/std)\n"
        "        if self.train and USE_AUGMENTATION:\n"
        "            temporal = augment_temporal(temporal)   # time mask/noise\n"
        "            spectral = augment_spectral(spectral)   # time mask/frequency mask\n"
        "        stats = stats_scaler.transform(X_stats[i:i+1])[0]\n"
        "        e2v = e2v_scaler.transform(X_e2v[i:i+1])[0]\n"
        "        return {temporal, spectral, stats, e2v, label}"
    )
    add_p(
        doc,
        "Điểm quan trọng là scaler chỉ fit trên train split. Validation/test chỉ transform bằng scaler đã học từ train. Augmentation cũng chỉ chạy khi self.train=True. Cách này tránh data leakage: mô hình không được nhìn thấy biến thể của validation/test trong training."
    )

    add_heading(doc, "3.2.6. Code forward pass của mô hình 06D", 2)
    add_code(
        doc,
        "def forward(self, temporal, spectral, stats, e2v):\n"
        "    z_t, temporal_attn = self.temporal_branch(temporal)  # [B,132,T] -> [B,192]\n"
        "    z_s = self.spectral_branch(spectral)                # [B,3,96,T] -> [B,192]\n"
        "    z_e = self.e2v_branch(e2v)                           # [B,D_e2v] -> [B,192]\n"
        "    z_stats = self.stats_branch(stats)                   # [B,796] -> [B,128]\n\n"
        "    z_context, co_weights = self.co_attention(z_e, z_t, z_s)  # [B,192]\n"
        "    z = torch.cat([z_t, z_s, z_e, z_context, z_stats], dim=1)  # [B,896]\n"
        "    fused = self.fusion(z)                               # [B,896] -> [B,192]\n"
        "    logits = self.classifier(fused)                      # [B,192] -> [B,6]\n"
        "    return logits, fused, attention_info"
    )
    add_p(
        doc,
        "Forward pass này cho thấy 06D không phải bốn model rời rạc. Bốn branch tạo ra representation, sau đó emotion-guided co-attention tạo z_context, rồi fusion MLP gom tất cả thành fused representation trước khi classifier tạo logits cho 6 cảm xúc."
    )


def add_feature_extraction_pipeline_section(doc):
    add_heading(doc, "3.1.2. Quy trình trích xuất đặc trưng: làm như thế nào, vì sao làm như vậy, tham khảo từ đâu", 3)
    add_p(
        doc,
        "Phần feature extraction của 06D được thiết kế theo nguyên tắc: một tín hiệu speech emotion phải được nhìn ở nhiều mức khác nhau. Ở mức thấp, waveform chứa biên độ và nhịp nói; ở mức cepstral, MFCC mô tả màu giọng/timbre; ở mức động học, delta và delta-delta mô tả sự thay đổi theo thời gian; ở mức phổ, log-Mel giữ cấu trúc time-frequency; ở mức toàn cục, statistical vector tóm tắt đặc điểm của cả câu nói. Đây là lý do 06D không dùng một feature duy nhất."
    )
    add_table(
        doc,
        ["Bước", "Cách trích xuất trong 06D", "Tạo ra input nào?", "Vì sao cần?", "Nguồn tham khảo"],
        [
            [
                "1. Chuẩn hóa audio",
                "Đọc waveform, resample về 16 kHz, mono, pad/trim về TARGET_DURATION=3s trong bản 06D gốc.",
                "waveform y",
                "Giúp mọi mẫu có cùng sample rate và số frame tương đối ổn định để batch training không lệch kích thước.",
                "Quy ước preprocessing SER/audio ML; phù hợp cách emotion2vec nhận waveform 16 kHz."
            ],
            [
                "2. MFCC",
                "librosa.feature.mfcc(y, sr=16000, n_mfcc=40). MFCC được tính từ log-power Mel spectrogram rồi biến đổi DCT.",
                "40 channel MFCC theo thời gian",
                "MFCC nén hình dạng phổ thành cepstral coefficients, thể hiện timbre/spectral envelope. Đây là feature cổ điển và phổ biến trong speech/speaker/emotion recognition.",
                "librosa MFCC docs: https://librosa.org/doc/latest/generated/librosa.feature.mfcc.html; Ahmed et al. dùng MFCC trong SER: https://arxiv.org/abs/2112.05666"
            ],
            [
                "3. Delta và delta-delta",
                "librosa.feature.delta(mfcc, order=1) và librosa.feature.delta(mfcc, order=2). Delta là ước lượng đạo hàm cục bộ theo trục thời gian.",
                "40 delta + 40 delta-delta",
                "Cảm xúc không chỉ nằm ở màu giọng tĩnh mà còn ở cách giọng biến đổi trong câu. Delta mô tả tốc độ thay đổi; delta-delta mô tả gia tốc thay đổi.",
                "librosa delta docs: https://librosa.org/doc/latest/generated/librosa.feature.delta.html; hướng CNN-LSTM/GRU của Ahmed et al. nhấn mạnh local feature + temporal dependency."
            ],
            [
                "4. Low-level temporal cues",
                "Tính RMS, ZCR, spectral centroid, bandwidth, rolloff và 7 spectral contrast theo frame.",
                "12 channel bổ sung cho Branch A",
                "Nhóm này bổ sung loudness, texture voiced/unvoiced, độ sáng phổ, độ rộng phổ, năng lượng vùng cao và peak-valley contrast. Chúng giúp Branch A không chỉ phụ thuộc vào MFCC.",
                "librosa spectral feature docs; Ullah et al. dùng ZCR, energy, entropy/RMS, MFCC cho feature fusion 1D-CNN: https://doi.org/10.1109/ICIT56493.2022.9989197"
            ],
            [
                "5. Concatenate temporal",
                "np.vstack([mfcc, delta, delta2, rms, zcr, centroid, bandwidth, rolloff, contrast]).",
                "[132, T]",
                "132 = 40 + 40 + 40 + 1 + 1 + 1 + 1 + 1 + 7. T là số frame thời gian. Nhánh 1D-CNN/BiLSTM cần input dạng sequence nên giữ nguyên trục thời gian.",
                "Feature-level fusion trong CA-MSER và các paper SER multi-feature."
            ],
            [
                "6. Log-Mel spectrogram",
                "librosa.feature.melspectrogram(..., n_mels=96) rồi librosa.power_to_db(...).",
                "log-Mel [96, T]",
                "Log-Mel giữ rõ cấu trúc thời gian - tần số như một ảnh. Phép log nén dynamic range, làm vùng năng lượng yếu vẫn còn nhìn được thay vì bị vùng năng lượng mạnh lấn át.",
                "librosa melspectrogram docs: https://librosa.org/doc/latest/generated/librosa.feature.melspectrogram.html; AST/2D-CNN dùng spectrogram như ảnh âm thanh."
            ],
            [
                "7. Delta log-Mel",
                "Tính delta và delta-delta của log-Mel, sau đó stack 3 kênh: log-Mel, delta log-Mel, delta-delta log-Mel.",
                "[3, 96, T]",
                "Giống ảnh 3 kênh nhưng không phải RGB: kênh 1 là năng lượng phổ, kênh 2 là tốc độ thay đổi năng lượng, kênh 3 là gia tốc thay đổi năng lượng.",
                "librosa delta docs; ý tưởng multi-channel spectrogram trong các hướng CNN spectrogram SER."
            ],
            [
                "8. Statistical vector",
                "Với 132 temporal channels, tính mean, std, min, max, median, IQR theo thời gian; cộng thêm waveform-level stats.",
                "[796]",
                "MLP/SVM không xử lý sequence biến thiên độ dài tốt như RNN/CNN nên cần vector cố định. Mean/std/min/max/median/IQR giữ thông tin phân bố toàn clip.",
                "Feature engineering SER; Chowdhury et al. DCRF-BiLSTM nhấn mạnh handcrafted statistical/acoustic features trên nhiều dataset: https://arxiv.org/abs/2507.07046"
            ],
            [
                "9. emotion2vec embedding",
                "Đưa waveform 16 kHz vào frozen emotion2vec, lấy utterance-level embedding rồi qua adapter MLP.",
                "X_e2v / z_e2v",
                "Feature thủ công có giới hạn; emotion2vec đã học representation cảm xúc từ dữ liệu lớn, nên bổ sung high-level emotion cue mà không cần fine-tune nặng.",
                "emotion2vec paper: https://arxiv.org/abs/2312.15185; ModelScope model: https://www.modelscope.cn/models/iic/emotion2vec_base"
            ],
        ],
        widths=[0.75, 2.05, 0.95, 2.0, 1.55],
        font_size=6.8,
    )
    add_p(
        doc,
        "Điểm quan trọng là 06D không trích xuất feature sau khi nhìn label test. Feature extraction là phép biến đổi tín hiệu audio thành representation; scaler và model selection mới là phần phải học từ train. Vì vậy trong implementation, scaler luôn fit trên train split, còn validation/test chỉ được transform bằng scaler đó."
    )
    add_code(
        doc,
        "# Input waveform y: 16 kHz, mono, fixed duration\n"
        "mfcc = librosa.feature.mfcc(y=y, sr=16000, n_mfcc=40)\n"
        "delta = librosa.feature.delta(mfcc, order=1)\n"
        "delta2 = librosa.feature.delta(mfcc, order=2)\n"
        "low = [rms, zcr, centroid, bandwidth, rolloff, contrast_7]\n"
        "X_temporal = vstack([mfcc, delta, delta2, *low])       # [132, T]\n\n"
        "mel = librosa.feature.melspectrogram(y=y, sr=16000, n_mels=96)\n"
        "logmel = librosa.power_to_db(mel, ref=np.max)\n"
        "X_spectral = stack([logmel, delta(logmel), delta2(logmel)])  # [3, 96, T]\n\n"
        "X_stats = concat([mean/std/min/max/median/IQR over 132 channels, waveform_stats])  # [796]\n"
        "X_e2v = frozen_emotion2vec(waveform_16k)  # pretrained emotion representation"
    )


def add_augmentation_detail_section(doc):
    add_heading(doc, "3.2.7. Data augmentation: làm ra sao, áp dụng ở đâu, vì sao không gây leakage", 2)
    add_p(
        doc,
        "Augmentation trong 06D là train-only augmentation, nghĩa là chỉ biến đổi dữ liệu khi sample thuộc train split. Validation và test luôn giữ nguyên. Lý do là nếu augment trước khi chia train/validation/test, một bản biến thể của cùng câu nói có thể rơi vào train còn bản gốc rơi vào test; khi đó mô hình đã gián tiếp nhìn thấy test, làm kết quả bị data leakage."
    )
    add_table(
        doc,
        ["Loại augmentation", "Cách làm trong 06D", "Áp dụng cho branch nào?", "Tại sao dùng?", "Tham khảo"],
        [
            [
                "Temporal noise/dropout nhẹ",
                "Thêm nhiễu nhỏ hoặc mask một số frame/channel trong temporal tensor sau khi đã split.",
                "Branch A",
                "Buộc 1D-CNN/BiLSTM không phụ thuộc quá mạnh vào vài frame hoặc vài channel quá dễ học, giảm overfit.",
                "Ý tưởng regularization cho acoustic sequence; gần với masking trong SpecAugment."
            ],
            [
                "Time masking",
                "Che một đoạn frame liên tiếp trên temporal hoặc spectrogram.",
                "Branch A và Branch B",
                "Mô phỏng việc một đoạn tín hiệu không rõ/không quan trọng; buộc model học ngữ cảnh rộng hơn thay vì memorizing đoạn ngắn.",
                "SpecAugment dùng time masking trực tiếp trên feature input: https://arxiv.org/abs/1904.08779"
            ],
            [
                "Frequency masking",
                "Che một vùng mel-bin liên tiếp trên spectrogram.",
                "Branch B",
                "Buộc CNN không chỉ bám vào một dải tần cố định; hữu ích vì speaker/domain có thể khác nhau về formant/microphone.",
                "SpecAugment frequency masking: https://arxiv.org/abs/1904.08779"
            ],
            [
                "Feature scaling sau split",
                "Fit mean/std hoặc StandardScaler bằng train_idx; validation/test chỉ transform.",
                "Tất cả branch có feature numeric",
                "Không phải augmentation nhưng là bước chống leakage rất quan trọng. Nếu fit scaler trên toàn bộ data thì test distribution đã bị đưa vào train pipeline.",
                "ML pipeline best practice; áp dụng trong notebook 06D."
            ],
            [
                "Không augment Branch C trong 06D gốc",
                "emotion2vec embedding được trích xuất từ waveform gốc và giữ frozen.",
                "Branch C",
                "Giữ pretrained emotion representation ổn định; tránh làm cache quá lớn và tránh tăng thời gian chạy Kaggle.",
                "emotion2vec được dùng như frozen representation: https://arxiv.org/abs/2312.15185"
            ],
        ],
        widths=[1.2, 1.75, 1.1, 2.0, 1.45],
        font_size=7.0,
    )
    add_p(
        doc,
        "Nói ngắn gọn, augmentation của 06D không nhằm tạo thêm label mới, mà nhằm làm model bớt nhạy với nhiễu nhỏ, đoạn mất thông tin và khác biệt dải tần. Đây là lý do augmentation được đặt trong SERDataset.__getitem__ với điều kiện self.train=True, thay vì lưu augment cố định trước khi chia dữ liệu."
    )
    add_code(
        doc,
        "class SERDataset(Dataset):\n"
        "    def __getitem__(self, idx):\n"
        "        temporal = normalize(X_temporal[i])\n"
        "        spectral = normalize(X_spectral[i])\n"
        "        if self.train and USE_AUGMENTATION:\n"
        "            temporal = augment_temporal(temporal)   # noise/time mask on [132, T]\n"
        "            spectral = augment_spectral(spectral)   # time/frequency mask on [3, 96, T]\n"
        "        return temporal, spectral, stats, e2v, label\n\n"
        "# validation/test: self.train=False -> no augmentation, only deterministic transform"
    )


def add_reference_for_extraction_augmentation_section(doc):
    add_heading(doc, "3.2.8. Các ý tưởng feature/augmentation được lấy từ đâu và được adapt như thế nào", 2)
    add_table(
        doc,
        ["Nguồn", "Ý tưởng gốc", "06D lấy phần nào?", "Vì sao phù hợp với đề tài"],
        [
            [
                "librosa documentation\nhttps://librosa.org/doc/latest/generated/librosa.feature.mfcc.html",
                "MFCC là Mel-frequency cepstral coefficients; docs cũng mô tả tham số n_mfcc và cách MFCC có thể sinh từ log-power Mel spectrogram.",
                "Dùng librosa để tính MFCC=40, log-Mel, RMS, ZCR, spectral centroid/bandwidth/rolloff/contrast.",
                "Đây là thư viện chuẩn, dễ chạy trên Kaggle/Colab, không cần GPU, phù hợp đồ án sinh viên."
            ],
            [
                "librosa delta docs\nhttps://librosa.org/doc/latest/generated/librosa.feature.delta.html",
                "Delta là ước lượng đạo hàm cục bộ của feature theo một trục; order=1 là delta, order=2 là delta-delta.",
                "Tạo delta/delta-delta cho MFCC và log-Mel.",
                "SER cần dynamics vì cảm xúc thể hiện qua cách giọng thay đổi, không chỉ qua một frame tĩnh."
            ],
            [
                "SpecAugment\nhttps://arxiv.org/abs/1904.08779",
                "Che block theo trục thời gian và trục tần số trực tiếp trên acoustic feature input.",
                "06D dùng biến thể nhẹ: time mask cho temporal/spectral và frequency mask cho log-Mel.",
                "Tăng robustness mà không phải thu thêm dữ liệu; đặc biệt hữu ích khi dataset nhỏ và dễ overfit speaker."
            ],
            [
                "Ahmed et al. ensemble 1D-CNN/CNN-LSTM/CNN-GRU\nhttps://arxiv.org/abs/2112.05666",
                "Dùng MFCC, log-Mel, ZCR, chromagram, RMS và các mô hình temporal/recurrent để học speech emotion.",
                "06D lấy ý tưởng local acoustic feature + recurrent context, nhưng gom vào một TemporalBranch 1D-CNN + BiLSTM + attention.",
                "Giúp mô hình học cả pattern cục bộ và diễn tiến cảm xúc trong câu nói."
            ],
            [
                "Ullah et al. feature fusion 1D-CNN\nhttps://doi.org/10.1109/ICIT56493.2022.9989197",
                "Khai thác nhiều acoustic feature như energy/RMS/ZCR/MFCC rồi fusion cho 1D-CNN.",
                "06D đưa RMS/ZCR/spectral shape/contrast vào temporal sequence cùng MFCC/delta.",
                "Làm Branch A giàu cue hơn MFCC thuần."
            ],
            [
                "CA-MSER\nhttps://arxiv.org/abs/2203.15326",
                "Dùng multi-level acoustic information và co-attention để kết hợp nhiều nguồn biểu diễn.",
                "06D dùng multi-representation: temporal, spectrogram, emotion2vec, stats; dùng emotion-guided co-attention giữa e2v và acoustic embeddings.",
                "Phù hợp mục tiêu học thuật: không chỉ train model đơn mà giải thích được mỗi representation đóng góp gì."
            ],
            [
                "emotion2vec\nhttps://arxiv.org/abs/2312.15185",
                "Self-supervised pretrained speech emotion representation, dùng cho nhiều tác vụ emotion.",
                "06D dùng frozen emotion2vec embedding và adapter MLP, không fine-tune toàn bộ để tiết kiệm tài nguyên.",
                "Giữ hướng có khả năng demo/realtime hơn: rút embedding rồi classifier nhẹ, ít phụ thuộc GPU mạnh."
            ],
        ],
        widths=[1.55, 1.9, 1.9, 1.85],
        font_size=6.9,
    )
    add_p(
        doc,
        "Tóm lại, feature extraction và augmentation trong 06D không phải phần phụ của notebook. Đây là phần quyết định model nhìn thấy tín hiệu cảm xúc như thế nào. Nếu feature quá nghèo, mô hình dù sâu vẫn khó học; nếu augmentation sai thời điểm, kết quả có thể bị leakage. Vì vậy 06D đặt feature extraction thành bước reproducible có cache, còn augmentation đặt trong train dataset để giữ validation/test sạch."
    )


def add_implementation_architecture_walkthrough(doc):
    add_heading(doc, "3.2.1. Code kiến trúc 06D theo từng khối thuật toán", 2)
    add_p(
        doc,
        "Mục này viết lại kiến trúc đúng theo class trong notebook 06D. Lưu ý: bản 06D hiện tại dùng BiLSTM, không dùng BiGRU; và Conv1D thường, không dùng dilated Conv1D. BiGRU/dilated Conv1D là hướng nâng cấp có thể thử sau, còn phần dưới đây mô tả đúng code đang chạy để báo cáo không bị lệch implementation."
    )
    add_code(
        doc,
        "Input audio 16 kHz\n"
        "  -> feature cache\n"
        "  -> X_temporal [B, 132, T]\n"
        "  -> X_spectral [B, 3, 96, T]\n"
        "  -> X_stats [B, 796]\n"
        "  -> X_e2v [B, D_e2v]\n\n"
        "Branch A - TemporalBranch:\n"
        "  Input: X_temporal [B, 132, T]\n"
        "  1D-CNN:\n"
        "    Conv1d(132 -> 128, kernel=7, padding=3, bias=False)\n"
        "    BatchNorm1d(128) -> GELU -> Dropout\n"
        "    Conv1d(128 -> 160, kernel=5, padding=2, bias=False)\n"
        "    BatchNorm1d(160) -> GELU -> MaxPool1d(2) -> Dropout\n"
        "  BiLSTM:\n"
        "    LSTM(input_size=160, hidden_size=96, bidirectional=True)\n"
        "    output: [B, T/2, 192]\n"
        "  Attention pooling:\n"
        "    Linear(192 -> 96) -> Tanh -> Linear(96 -> 1)\n"
        "    softmax over time frames\n"
        "  Projection:\n"
        "    LayerNorm(192) -> Linear(192 -> 192) -> GELU -> Dropout\n"
        "  Output: z_temporal [B, 192]\n\n"
        "Branch B - SpectralBranch:\n"
        "  Input: X_spectral [B, 3, 96, T]\n"
        "  ResidualSEBlock(3 -> 32) -> MaxPool2d(2)\n"
        "  ResidualSEBlock(32 -> 64) -> MaxPool2d(2)\n"
        "  ResidualSEBlock(64 -> 128)\n"
        "  AdaptiveAvgPool2d((1,1)) -> Flatten\n"
        "  LayerNorm(128) -> Linear(128 -> 192) -> GELU -> Dropout\n"
        "  Output: z_spectral [B, 192]\n\n"
        "Branch C - Emotion2Vec adapter:\n"
        "  Input: X_e2v [B, D_e2v]\n"
        "  Linear(D_e2v -> 384) -> LayerNorm -> GELU -> Dropout\n"
        "  Linear(384 -> 192) -> GELU -> Dropout\n"
        "  Output: z_e2v [B, 192]\n\n"
        "Branch D - Stats branch:\n"
        "  Input: X_stats [B, 796]\n"
        "  Linear(796 -> 256) -> LayerNorm -> GELU -> Dropout\n"
        "  Linear(256 -> 128) -> GELU -> Dropout\n"
        "  Output: z_stats [B, 128]\n\n"
        "Co-attention + fusion:\n"
        "  q  = z_e2v.unsqueeze(1)                       # [B, 1, 192]\n"
        "  kv = stack([z_temporal, z_spectral], dim=1)   # [B, 2, 192]\n"
        "  z_context = MultiheadAttention(q, kv, kv)     # [B, 192]\n"
        "  z = concat([z_temporal, z_spectral, z_e2v, z_context, z_stats])\n"
        "    = [B, 192+192+192+192+128] = [B, 896]\n"
        "  Fusion MLP: Linear(896 -> 384) -> LayerNorm -> GELU -> Dropout -> Linear(384 -> 192)\n"
        "  Classifier: Linear(192 -> 6)\n"
        "  Output: logits [B, 6] -> softmax -> p_deep [B, 6]"
    )
    add_table(
        doc,
        ["Khối thuật toán", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng / tham khảo"],
        [
            [
                "1D-CNN",
                "Convolution 1 chiều trượt kernel dọc theo trục thời gian của X_temporal. Kernel nhìn một cửa sổ frame ngắn và học pattern cục bộ như năng lượng tăng, MFCC đổi nhanh, ZCR/spectral cue xuất hiện.",
                "[B,132,T] -> [B,128,T] -> [B,160,T]",
                "1D-CNN nhẹ hơn 2D-CNN và phù hợp chuỗi MFCC/acoustic feature. Ahmed et al. dùng 1D-CNN/CNN-LSTM/CNN-GRU cho SER: https://arxiv.org/abs/2112.05666"
            ],
            [
                "BatchNorm + GELU + Dropout",
                "BatchNorm ổn định phân phối feature map; GELU tạo phi tuyến mềm; Dropout tắt ngẫu nhiên một phần activation khi train.",
                "feature maps -> normalized/nonlinear/regularized maps",
                "Giúp train ổn định và giảm overfit, đặc biệt khi dataset SER không quá lớn."
            ],
            [
                "MaxPool1d",
                "Giảm chiều thời gian sau Conv1D bằng cách giữ activation nổi bật trong mỗi vùng nhỏ.",
                "[B,160,T] -> [B,160,T/2]",
                "Giảm nhiễu và chi phí tính toán; giữ pattern quan trọng trước khi đưa vào BiLSTM."
            ],
            [
                "BiLSTM",
                "LSTM hai chiều. Forward LSTM đọc từ đầu câu tới hiện tại; backward LSTM đọc từ cuối câu về hiện tại; output hai hướng được concat.",
                "[B,T/2,160] -> [B,T/2,192]",
                "Cảm xúc cần ngữ cảnh cả câu. CA-MSER dùng CNN + BiLSTM trong SER multi-level acoustic: https://arxiv.org/abs/2203.15326; DCRF-BiLSTM dùng BiLSTM cho SER: https://arxiv.org/abs/2507.07046"
            ],
            [
                "Attention pooling",
                "Tính score cho từng frame, softmax thành trọng số, rồi lấy tổng có trọng số theo thời gian.",
                "[B,T/2,192] -> [B,192]",
                "Không phải frame nào cũng chứa cảm xúc rõ. Attention giúp tập trung vào đoạn phát âm giàu thông tin. Attention reference: https://arxiv.org/abs/1706.03762"
            ],
            [
                "ResidualSEBlock 2D",
                "Hai Conv2D học pattern trên spectrogram; residual shortcut giữ thông tin gốc; SE2D tính trọng số channel rồi nhân lại feature maps.",
                "[B,3,96,T] -> [B,128,H',W']",
                "Log-Mel là ảnh time-frequency nên dùng 2D-CNN. Residual học sâu ổn định hơn; SE nhấn mạnh channel chứa cue cảm xúc. ResNet: https://arxiv.org/abs/1512.03385; SENet: https://arxiv.org/abs/1709.01507"
            ],
            [
                "AdaptiveAvgPool2d + projection",
                "Gom toàn bộ feature map spectrogram thành vector cố định rồi đưa về 192 chiều.",
                "[B,128,H',W'] -> [B,128] -> [B,192]",
                "Fusion cần vector cố định cùng kích thước với z_temporal và z_e2v."
            ],
            [
                "Emotion2Vec adapter MLP",
                "Nhận embedding pretrained X_e2v, qua hai lớp Linear + LayerNorm/GELU/Dropout để project về 192 chiều.",
                "[B,D_e2v] -> [B,192]",
                "emotion2vec cung cấp high-level emotion representation. Giữ frozen để tiết kiệm GPU và giảm overfit: https://arxiv.org/abs/2312.15185"
            ],
            [
                "Stats MLP",
                "Nhận vector thống kê thủ công rồi nén thành embedding phụ 128 chiều.",
                "[B,796] -> [B,128]",
                "Handcrafted/global statistics bổ sung góc nhìn ổn định cho fusion và stacking. Feature engineering SER: https://arxiv.org/abs/2507.07046"
            ],
            [
                "Co-attention",
                "z_e2v làm query; z_temporal và z_spectral làm key/value. Attention học temporal/spectral cue nào phù hợp với emotion embedding.",
                "q [B,1,192], kv [B,2,192] -> z_context [B,192]",
                "Dựa trên ý tưởng co-attention/multi-level fusion của CA-MSER: https://arxiv.org/abs/2203.15326"
            ],
            [
                "Fusion MLP",
                "Concat 5 vector rồi học mapping phi tuyến về fused representation.",
                "[B,896] -> [B,192]",
                "Kết hợp bốn nhánh và context attention thành một vector quyết định chung."
            ],
            [
                "Classifier",
                "Linear layer biến fused representation thành 6 logits; softmax đổi logits thành xác suất.",
                "[B,192] -> [B,6]",
                "Tạo xác suất cho 6 emotion: neutral, happy, sad, angry, fear, disgust."
            ],
        ],
        widths=[1.15, 2.35, 1.55, 2.25],
        font_size=6.5,
    )
    add_p(
        doc,
        "So với ví dụ BiGRU/dilated Conv1D, 06D hiện tại chọn BiLSTM vì code đã ổn định và có reference SER dùng BiLSTM. Nếu cần future work, có thể thử thay BiLSTM bằng BiGRU để nhẹ hơn, hoặc thêm dilated Conv1D để tăng receptive field mà không tăng nhiều tham số."
    )


def add_implementation_architecture_walkthrough_v2(doc):
    add_heading(doc, "3.2.1. Code kiến trúc 06D theo từng khối độc lập", 2)
    add_p(
        doc,
        "Phần này trình bày implementation đúng theo notebook 06D nhưng gom lại theo từng block lớn. Mỗi block gồm: code-level architecture, feature/input đi vào block, cách block hoạt động, và reference liên quan. Cách viết này giúp người đọc đối chiếu thẳng từ sơ đồ kiến trúc sang notebook."
    )
    add_p(
        doc,
        "Ký hiệu shape: B là batch size, tức số audio được xử lý cùng lúc trong một batch; T là số frame thời gian sau khi chia audio thành các cửa sổ ngắn; D là số chiều embedding. Ví dụ [B,132,T] nghĩa là một batch có B mẫu, mỗi mẫu là chuỗi thời gian gồm T frame, mỗi frame có 132 đặc trưng acoustic."
    )

    add_heading(doc, "3.2.1.1. Input, split và feature cache", 3)
    add_code(
        doc,
        "metadata.csv + audio_16k/*.wav\n"
        "  -> load_audio_fixed(sr=16000, duration=3s)\n"
        "  -> split protocol: combined_random / combined_strict_no_tess / single-dataset\n"
        "  -> extract features once and save cache\n\n"
        "X_temporal : [B, 132, T]      # Branch A\n"
        "X_spectral : [B, 3, 96, T]    # Branch B\n"
        "X_e2v      : [B, D_e2v]       # Branch C\n"
        "X_stats    : [B, 796]         # Branch D\n"
        "labels     : [B]              # 6 emotion classes"
    )
    add_table(
        doc,
        ["Thành phần", "Feature/Input", "Giải thích", "Reference"],
        [
            ["Audio 16 kHz", "waveform 1-D", "Tín hiệu giọng nói sau khi chuẩn hóa sample rate. 06D pad/trim về 3 giây để tạo batch ổn định.", "librosa load/resample: https://librosa.org/doc/latest/generated/librosa.load.html"],
            ["Split trước khi train", "train/validation/test indices", "Chia dữ liệu trước rồi mới train/fit scaler/stacking. Validation dùng chọn epoch và stacking; test chỉ báo cáo cuối để tránh leakage.", "Good ML evaluation practice; scikit-learn model selection: https://scikit-learn.org/stable/model_selection.html"],
            ["Feature cache", "npz arrays", "Lưu feature đã trích xuất để chạy lại không phải extract audio/emotion2vec từ đầu.", "Notebook 06D: build_or_load_feature_cache"],
            ["Label space", "6 labels", "Dùng 6 cảm xúc chung: neutral, happy, sad, angry, fear, disgust để gộp 4 dataset.", "Roadmap/project preprocessing"],
        ],
        widths=[1.35, 1.35, 3.05, 1.85],
        font_size=7.4,
    )

    add_heading(doc, "3.2.1.2. Branch A - Temporal acoustic branch", 3)
    add_code(
        doc,
        "Input: X_temporal [B, 132, T]\n"
        "132 = 40 MFCC + 40 delta MFCC + 40 delta-delta MFCC\n"
        "    + 1 RMS + 1 ZCR + 1 centroid + 1 bandwidth + 1 rolloff + 7 spectral contrast\n\n"
        "TemporalBranch:\n"
        "  Conv1d(132 -> 128, kernel=7, padding=3, bias=False)\n"
        "  BatchNorm1d(128) -> GELU -> Dropout\n"
        "  Conv1d(128 -> 160, kernel=5, padding=2, bias=False)\n"
        "  BatchNorm1d(160) -> GELU -> MaxPool1d(2) -> Dropout\n"
        "  transpose: [B, 160, T/2] -> [B, T/2, 160]\n"
        "  BiLSTM(input_size=160, hidden_size=96, bidirectional=True)\n"
        "  AttentionPooling: Linear(192->96) -> Tanh -> Linear(96->1) -> softmax(time)\n"
        "  Projection: LayerNorm(192) -> Linear(192->192) -> GELU -> Dropout\n"
        "Output: z_temporal [B, 192]"
    )
    add_p(
        doc,
        "Feature của Branch A là chuỗi acoustic theo thời gian. MFCC mô tả timbre/spectral envelope, tức màu giọng và hình dạng phổ đã được nén theo thang Mel. Delta MFCC là tốc độ thay đổi của MFCC theo thời gian; delta-delta MFCC là gia tốc thay đổi. RMS/energy mô tả độ lớn tiếng, ZCR mô tả độ sắc/voiced-unvoiced, còn centroid/bandwidth/rolloff/contrast mô tả độ sáng, độ rộng và texture của phổ. Vì Branch A có BiLSTM và attention theo trục thời gian, các feature này được giữ ở dạng sequence [132,T] thay vì chỉ lấy mean/std."
    )
    add_table(
        doc,
        ["Khối trong Branch A", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng / reference"],
        [
            ["Conv1D", "Kernel 1 chiều trượt dọc theo trục thời gian để học pattern cục bộ như năng lượng tăng, MFCC đổi nhanh, hoặc ZCR/spectral cue xuất hiện trong vài frame.", "[B,132,T] -> [B,128,T] -> [B,160,T]", "Phù hợp chuỗi MFCC/acoustic feature và nhẹ hơn 2D-CNN. Ahmed et al. dùng 1D-CNN/CNN-LSTM/CNN-GRU cho SER: https://arxiv.org/abs/2112.05666"],
            ["BatchNorm + GELU + Dropout", "Ổn định phân phối activation, thêm phi tuyến mềm và giảm overfit bằng cách tắt ngẫu nhiên một phần activation khi train.", "feature maps -> regularized maps", "Các block deep learning tiêu chuẩn để train ổn định trên dataset nhỏ."],
            ["MaxPool1d", "Giữ activation nổi bật trong vùng ngắn và giảm số frame thời gian.", "[B,160,T] -> [B,160,T/2]", "Giảm nhiễu và giảm chi phí trước khi đưa vào recurrent layer."],
            ["BiLSTM", "Forward LSTM đọc từ đầu câu tới hiện tại; backward LSTM đọc từ cuối câu về hiện tại; hai hướng được concat.", "[B,T/2,160] -> [B,T/2,192]", "Cảm xúc phụ thuộc ngữ cảnh cả câu. CA-MSER dùng CNN+BiLSTM trong SER: https://arxiv.org/abs/2203.15326; DCRF-BiLSTM dùng BiLSTM cho SER: https://arxiv.org/abs/2507.07046"],
            ["Attention pooling", "Tính score từng frame, softmax thành trọng số, rồi lấy tổng có trọng số theo thời gian.", "[B,T/2,192] -> [B,192]", "Không phải mọi frame đều chứa cảm xúc rõ. Attention giúp tập trung vào đoạn phát âm giàu thông tin: https://arxiv.org/abs/1706.03762"],
            ["Projection", "Chuẩn hóa và chiếu vector về 192 chiều.", "[B,192] -> z_temporal [B,192]", "192 là embedding dimension chung để fusion với z_spectral và z_e2v."],
        ],
        widths=[1.25, 2.45, 1.45, 2.55],
        font_size=6.8,
    )

    add_heading(doc, "3.2.1.3. Branch B - Spectrogram CNN-SE branch", 3)
    add_code(
        doc,
        "Input: X_spectral [B, 3, 96, T]\n"
        "3 channels = log-Mel + delta log-Mel + delta-delta log-Mel\n"
        "96 = number of Mel bins\n\n"
        "SpectralBranch:\n"
        "  ResidualSEBlock(3 -> 32) -> MaxPool2d(2)\n"
        "  ResidualSEBlock(32 -> 64) -> MaxPool2d(2)\n"
        "  ResidualSEBlock(64 -> 128)\n"
        "  AdaptiveAvgPool2d((1,1)) -> Flatten\n"
        "  LayerNorm(128) -> Linear(128->192) -> GELU -> Dropout\n"
        "Output: z_spectral [B, 192]"
    )
    add_p(
        doc,
        "Branch B dùng log-Mel như một ảnh time-frequency: trục ngang là thời gian, trục dọc là Mel bin/tần số cảm nhận, màu là năng lượng sau phép log. Delta log-Mel và delta-delta log-Mel bổ sung tốc độ/gia tốc thay đổi của năng lượng trên từng vùng tần số. Khác với MFCC đã nén phổ thành cepstral coefficients, log-Mel giữ vị trí tần số rõ hơn nên phù hợp với 2D-CNN."
    )
    add_table(
        doc,
        ["Khối trong Branch B", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng / reference"],
        [
            ["Log-Mel 3 kênh", "Tạo ảnh phổ gồm log-Mel, delta log-Mel, delta-delta log-Mel.", "[B,3,96,T]", "Giữ cấu trúc thời gian-tần số cho CNN. Librosa Mel: https://librosa.org/doc/latest/generated/librosa.feature.melspectrogram.html"],
            ["2D-CNN", "Kernel 2D quét trên mặt phẳng time-frequency để học pattern cục bộ như formant, vùng năng lượng sáng/tối, hoặc texture phổ.", "[B,3,96,T] -> feature maps", "Spectrogram thường được xử lý như ảnh acoustic. AST cũng dùng spectrogram patch làm input: https://arxiv.org/abs/2104.01778"],
            ["Residual block", "Học F(x) rồi cộng shortcut x + F(x), giúp giữ thông tin gốc và train sâu ổn định hơn.", "maps -> residual maps", "ResNet: https://arxiv.org/abs/1512.03385"],
            ["SE2D/channel attention", "Squeeze global info của feature map, tạo trọng số channel bằng sigmoid, rồi nhân lại vào từng channel.", "maps -> channel-reweighted maps", "Nhấn mạnh feature map có cue cảm xúc mạnh. SENet: https://arxiv.org/abs/1709.01507"],
            ["Global pooling + projection", "Gom toàn bộ feature map thành vector cố định rồi chiếu về 192 chiều.", "[B,128,H,W] -> z_spectral [B,192]", "Fusion cần vector cố định, cùng kích thước với Branch A/C."],
        ],
        widths=[1.25, 2.45, 1.45, 2.55],
        font_size=6.8,
    )

    add_heading(doc, "3.2.1.4. Branch C - Frozen emotion2vec branch", 3)
    add_code(
        doc,
        "Input: raw waveform 16 kHz\n"
        "  -> frozen emotion2vec_base encoder\n"
        "  -> X_e2v [B, D_e2v]\n\n"
        "Emotion2Vec adapter MLP:\n"
        "  Linear(D_e2v -> 384) -> LayerNorm -> GELU -> Dropout\n"
        "  Linear(384 -> 192) -> GELU -> Dropout\n"
        "Output: z_e2v [B, 192]\n"
        "Extra heads: emotion2vec_logreg, emotion2vec_mlp -> p_e2v [B,6]"
    )
    add_p(
        doc,
        "Branch C đưa waveform vào pretrained emotion2vec để lấy high-level emotion representation. Trong 06D, emotion2vec được giữ frozen, nghĩa là trọng số pretrained không cập nhật khi train. Chỉ adapter MLP và classifier phụ được train. Cách này phù hợp tài nguyên Kaggle/Colab và giảm nguy cơ overfit vào speaker/dataset style khi dataset SER không quá lớn."
    )
    add_table(
        doc,
        ["Khối trong Branch C", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng / reference"],
        [
            ["emotion2vec encoder", "Model pretrained đọc waveform và xuất embedding cảm xúc mức cao.", "waveform -> X_e2v [B,D_e2v]", "emotion2vec học representation cho speech emotion: https://arxiv.org/abs/2312.15185"],
            ["Frozen setting", "Không fine-tune encoder, chỉ dùng nó như extractor.", "fixed encoder weights", "Tiết kiệm GPU, tránh học thuộc speaker/dataset khi dữ liệu nhỏ."],
            ["Adapter MLP", "Nén/biến đổi embedding gốc về không gian 192-D.", "[B,D_e2v] -> z_e2v [B,192]", "Đưa Branch C về cùng kích thước với Branch A/B để co-attention và fusion."],
            ["p_e2v phụ", "Logistic regression/MLP tạo xác suất trực tiếp từ embedding emotion2vec.", "X_e2v -> p_e2v [B,6]", "Dùng như model phụ trong stacking, tương tự ensemble theo validation."],
        ],
        widths=[1.25, 2.45, 1.45, 2.55],
        font_size=6.8,
    )

    add_heading(doc, "3.2.1.5. Branch D - Statistical branch và RBF-SVM", 3)
    add_code(
        doc,
        "Input: temporal features [132, T]\n"
        "For each of 132 channels:\n"
        "  mean, std, min, max, median, IQR\n"
        "Extra waveform stats:\n"
        "  energy, max_abs, waveform_std, mean_abs\n"
        "X_stats size: 132*6 + 4 = 796\n\n"
        "Stats MLP:\n"
        "  Linear(796 -> 256) -> LayerNorm -> GELU -> Dropout\n"
        "  Linear(256 -> 128) -> GELU -> Dropout\n"
        "Output: z_stats [B, 128]\n\n"
        "RBF-SVM side model:\n"
        "  StandardScaler -> SVC(kernel='rbf', probability=True)\n"
        "Output: p_stats [B, 6]"
    )
    add_p(
        doc,
        "Branch D không giữ toàn bộ chuỗi thời gian mà tóm tắt mỗi feature channel thành thống kê toàn clip. Mean cho biết mức trung bình, std/IQR cho biết độ dao động, min/max cho biết biên độ cực trị, median giúp ổn định hơn khi có outlier. RBF-SVM được dùng như một model cổ điển để khai thác feature thủ công; tuy nhiên trong kết quả 06D, nhánh này đang yếu nhất, nghĩa là feature thống kê hiện chưa đủ giàu hoặc chưa được chọn/lọc tốt như các paper feature engineering mạnh."
    )
    add_table(
        doc,
        ["Khối trong Branch D", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng / reference"],
        [
            ["Statistical extraction", "Lấy mean/std/min/max/median/IQR theo trục thời gian cho từng acoustic channel.", "[132,T] -> [796]", "Biến sequence thành vector cố định cho MLP/SVM; feature engineering SER thường dùng thống kê thủ công."],
            ["StandardScaler", "Chuẩn hóa từng chiều feature về scale tương đương.", "X_stats -> scaled X_stats", "SVM RBF dựa trên khoảng cách nên rất nhạy với scale. SVC: https://scikit-learn.org/stable/modules/generated/sklearn.svm.SVC.html"],
            ["Stats MLP", "Học projection phi tuyến từ stats vector sang embedding phụ.", "[B,796] -> z_stats [B,128]", "Cho stats tham gia deep fusion dưới dạng vector."],
            ["RBF-SVM", "Kernel RBF đo tương đồng phi tuyến giữa các vector thống kê.", "X_stats -> p_stats [B,6]", "Dùng như classical side model cho stacking; liên hệ hướng handcrafted feature trong DCRF-BiLSTM: https://arxiv.org/abs/2507.07046"],
        ],
        widths=[1.25, 2.45, 1.45, 2.55],
        font_size=6.8,
    )

    add_heading(doc, "3.2.1.6. Fusion, co-attention, classifier và stacking", 3)
    add_code(
        doc,
        "Inputs:\n"
        "  z_temporal [B,192]\n"
        "  z_spectral [B,192]\n"
        "  z_e2v      [B,192]\n"
        "  z_stats    [B,128]\n\n"
        "EmotionGuidedCoAttention:\n"
        "  q  = z_e2v.unsqueeze(1)                       # [B,1,192]\n"
        "  kv = stack([z_temporal, z_spectral], dim=1)   # [B,2,192]\n"
        "  z_context = MultiheadAttention(q, kv, kv)     # [B,192]\n\n"
        "Fusion:\n"
        "  z = concat([z_temporal, z_spectral, z_e2v, z_context, z_stats])\n"
        "  z shape = [B, 192+192+192+192+128] = [B,896]\n"
        "  FusionMLP: Linear(896->384) -> LayerNorm -> GELU -> Dropout -> Linear(384->192)\n"
        "  Classifier: Linear(192->6) -> logits -> softmax -> p_deep [B,6]\n\n"
        "Validation-weighted stacking:\n"
        "  p_final = w1*p_deep + w2*p_stats + w3*p_e2v_logreg + w4*p_e2v_mlp\n"
        "  weights are selected on validation, not on test"
    )
    add_p(
        doc,
        "Trong co-attention, query có thể hiểu là câu hỏi, còn key/value là nơi chứa thông tin để trả lời. 06D dùng z_e2v làm query vì đây là embedding pretrained có cue cảm xúc tổng quát; nó hỏi hai nhánh acoustic z_temporal và z_spectral xem cue nào phù hợp. Sau đó model concat cả embedding gốc và context mới để classifier ra p_deep. Cuối cùng, stacking kết hợp p_deep với p_stats và p_e2v. Stacking khác với gated fusion ở chỗ nó kết hợp xác suất ở tầng quyết định cuối, thường dễ kiểm soát và dễ báo cáo hơn."
    )
    add_table(
        doc,
        ["Khối cuối", "Hoạt động ra sao?", "Input -> Output", "Vì sao dùng / reference"],
        [
            ["Co-attention", "z_e2v làm query, z_temporal/z_spectral làm key-value; attention học nhánh acoustic nào hữu ích với emotion cue.", "q [B,1,192], kv [B,2,192] -> z_context [B,192]", "Dựa trên ý tưởng co-attention/multi-level acoustic fusion trong CA-MSER: https://arxiv.org/abs/2203.15326"],
            ["Fusion MLP", "Concat nhiều representation rồi học mapping phi tuyến về fused embedding.", "[B,896] -> [B,192]", "Cho mô hình học quan hệ giữa temporal, spectral, pretrained và stats thay vì vote cứng."],
            ["Classifier", "Linear biến fused embedding thành 6 logits; softmax đổi logits thành xác suất.", "[B,192] -> p_deep [B,6]", "Output xác suất cho 6 emotion chung."],
            ["Stacking", "Kết hợp xác suất từ deep model, SVM và emotion2vec heads theo validation.", "nhiều p[6] -> p_final[6]", "Ahmed et al. dùng weighted ensemble giữa các biến thể CNN/recurrent: https://arxiv.org/abs/2112.05666"],
        ],
        widths=[1.25, 2.45, 1.45, 2.55],
        font_size=6.8,
    )

    add_heading(doc, "3.2.1.7. Train/evaluation output trong notebook", 3)
    add_table(
        doc,
        ["Bước", "Trong notebook 06D", "Ý nghĩa"],
        [
            ["Train deep model", "train_deep_model + run_epoch", "Train co-attention model, chọn best epoch bằng validation macro-F1."],
            ["Train side models", "stats_rbf_svm, emotion2vec_logreg, emotion2vec_mlp", "Tạo xác suất phụ cho stacking."],
            ["Report", "metrics.csv, per_dataset.csv, predictions.csv", "Lưu accuracy, macro-F1, UAR, confusion matrix và prediction từng mẫu."],
            ["Confusion matrix", "plot_confusion_matrix(y_true, y_pred)", "Cho biết lớp thật ở hàng, lớp dự đoán ở cột; đường chéo càng đậm càng tốt."],
        ],
        widths=[1.35, 2.65, 3.0],
        font_size=7.6,
    )


def add_15_method_selection_deep_justification(doc):
    add_heading(doc, "1.5.3. Giải thích sâu lý do chọn từng method/block trong 06D", 3)
    add_p(
        doc,
        "Cách chọn kiến trúc 06D đi từ câu hỏi chính của bài toán: cảm xúc trong giọng nói nằm ở đâu? Câu trả lời là không nằm ở một điểm duy nhất. Nó nằm trong cường độ giọng, màu giọng, tốc độ thay đổi của phổ, cấu trúc thời gian - tần số, đặc điểm toàn câu và cả high-level emotion representation đã học từ dữ liệu lớn. Vì vậy 06D không chọn một model đơn, mà chọn multi-representation architecture. Bảng dưới đây giải thích từng block được chọn vì lý do gì và lấy ý tưởng từ nguồn nào."
    )
    add_table(
        doc,
        ["Block/method trong 06D", "Thuật ngữ cần hiểu", "Vì sao chọn block này?", "Nguồn tham khảo / link", "Hình minh họa trong doc"],
        [
            [
                "Preprocess audio 16 kHz",
                "Resample, mono, pad/trim",
                "Chuẩn hóa định dạng audio để feature extraction ổn định. Nếu mỗi file khác sample rate/độ dài, batch training sẽ khó kiểm soát và feature theo frame không đồng nhất.",
                "Audio ML preprocessing; emotion2vec nhận waveform speech 16 kHz: https://arxiv.org/abs/2312.15185",
                "Hình 1, Hình 2"
            ],
            [
                "MFCC + delta + delta-delta",
                "MFCC = cepstral coefficients; delta = tốc độ thay đổi; delta-delta = gia tốc thay đổi",
                "MFCC mô tả timbre/spectral envelope; delta và delta-delta bổ sung dynamics. SER cần dynamics vì cùng một timbre nhưng cách biến đổi theo thời gian có thể khác cảm xúc.",
                "librosa MFCC: https://librosa.org/doc/latest/generated/librosa.feature.mfcc.html; librosa delta: https://librosa.org/doc/latest/generated/librosa.feature.delta.html",
                "Hình 3"
            ],
            [
                "Temporal Branch: 1D-CNN",
                "1D convolution, kernel, local receptive field, feature map",
                "1D-CNN quét theo trục thời gian của sequence [132,T] để bắt pattern ngắn: năng lượng tăng, MFCC đổi nhanh, ZCR/spectral cue xuất hiện trong vài frame.",
                "PyTorch Conv1d: https://docs.pytorch.org/docs/stable/generated/torch.nn.Conv1d.html; Ahmed et al. LFAB 1D-CNN: https://arxiv.org/abs/2112.05666",
                "Hình 1, Hình 3, Hình 5"
            ],
            [
                "Temporal Branch: BiLSTM",
                "Bidirectional LSTM, forward context, backward context, hidden state",
                "1D-CNN bắt pattern cục bộ; BiLSTM đặt pattern đó vào ngữ cảnh trước/sau của câu nói. Cảm xúc thường phụ thuộc diễn tiến cả câu, không chỉ một đoạn ngắn.",
                "PyTorch LSTM: https://docs.pytorch.org/docs/stable/generated/torch.nn.LSTM.html; Ahmed et al. CNN-LSTM/GRU: https://arxiv.org/abs/2112.05666",
                "Hình 1"
            ],
            [
                "Attention pooling",
                "Attention weight, weighted sum, important frame",
                "Không phải frame nào cũng giàu cảm xúc. Attention học trọng số frame, giúp model tập trung vào đoạn có tín hiệu cảm xúc rõ hơn thay vì average toàn bộ sequence.",
                "Attention Is All You Need: https://arxiv.org/abs/1706.03762; PyTorch MultiheadAttention: https://docs.pytorch.org/docs/stable/generated/torch.nn.MultiheadAttention.html",
                "Hình 1"
            ],
            [
                "Spectrogram Branch: log-Mel + 2D-CNN",
                "log-Mel, mel bin, time-frequency image, 2D convolution",
                "Log-Mel giữ cấu trúc thời gian - tần số như một ảnh; 2D-CNN học pattern cục bộ theo cả thời gian và tần số, ví dụ vùng formant mạnh hoặc vùng phổ thay đổi nhanh.",
                "librosa melspectrogram: https://librosa.org/doc/latest/generated/librosa.feature.melspectrogram.html; AST spectrogram patch idea: https://arxiv.org/abs/2104.01778",
                "Hình 1, Hình 6"
            ],
            [
                "Residual CNN + SE attention",
                "Residual connection, channel attention, squeeze-and-excitation",
                "Residual connection giúp train CNN sâu ổn định hơn; SE học trọng số cho từng channel feature map, nhấn mạnh channel chứa cue cảm xúc rõ.",
                "ResNet: https://arxiv.org/abs/1512.03385; SENet: https://arxiv.org/abs/1709.01507",
                "Hình 1, Hình 6"
            ],
            [
                "Frozen emotion2vec branch",
                "Pretrained embedding, frozen encoder, adapter MLP",
                "Feature thủ công có giới hạn. emotion2vec đã học representation cảm xúc từ dữ liệu lớn, nên dùng làm high-level cue. Frozen encoder giúp tiết kiệm GPU và giảm rủi ro overfit dataset nhỏ.",
                "emotion2vec: https://arxiv.org/abs/2312.15185; ModelScope: https://www.modelscope.cn/models/iic/emotion2vec_base",
                "Hình 1, Hình 2"
            ],
            [
                "Statistical branch + RBF-SVM",
                "Handcrafted statistics, StandardScaler, kernel trick, probability",
                "Stats branch tóm tắt toàn clip bằng mean/std/min/max/median/IQR. RBF-SVM là model cổ điển mạnh cho vector thủ công, bổ sung một hướng quyết định khác cho deep model.",
                "Chowdhury et al. feature engineering: https://arxiv.org/abs/2507.07046; scikit-learn SVC: https://scikit-learn.org/stable/modules/generated/sklearn.svm.SVC.html",
                "Hình 1, Hình 5"
            ],
            [
                "Co-attention + stacking",
                "Query/key/value, co-attention, probability ensemble, validation weighting",
                "Co-attention cho emotion2vec embedding hỏi temporal/spectral cue nào quan trọng. Stacking kết hợp p_deep, p_stats và p_e2v để giảm phụ thuộc vào một nhánh đơn.",
                "CA-MSER co-attention: https://arxiv.org/abs/2203.15326; Ahmed weighted ensemble: https://arxiv.org/abs/2112.05666",
                "Hình 1"
            ],
        ],
        widths=[1.15, 1.25, 2.05, 2.0, 0.9],
        font_size=6.7,
    )
    add_p(
        doc,
        "Nhìn theo bảng trên, mỗi thuật ngữ trong 06D đều có vai trò cụ thể. 1D-CNN không được đưa vào chỉ vì nó là deep learning, mà vì input Branch A là chuỗi feature theo thời gian. 2D-CNN không dùng cho MFCC sequence, mà dùng cho log-Mel vì log-Mel là bản đồ time-frequency. BiLSTM không thay thế CNN, mà bổ sung ngữ cảnh dài hạn. emotion2vec không thay thế feature thủ công, mà cung cấp high-level pretrained representation. RBF-SVM không phải main deep model, mà là nhánh phụ cho handcrafted statistics và stacking."
    )


def add_31_block_code_deep_explanation(doc):
    add_heading(doc, "3.1.3. Giải thích từng block code trong kiến trúc 06D", 3)
    add_p(
        doc,
        "Phần này đọc từ trái sang phải theo Hình 1: audio đầu vào -> preprocess -> feature extraction -> bốn branch -> co-attention/fusion -> classifier/stacking -> output. Mỗi block bên dưới đều gắn với code trong notebook 06D, input/output tensor, thuật ngữ chính, lý do dùng và reference."
    )

    add_heading(doc, "3.1.3.1. Input pipeline và split protocol", 4)
    add_table(
        doc,
        ["Block code", "Input -> Output", "Thuật ngữ", "Giải thích và lý do dùng"],
        [
            [
                "find_ser_processed(), read metadata.csv",
                "ser_processed folder -> metadata dataframe",
                "metadata, label space, dataset_id, speaker_id",
                "Metadata cho biết mỗi file thuộc dataset nào, speaker nào và emotion nào. Đây là nền để chia combined_random, strict và per-dataset."
            ],
            [
                "load_audio_fixed(path)",
                "wav file -> waveform 16 kHz fixed length",
                "resample, mono, pad/trim",
                "Đưa audio về cùng chuẩn để MFCC/log-Mel/RMS/ZCR có frame alignment ổn định. Nếu không chuẩn hóa, mỗi feature có thể có số frame khác nhau."
            ],
            [
                "make_protocol_splits()",
                "metadata -> train/validation/test indices",
                "combined_random, combined_strict_no_tess, single dataset",
                "combined_random dùng để so sánh với nhiều paper random split; strict dùng để kiểm tra generalization khi giảm speaker overlap; single dataset dùng để so với paper báo cáo riêng từng corpus."
            ],
        ],
        widths=[1.6, 1.45, 1.5, 2.55],
        font_size=7.2,
    )
    add_code(
        doc,
        "metadata = pd.read_csv(SER_PROCESSED_DIR / 'metadata.csv')\n"
        "metadata = metadata[metadata['emotion'].isin(COMMON_EMOTIONS)]\n"
        "y = load_audio_fixed(audio_path, sr=16000, duration=3.0)\n"
        "split_map = build_split(metadata, protocol='combined_random' or 'combined_strict_no_tess')"
    )

    add_heading(doc, "3.1.3.2. Unified feature extraction block", 4)
    add_p(
        doc,
        "Unified feature extraction nghĩa là từ một waveform, notebook tạo ra đầy đủ feature cho cả bốn nhánh. Cách này giúp các branch nhìn cùng một sample nhưng ở những representation khác nhau."
    )
    add_table(
        doc,
        ["Feature output", "Shape", "Dùng cho branch", "Đặc trưng bên trong", "Lý do chọn"],
        [
            ["X_temporal", "[132,T]", "Branch A", "40 MFCC + 40 delta + 40 delta2 + RMS + ZCR + centroid + bandwidth + rolloff + 7 contrast", "Giữ sequence theo thời gian để 1D-CNN/BiLSTM học dynamics."],
            ["X_spectral", "[3,96,T]", "Branch B", "log-Mel + delta log-Mel + delta-delta log-Mel", "Giữ ảnh time-frequency để 2D-CNN học pattern theo thời gian và tần số."],
            ["X_e2v", "[D_e2v]", "Branch C", "frozen emotion2vec utterance embedding", "Cung cấp high-level emotion representation đã học trước."],
            ["X_stats", "[796]", "Branch D", "mean/std/min/max/median/IQR của 132 channels + waveform stats", "Tạo vector cố định cho MLP/RBF-SVM, bổ sung global descriptors."],
        ],
        widths=[1.0, 0.75, 1.0, 2.75, 1.65],
        font_size=7.1,
    )
    add_code(
        doc,
        "X_temporal = vstack([mfcc, delta, delta2, rms, zcr, centroid, bandwidth, rolloff, contrast])\n"
        "X_spectral = stack([logmel, delta(logmel), delta2(logmel)])\n"
        "X_e2v = emotion2vec_encoder(waveform_16k)\n"
        "X_stats = statistics_over_time(X_temporal)"
    )
    add_p(
        doc,
        "Hình minh họa: Hình 2 cho waveform/RMS, Hình 3 cho MFCC-delta-delta2, Hình 4 cho ZCR, Hình 5 cho spectral centroid/bandwidth/rolloff, Hình 6 cho log-Mel và spectral contrast."
    )

    add_heading(doc, "3.1.3.3. Branch A - Temporal acoustic branch", 4)
    add_table(
        doc,
        ["Sub-block", "Code-level", "Input -> Output", "Thuật ngữ", "Tác dụng"],
        [
            ["Input", "temporal tensor", "[B,132,T]", "B=batch size; 132=feature channels; T=time frames", "Mỗi sample là một chuỗi feature theo thời gian. Đây là lý do phải dùng mô hình temporal."],
            ["1D-CNN", "Conv1d(132,128,kernel=7)", "[B,132,T] -> [B,128,T]", "1D convolution, kernel, feature map", "Quét dọc thời gian để học pattern ngắn như tăng energy, đổi MFCC, tăng ZCR."],
            ["BatchNorm + GELU", "BatchNorm1d + GELU", "[B,128,T] -> [B,128,T]", "normalization, activation", "Ổn định phân phối feature và thêm phi tuyến để model học quan hệ phức tạp."],
            ["MaxPool", "MaxPool1d(2)", "[B,160,T] -> [B,160,T/2]", "downsampling", "Giảm số frame, giữ pattern nổi bật, giảm nhiễu và chi phí."],
            ["BiLSTM", "bidirectional=True", "[B,T/2,160] -> [B,T/2,192]", "forward/backward context", "Học quan hệ dài hạn theo hai chiều thời gian của câu nói."],
            ["AttentionPooling", "softmax over time", "[B,T/2,192] -> [B,192]", "attention weights", "Chọn frame giàu cảm xúc hơn thay vì average cứng toàn bộ câu."],
        ],
        widths=[1.0, 1.35, 1.2, 1.55, 2.05],
        font_size=6.9,
    )
    add_p(
        doc,
        "Tại sao các feature của Branch A không đưa sang Branch B? Vì Branch A cần sequence theo kênh [132,T] để convolution 1D và BiLSTM học diễn tiến theo thời gian. Branch B cần ảnh [3,96,T] để CNN 2D học pattern trên mặt phẳng thời gian - tần số. Hai dạng input phục vụ hai kiểu học khác nhau."
    )

    add_heading(doc, "3.1.3.4. Branch B - Spectrogram CNN-SE branch", 4)
    add_table(
        doc,
        ["Sub-block", "Code-level", "Input -> Output", "Thuật ngữ", "Tác dụng"],
        [
            ["Input", "spectral tensor", "[B,3,96,T]", "3 channels, 96 mel bins, T frames", "Ba kênh không phải RGB; đó là log-Mel, delta log-Mel, delta-delta log-Mel."],
            ["2D-CNN", "Conv2d over time-frequency", "[B,3,96,T] -> feature maps", "2D convolution, local patch", "Học vùng phổ cục bộ: formant, harmonic texture, vùng năng lượng theo thời gian."],
            ["Residual block", "x + F(x)", "feature maps -> feature maps", "skip connection", "Giữ thông tin gốc và giúp train CNN sâu ổn định hơn."],
            ["SE2D", "GlobalAvgPool -> Linear -> Sigmoid", "feature maps -> weighted maps", "channel attention", "Tính trọng số cho từng channel feature map, nhấn mạnh map hữu ích cho emotion."],
            ["Global pooling", "AdaptiveAvgPool2d", "maps -> vector", "spatial pooling", "Tóm tắt ảnh spectrogram thành embedding cố định z_spectral."],
        ],
        widths=[1.0, 1.35, 1.35, 1.45, 2.0],
        font_size=6.9,
    )
    add_p(
        doc,
        "Tại sao dùng log-Mel thay vì Mel thường? Vì năng lượng speech có dynamic range lớn; phép log nén thang năng lượng, giúp vùng yếu vẫn còn thông tin và CNN không bị vùng rất mạnh lấn át. Tại sao n_mels=96? Đây là mức đủ chi tiết cho time-frequency pattern nhưng vẫn nhẹ hơn 128 mel bins, phù hợp Kaggle/Colab."
    )

    add_heading(doc, "3.1.3.5. Branch C - Frozen emotion2vec branch", 4)
    add_table(
        doc,
        ["Sub-block", "Code-level", "Input -> Output", "Thuật ngữ", "Tác dụng"],
        [
            ["Waveform input", "raw waveform 16 kHz", "[samples] -> encoder input", "raw speech", "Giữ toàn bộ tín hiệu âm thanh để pretrained model tự rút representation."],
            ["Frozen emotion2vec", "AutoModel / emotion2vec_base", "waveform -> hidden states", "self-supervised pretrained encoder", "Dùng knowledge đã học trước về emotion-related speech representation."],
            ["Pooling", "mean over frames", "hidden states -> X_e2v", "utterance embedding", "Tóm tắt toàn câu nói thành vector cảm xúc mức cao."],
            ["Adapter MLP", "Linear -> GELU -> Dropout", "X_e2v -> z_e2v [B,192]", "projection/adaptation", "Đưa embedding về cùng không gian 192-D để fusion với các branch khác."],
        ],
        widths=[1.0, 1.45, 1.3, 1.55, 1.95],
        font_size=6.9,
    )
    add_p(
        doc,
        "Tại sao frozen thay vì fine-tune? Vì đồ án chạy trên Kaggle/Colab, dữ liệu SER không quá lớn, fine-tune toàn bộ pretrained model dễ tốn tài nguyên và overfit. Frozen embedding giữ giá trị ứng dụng thực tế hơn: khi demo, hệ thống chỉ cần rút embedding rồi chạy classifier nhẹ."
    )

    add_heading(doc, "3.1.3.6. Branch D - Statistical branch và RBF-SVM", 4)
    add_table(
        doc,
        ["Sub-block", "Code-level", "Input -> Output", "Thuật ngữ", "Tác dụng"],
        [
            ["Statistics", "mean/std/min/max/median/IQR", "[132,T] -> [792] + waveform stats", "global descriptor", "Tóm tắt phân bố toàn clip, không phụ thuộc số frame."],
            ["Scaler", "StandardScaler.fit(train)", "raw stats -> scaled stats", "standardization", "RBF-SVM dựa trên khoảng cách nên cần các feature cùng scale."],
            ["Stats MLP", "Linear layers", "[B,796] -> z_stats [B,128]", "bottleneck embedding", "Biến vector thủ công thành embedding nhỏ để ghép với deep fusion."],
            ["RBF-SVM", "SVC(kernel='rbf', probability=True)", "X_stats -> p_stats[6]", "kernel trick", "Tạo xác suất phụ từ handcrafted feature, dùng trong stacking."],
        ],
        widths=[1.0, 1.5, 1.45, 1.35, 1.9],
        font_size=6.9,
    )
    add_p(
        doc,
        "Tại sao không giữ nguyên sequence cho SVM? Vì SVM yêu cầu vector cố định cho mỗi sample. Do đó sequence [132,T] được tóm tắt thành các thống kê theo thời gian. Nhược điểm là mất thông tin thứ tự frame, nên Branch D chỉ là nhánh phụ, không thay thế Branch A."
    )

    add_heading(doc, "3.1.3.7. Fusion, classifier, stacking và output", 4)
    add_table(
        doc,
        ["Sub-block", "Code-level", "Input -> Output", "Thuật ngữ", "Tác dụng"],
        [
            ["Co-attention", "query=z_e2v; key/value=[z_temporal,z_spectral]", "3 embeddings -> z_context", "query/key/value", "Dùng emotion2vec làm câu hỏi để chọn acoustic cue quan trọng từ temporal/spectral."],
            ["Concat", "torch.cat([...], dim=1)", "192+192+192+192+128 -> 896", "feature fusion", "Gom representation từ bốn góc nhìn và context attention."],
            ["Fusion MLP", "Linear(896,384)->LayerNorm->GELU->Dropout->Linear(384,192)", "[B,896] -> [B,192]", "projection, regularization", "Nén vector ghép thành fused representation ổn định trước classifier."],
            ["Classifier", "Linear(192,6)", "[B,192] -> logits [B,6]", "logits, softmax", "Tạo xác suất cho 6 emotion: neutral, happy, sad, angry, fear, disgust."],
            ["Stacking", "combine p_deep, p_stats, p_e2v using validation", "probabilities -> p_final[6]", "validation-weighted ensemble", "Kết hợp nhánh deep và nhánh phụ, tránh chọn weight bằng test set."],
        ],
        widths=[1.0, 1.65, 1.35, 1.25, 2.0],
        font_size=6.9,
    )
    add_code(
        doc,
        "z_t = TemporalBranch(X_temporal)      # [B,132,T] -> [B,192]\n"
        "z_s = SpectralBranch(X_spectral)      # [B,3,96,T] -> [B,192]\n"
        "z_e = AdapterMLP(X_e2v)               # [B,D_e2v] -> [B,192]\n"
        "z_d = StatsMLP(X_stats)               # [B,796] -> [B,128]\n"
        "z_ctx = CoAttention(query=z_e, keys=[z_t, z_s])\n"
        "z = concat([z_t, z_s, z_e, z_ctx, z_d])  # [B,896]\n"
        "p_deep = softmax(Classifier(FusionMLP(z)))\n"
        "p_final = weighted_stack(p_deep, p_stats, p_e2v)"
    )


def add_15_31_compact_walkthrough(doc):
    add_heading(doc, "1.5 & 3.1. Proposed solution and architecture workflow", 1)
    add_p(
        doc,
        "Hai mục 1.5 và 3.1 được gộp thành một mạch: chọn mô hình đề xuất 06D, giải thích mô hình lấy ý tưởng từ đâu, rồi đi từng khối trong kiến trúc. Mỗi khối dưới đây trình bày đúng thứ tự trong sơ đồ: pipeline -> Branch A -> Branch B -> Branch C -> Branch D -> fusion/final decision -> output."
    )

    add_heading(doc, "1.5. Proposed solution: 06D Emotion2Vec Co-Attention SER", 2)
    add_p(
        doc,
        "Giải pháp được chọn là 06D Emotion2Vec Co-Attention SER. Lý do chọn mô hình này là cảm xúc trong giọng nói không nằm trong một feature đơn lẻ: MFCC thể hiện màu giọng, RMS/ZCR/spectral features thể hiện năng lượng và texture, log-Mel giữ cấu trúc thời gian - tần số, emotion2vec cung cấp representation cảm xúc đã học trước, còn statistical vector tóm tắt toàn clip. Vì vậy 06D dùng bốn nhánh biểu diễn và kết hợp chúng bằng co-attention + stacking."
    )
    add_table(
        doc,
        ["Ý tưởng chính", "06D dùng như thế nào?", "Lý do chọn", "Reference"],
        [
            ["Multi-representation SER", "Temporal + spectrogram + pretrained emotion embedding + statistics", "Một nguồn feature không đủ bao quát emotion signal.", "CA-MSER: https://arxiv.org/abs/2203.15326"],
            ["Temporal CNN/RNN", "1D-CNN + BiLSTM + attention trên [132,T]", "Học local acoustic pattern và diễn tiến theo thời gian.", "Ahmed et al.: https://arxiv.org/abs/2112.05666"],
            ["Spectrogram CNN-SE", "2D-CNN/ResidualSEBlock trên log-Mel 3 kênh", "Học pattern time-frequency và tự cân trọng số channel.", "ResNet: https://arxiv.org/abs/1512.03385; SENet: https://arxiv.org/abs/1709.01507"],
            ["Frozen pretrained branch", "emotion2vec embedding + adapter MLP", "Tận dụng pretrained emotion representation nhưng không fine-tune nặng.", "emotion2vec: https://arxiv.org/abs/2312.15185"],
            ["Stats + SVM/stacking", "Stats MLP + RBF-SVM probability", "Bổ sung hướng feature engineering cổ điển và xác suất phụ.", "Chowdhury et al.: https://arxiv.org/abs/2507.07046"],
        ],
        widths=[1.35, 2.0, 1.9, 1.75],
        font_size=7.1,
    )

    add_heading(doc, "3.1.1. Pipeline block: audio -> split -> feature extraction", 2)
    add_p(
        doc,
        "Pipeline nhận audio 16 kHz, đọc metadata, chia train/validation/test theo protocol, rồi mới trích xuất và chuẩn hóa feature. Split phải có trước augmentation/model selection để tránh data leakage: validation/test không được có bản augment hoặc scaler được fit từ chính chúng."
    )
    add_code(
        doc,
        "audio 16 kHz -> load_audio_fixed()\n"
        "metadata -> split_map(train/validation/test)\n"
        "waveform -> X_temporal [132,T], X_spectral [3,96,T], X_e2v [D_e2v], X_stats [796]\n"
        "train_idx -> fit scaler; validation/test -> transform only"
    )
    add_table(
        doc,
        ["Bên trong pipeline", "Giải thích", "Vì sao dùng / reference"],
        [
            ["Audio 16 kHz", "Resample, mono, pad/trim để mỗi sample có cùng chuẩn đầu vào.", "Ổn định số frame cho MFCC/log-Mel và phù hợp pretrained speech branch."],
            ["Split protocol", "combined_random để so sánh với nhiều paper random split; strict để kiểm tra generalization theo speaker/domain.", "Tránh đánh giá quá lạc quan do speaker overlap."],
            ["Feature extraction", "Tạo bốn representation từ cùng waveform: temporal, spectrogram, pretrained embedding, stats.", "Multi-level acoustic information theo hướng CA-MSER."],
            ["Train-only scaler/augmentation", "Scaler fit trên train; augmentation chỉ chạy khi self.train=True.", "Tránh leakage; SpecAugment masking tham khảo: https://arxiv.org/abs/1904.08779"],
        ],
        widths=[1.35, 2.65, 2.85],
        font_size=7.4,
    )
    add_image(doc, DOMAIN_DATASET_IMG, "Minh chứng domain/dataset: phân bố sample, emotion, duration và RMS khác nhau giữa RAVDESS, CREMA-D, TESS, SAVEE.", width=6.8)
    add_p(
        doc,
        "Cách đọc hình domain/dataset: Panel A có trục ngang là dataset/domain và trục dọc là số utterance, cho thấy CREMA-D lớn hơn hẳn các bộ còn lại. Panel B có trục ngang là dataset, trục dọc là emotion, màu/ô số là số mẫu mỗi emotion. Panel C có trục ngang là dataset và trục dọc là duration, cho thấy độ dài câu nói khác nhau giữa corpus. Panel D có trục ngang là dataset và trục dọc là RMS energy, cho thấy năng lượng/recording style cũng khác nhau. Đây là lý do bài toán cần combined_random, strict và per-dataset evaluation: nếu không, mô hình có thể học dấu hiệu domain/speaker thay vì emotion."
    )

    add_heading(doc, "3.1.2. Branch A - Temporal acoustic branch", 2)
    add_image(doc, EVIDENCE_WAVE_RMS, "Minh chứng Branch A: waveform và RMS trên 6 emotion.", width=6.6)
    add_image(doc, EVIDENCE_MFCC, "Minh chứng Branch A: MFCC, delta MFCC và delta-delta MFCC trên 6 emotion.", width=6.6)
    add_image(doc, EVIDENCE_ZCR, "Minh chứng Branch A: ZCR trên 6 emotion.", width=6.5)
    add_image(doc, EVIDENCE_SPECTRAL, "Minh chứng Branch A: spectral centroid, bandwidth và rolloff trên 6 emotion.", width=6.5)
    add_p(
        doc,
        "Branch A là nhánh temporal acoustic, tức nhánh học chuỗi đặc trưng âm học theo thời gian. Input của nhánh này là X_temporal [B,132,T]: B là batch size, T là số frame thời gian, 132 là số kênh feature. 132 kênh này gồm 40 MFCC, 40 delta MFCC, 40 delta-delta MFCC, RMS, ZCR, spectral centroid, bandwidth, rolloff và 7 spectral contrast. Vì Branch A giữ nguyên trục thời gian nên nó phù hợp với 1D-CNN, BiLSTM và attention pooling."
    )
    add_p(
        doc,
        "Trong hình minh họa MFCC/delta/delta-delta, trục ngang là thời gian, trục dọc là chỉ số hệ số MFCC, màu thể hiện giá trị hệ số. MFCC cho thấy hình dạng phổ/màu giọng ở từng frame. Delta làm nổi bật vùng MFCC thay đổi nhanh giữa các frame. Delta-delta cho thấy sự thay đổi đó đang tăng tốc hay chậm lại. Vì cảm xúc không chỉ là giọng ở một thời điểm mà là cách giọng biến đổi trong cả câu, ba nhóm này được đưa cùng vào Branch A."
    )
    add_table(
        doc,
        ["Hình/feature", "Trục ngang", "Trục dọc / màu", "Cách đọc khác biệt giữa 6 emotion"],
        [
            ["Waveform", "Thời gian của utterance.", "Biên độ âm thanh, dao động quanh 0.", "Đoạn biên độ lớn thường gợi cường độ mạnh; khoảng gần 0 là pause/silence. Angry/happy thường có cụm biên độ dày hơn, sad/neutral có thể mềm hơn."],
            ["RMS", "Frame thời gian.", "Năng lượng/loudness mỗi frame.", "RMS cho thấy lúc nào giọng mạnh/yếu. Nó giúp model thấy emotional intensity nhưng không quyết định một mình vì speaker/microphone cũng ảnh hưởng energy."],
            ["MFCC", "Frame thời gian.", "Chỉ số hệ số MFCC; màu là giá trị hệ số.", "MFCC mô tả timbre/spectral envelope. Hình khác nhau ở vân màu và vùng hệ số thay đổi, phản ánh màu giọng khác nhau giữa emotion."],
            ["Delta MFCC", "Frame thời gian.", "Chỉ số hệ số MFCC; màu là tốc độ thay đổi.", "Vùng màu mạnh cho thấy MFCC đổi nhanh. Điều này quan trọng cho cảm xúc có biến động giọng nhanh hoặc nhấn mạnh."],
            ["Delta-delta MFCC", "Frame thời gian.", "Chỉ số hệ số MFCC; màu là gia tốc thay đổi.", "Cho biết sự thay đổi MFCC đang tăng/giảm nhanh thế nào, hữu ích với các đoạn bùng năng lượng hoặc chuyển tone đột ngột."],
            ["ZCR", "Frame thời gian.", "Tỷ lệ zero-crossing.", "ZCR cao hơn khi tín hiệu có nhiều thành phần sắc/vô thanh/high-frequency. Nó giúp phân biệt texture giọng, nhưng phải học cùng MFCC/RMS."],
            ["Centroid/bandwidth/rolloff", "Frame thời gian.", "Giá trị tần số hoặc độ rộng phổ.", "Centroid thể hiện độ sáng; bandwidth thể hiện phổ hẹp/rộng; rolloff thể hiện năng lượng lan tới vùng cao. Các đường khác nhau cho thấy spectral shape khác nhau."],
            ["Spectral contrast", "Frame thời gian.", "7 dải tần; màu/giá trị là độ chênh peak-valley.", "Contrast cho biết cấu trúc phổ rõ/mờ ở từng dải tần, bổ sung thông tin texture cho Branch A."],
        ],
        widths=[1.25, 1.25, 1.65, 2.85],
        font_size=6.7,
    )
    add_code(
        doc,
        "X_temporal = vstack([mfcc40, delta40, delta2_40, rms, zcr, centroid, bandwidth, rolloff, contrast7])\n"
        "Conv1d blocks -> BiLSTM -> AttentionPooling -> z_temporal [B,192]"
    )
    add_table(
        doc,
        ["Thành phần", "Nó là gì?", "Tác dụng trong Branch A", "Reference"],
        [
            ["MFCC", "Mel-Frequency Cepstral Coefficients. Đây là các hệ số cepstral lấy từ log-Mel spectrum, thường dùng để mô tả spectral envelope/timbre.", "Màu giọng và cách cộng hưởng âm thay đổi theo emotion. Angry/happy/sad có thể khác nhau ở độ sáng, độ căng và envelope phổ.", "librosa MFCC: https://librosa.org/doc/latest/generated/librosa.feature.mfcc.html"],
            ["Delta MFCC", "Đạo hàm bậc 1 của MFCC theo thời gian, biểu diễn tốc độ thay đổi của từng hệ số MFCC.", "Cho model biết màu giọng đang thay đổi nhanh hay chậm. Đây là cue quan trọng vì cảm xúc thể hiện qua dynamics, không chỉ trạng thái tĩnh.", "librosa delta: https://librosa.org/doc/latest/generated/librosa.feature.delta.html"],
            ["Delta-delta MFCC", "Đạo hàm bậc 2 của MFCC, biểu diễn gia tốc thay đổi của MFCC.", "Bắt các chuyển động phổ nhanh/chậm lại, ví dụ giọng căng hoặc bùng năng lượng trong một đoạn ngắn.", "librosa delta order=2: https://librosa.org/doc/latest/generated/librosa.feature.delta.html"],
            ["RMS", "Root Mean Square energy theo frame.", "Thể hiện loudness/cường độ. Angry/happy thường có energy mạnh hơn; sad/neutral có thể thấp hơn, nhưng phải học cùng feature khác để tránh suy luận đơn giản.", "librosa RMS: https://librosa.org/doc/latest/generated/librosa.feature.rms.html"],
            ["ZCR", "Zero-Crossing Rate, số lần waveform đổi dấu trong một frame.", "Bổ sung cue về voiced/unvoiced, độ sắc, độ nhiễu và high-frequency texture.", "librosa ZCR: https://librosa.org/doc/latest/generated/librosa.feature.zero_crossing_rate.html"],
            ["Spectral centroid/bandwidth/rolloff/contrast", "Các đặc trưng hình dạng phổ: độ sáng, độ rộng, vùng năng lượng cao và peak-valley contrast.", "Giúp Branch A thấy giọng sáng/gắt/mềm/trầm, không chỉ dựa vào MFCC.", "Ullah et al. dùng acoustic feature fusion cho SER: https://doi.org/10.1109/ICIT56493.2022.9989197"],
            ["1D-CNN", "Convolution 1 chiều trượt trên trục thời gian của chuỗi [132,T].", "Bắt pattern cục bộ như energy tăng, MFCC đổi nhanh, hoặc spectral cue xuất hiện trong vài frame.", "Ahmed et al. dùng 1D-CNN/LSTM/GRU trong SER: https://arxiv.org/abs/2112.05666"],
            ["BiLSTM", "Bidirectional LSTM gồm forward LSTM đọc quá khứ -> hiện tại và backward LSTM đọc tương lai -> hiện tại.", "Cảm xúc phụ thuộc ngữ cảnh cả câu. BiLSTM giúp mỗi frame được hiểu dựa trên những gì xảy ra trước và sau nó.", "CA-MSER dùng CNN + BiLSTM cho multi-level acoustic SER: https://arxiv.org/abs/2203.15326; DCRF-BiLSTM: https://arxiv.org/abs/2507.07046"],
            ["Attention pooling", "Tính trọng số cho từng frame rồi lấy tổng có trọng số.", "Không phải frame nào cũng giàu cảm xúc. Attention chọn đoạn quan trọng thay vì average cứng toàn bộ câu.", "Attention: https://arxiv.org/abs/1706.03762"],
        ],
        widths=[1.15, 1.9, 2.45, 1.5],
        font_size=6.6,
    )
    add_image(doc, BILSTM_IMG, "Minh chứng BiLSTM: forward context và backward context giúp mỗi frame hiểu cả ngữ cảnh trước/sau.", width=6.7)
    add_p(
        doc,
        "BiLSTM phù hợp với SER vì câu nói là chuỗi thời gian. Nếu chỉ dùng 1D-CNN, model chủ yếu thấy pattern cục bộ trong một vùng ngắn. BiLSTM bổ sung long-range context: forward LSTM học diễn tiến từ đầu câu tới hiện tại, backward LSTM học thông tin từ cuối câu quay lại hiện tại. Khi concat hai hướng, mỗi frame có representation giàu hơn trước khi attention pooling chọn frame quan trọng."
    )
    add_p(
        doc,
        "z_temporal [B,192] là output cuối của Branch A sau attention pooling và projection. B vẫn là số mẫu trong batch; 192 không phải số feature gốc mà là embedding dimension chung. Nhóm chọn 192 để các nhánh temporal, spectral và emotion2vec có cùng kích thước khi fusion: đủ lớn để giữ thông tin cảm xúc, nhưng vẫn nhẹ hơn các embedding quá lớn trên Kaggle/Colab."
    )

    add_heading(doc, "3.1.3. Branch B - Spectrogram CNN-SE branch", 2)
    add_image(doc, EVIDENCE_LOGMEL, "Minh chứng Branch B: log-Mel spectrogram và spectral contrast trên 6 emotion.", width=6.6)
    add_p(
        doc,
        "Branch B dùng X_spectral [B,3,96,T]. Ba kênh không phải RGB mà là log-Mel, delta log-Mel và delta-delta log-Mel. Log-Mel giữ bản đồ năng lượng theo thời gian - tần số; delta/delta-delta log-Mel cho biết vùng tần số nào đang thay đổi nhanh. n_mels=96 là mức cân bằng giữa chi tiết phổ và chi phí tính toán trên Kaggle."
    )
    add_p(
        doc,
        "Cách đọc hình log-Mel: trục ngang là thời gian, trục dọc là tần số/Mel bin, màu là mức năng lượng sau phép log. Vùng sáng thể hiện dải tần có năng lượng mạnh; vùng tối là ít năng lượng hoặc silence. So với MFCC, log-Mel giữ vị trí tần số rõ hơn, nên phù hợp với 2D-CNN. Spectral contrast trong cùng hình cho thấy độ chênh giữa đỉnh và đáy phổ ở các dải tần, giúp mô tả texture phổ."
    )
    add_code(
        doc,
        "logmel = power_to_db(melspectrogram(y, n_mels=96))\n"
        "X_spectral = stack([logmel, delta(logmel), delta2(logmel)])  # [3,96,T]\n"
        "ResidualSEBlock + 2D-CNN -> GlobalPooling -> z_spectral [B,192]"
    )
    add_table(
        doc,
        ["Thành phần", "Nó là gì?", "Tác dụng", "Reference"],
        [
            ["Log-Mel", "Ảnh time-frequency trên thang Mel sau phép log.", "Giữ vùng năng lượng, formant và texture phổ rõ hơn MFCC.", "librosa melspectrogram: https://librosa.org/doc/latest/generated/librosa.feature.melspectrogram.html"],
            ["2D-CNN", "Convolution trên mặt phẳng thời gian - tần số.", "Học pattern cục bộ trong spectrogram.", "AST dùng spectrogram patches: https://arxiv.org/abs/2104.01778"],
            ["Residual block", "Shortcut x + F(x).", "Giữ thông tin gốc và giúp train CNN sâu ổn định.", "ResNet: https://arxiv.org/abs/1512.03385"],
            ["SE/channel attention", "Squeeze global info, excite trọng số từng channel.", "Nhấn mạnh feature map có cue cảm xúc rõ.", "SENet: https://arxiv.org/abs/1709.01507"],
        ],
        widths=[1.25, 1.9, 2.1, 1.75],
        font_size=7.0,
    )
    add_p(
        doc,
        "z_spectral [B,192] là embedding cuối của Branch B sau 2D-CNN, SE/channel attention và global pooling. Nó không còn là ảnh spectrogram nữa, mà là vector tóm tắt các pattern time-frequency quan trọng cho emotion. Số 192 được chọn để khớp với z_temporal và z_e2v khi fusion."
    )

    add_heading(doc, "3.1.4. Branch C - Frozen emotion2vec branch", 2)
    add_image(doc, EVIDENCE_WAVE_RMS, "Minh chứng Branch C: waveform là đầu vào raw speech cho pretrained emotion representation.", width=6.6)
    add_p(
        doc,
        "Branch C đưa waveform 16 kHz vào emotion2vec_base để lấy embedding cảm xúc mức cao. Trong 06D, emotion2vec được giữ frozen: chỉ adapter MLP phía sau được train. Cách này tiết kiệm GPU, giảm rủi ro overfit dataset nhỏ và phù hợp mục tiêu demo thực tế vì hệ thống chỉ cần rút embedding rồi chạy classifier nhẹ."
    )
    add_code(
        doc,
        "waveform -> frozen emotion2vec_base -> X_e2v [B,D_e2v]\n"
        "Adapter MLP: Linear -> LayerNorm -> GELU -> Dropout -> z_e2v [B,192]"
    )
    add_table(
        doc,
        ["Thành phần", "Giải thích", "Vì sao dùng", "Reference"],
        [
            ["emotion2vec", "Pretrained speech emotion representation.", "Bổ sung high-level emotion cue mà feature thủ công khó tự học.", "emotion2vec: https://arxiv.org/abs/2312.15185"],
            ["Frozen encoder", "Không cập nhật trọng số pretrained model.", "Giảm chi phí, giảm overfit, dễ triển khai demo.", "ModelScope: https://www.modelscope.cn/models/iic/emotion2vec_base"],
            ["Adapter MLP", "Project embedding về 192-D.", "Đưa Branch C về cùng không gian với Branch A/B để fusion.", "MLP adapter là cách nhẹ để adapt pretrained embeddings."],
        ],
        widths=[1.3, 2.0, 2.15, 1.55],
        font_size=7.2,
    )
    add_p(
        doc,
        "X_e2v [B,D_e2v] là embedding gốc từ emotion2vec, trong đó D_e2v phụ thuộc pretrained model. Adapter MLP biến nó thành z_e2v [B,192]. z_e2v đại diện cho high-level emotion cue đã học trước, nên trong co-attention nó được dùng như query để hỏi các nhánh acoustic nên chú ý vào đâu."
    )

    add_heading(doc, "3.1.5. Branch D - Statistical branch", 2)
    add_image(doc, EVIDENCE_SPECTRAL, "Minh chứng Branch D: spectral centroid, bandwidth và rolloff được tóm tắt thành thống kê toàn clip.", width=6.6)
    add_p(
        doc,
        "Branch D tạo X_stats [B,796] bằng cách lấy mean, std, min, max, median và IQR trên 132 temporal channels, rồi cộng thêm thống kê waveform. Nhánh này không giữ thứ tự frame như Branch A, nhưng cung cấp global descriptors ổn định cho MLP và RBF-SVM."
    )
    add_code(
        doc,
        "X_stats = concat([mean, std, min, max, median, IQR over 132 channels, waveform_stats])  # [796]\n"
        "StatsMLP -> z_stats [B,128]\n"
        "StandardScaler + RBF-SVM(probability=True) -> p_stats [B,6]"
    )
    add_table(
        doc,
        ["Thành phần", "Giải thích", "Tác dụng", "Reference"],
        [
            ["Mean/std/min/max/median/IQR", "Các thống kê mô tả phân bố feature theo toàn clip.", "Biến sequence thành vector cố định cho MLP/SVM.", "Feature engineering SER: https://arxiv.org/abs/2507.07046"],
            ["StandardScaler", "Chuẩn hóa từng chiều feature.", "RBF-SVM dựa trên khoảng cách nên cần cùng scale.", "scikit-learn SVC: https://scikit-learn.org/stable/modules/generated/sklearn.svm.SVC.html"],
            ["RBF-SVM", "Kernel classifier phi tuyến cho vector thủ công.", "Tạo p_stats phụ cho stacking.", "Classical SER baseline/feature-engineering approach."],
        ],
        widths=[1.5, 2.1, 2.0, 1.4],
        font_size=7.2,
    )
    add_p(
        doc,
        "X_stats [B,796] là vector thống kê còn z_stats [B,128] là embedding nén từ stats branch. 128 nhỏ hơn 192 vì stats là global descriptor phụ trợ, không phải nhánh sequence/spectrogram chính. Nhánh này cung cấp một góc nhìn ổn định cho stacking và giúp báo cáo giải thích được feature engineering thủ công."
    )

    add_heading(doc, "3.1.6. Fusion and final decision block", 2)
    add_p(
        doc,
        "Fusion nhận bốn embedding: z_temporal, z_spectral, z_e2v và z_stats. Co-attention dùng z_e2v làm query để hỏi hai nhánh acoustic xem temporal/spectral cue nào quan trọng. Sau đó mô hình concat z_temporal, z_spectral, z_e2v, z_context và z_stats thành vector 896-D, đưa qua Fusion MLP và classifier để tạo p_deep. Cuối cùng stacking kết hợp p_deep, p_stats và p_e2v bằng validation, không dùng test để chọn weight."
    )
    add_code(
        doc,
        "z_context = CoAttention(query=z_e2v, key/value=[z_temporal, z_spectral])\n"
        "z = concat([z_temporal, z_spectral, z_e2v, z_context, z_stats])  # [B,896]\n"
        "p_deep = softmax(Classifier(FusionMLP(z)))\n"
        "p_final = stack(p_deep, p_stats, p_e2v)"
    )
    add_table(
        doc,
        ["Thành phần", "Giải thích", "Tác dụng", "Reference"],
        [
            ["Co-attention", "Một embedding làm query, embedding khác làm key/value.", "Cho pretrained emotion cue tương tác với acoustic cue.", "CA-MSER: https://arxiv.org/abs/2203.15326"],
            ["Fusion MLP", "Linear + LayerNorm + GELU + Dropout.", "Nén vector 896-D thành representation quyết định.", "PyTorch GELU: https://docs.pytorch.org/docs/stable/generated/torch.nn.GELU.html"],
            ["Stacking", "Kết hợp xác suất từ deep model và model phụ.", "Giảm phụ thuộc vào một nhánh đơn.", "Ahmed weighted ensemble: https://arxiv.org/abs/2112.05666"],
        ],
        widths=[1.3, 2.2, 2.1, 1.4],
        font_size=7.2,
    )

    add_heading(doc, "3.1.7. Output block", 2)
    add_p(
        doc,
        "Output cuối cùng là p_final[6], tức xác suất cho 6 cảm xúc chung: neutral, happy, sad, angry, fear và disgust. Label dự đoán là lớp có xác suất cao nhất; confidence là xác suất của lớp đó. Khi phát triển thành presentation feedback system, cùng pipeline có thể chạy theo từng đoạn audio để tạo emotion timeline, phát hiện đoạn nói quá căng, quá đơn điệu hoặc thiếu năng lượng."
    )


def build_doc():
    metrics = read_csv(REPORT_DIR / "06D_emotion2vec_coattention_metrics.csv")
    per_dataset = read_csv(REPORT_DIR / "06D_emotion2vec_coattention_per_dataset.csv")
    refs = read_csv(REPORT_DIR / "06D_reference_model_comparison.csv")
    single_metrics = read_csv(SINGLE_REPORT_DIR / "06D_emotion2vec_coattention_metrics.csv")
    single_split = read_csv(SINGLE_REPORT_DIR / "06D_single_dataset_split_summary.csv")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("06D Emotion2Vec Co-Attention SER - Full Detailed Feature & Block Analysis")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = INK
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Bản viết lại chi tiết cho Midterm Param: 1.5, 3.1, 3.2, 3.3")
    sr.font.size = Pt(11)
    sr.font.color.rgb = MUTED

    add_p(
        doc,
        "Tài liệu này viết lại phần proposed solution theo đúng kiến trúc 06D hiện tại. Nội dung bám vào script thuyết trình/ghi chú chi tiết của nhóm, nhưng đã chỉnh lại cho đúng notebook 06D: Branch A dùng TemporalBranch với Conv1D + BiLSTM + AttentionPooling; Branch B dùng SpectralBranch với ResidualSEBlock/SE2D; Branch C dùng frozen emotion2vec + MLP adapter; Branch D dùng statistical vector + MLPBranch và RBF-SVM; fusion dùng EmotionGuidedCoAttention và stacking."
    )
    add_image(doc, ARCH_IMG, "Hình 1. Kiến trúc 06D Emotion2Vec Co-Attention SER.", width=7.0)

    add_15_31_compact_walkthrough(doc)

    add_heading(doc, "1.5 & 3.1. Proposed solution, research origin and architecture workflow", 1)
    add_p(
        doc,
        "Hai mục 1.5 và 3.1 được trình bày chung vì chúng là một mạch lập luận liên tục: trước hết chọn 06D làm proposed solution, sau đó chỉ ra mô hình này lấy ý tưởng từ paper/reference nào, rồi mới đi vào workflow và kiến trúc chi tiết. Cách viết này giúp người đọc hiểu rằng kiến trúc không được ghép ngẫu nhiên, mà mỗi branch và mỗi phương pháp fusion đều có lý do học thuật."
    )

    add_heading(doc, "1.5. Select the proposed solution for the project and justify the selection", 2)
    add_p(
        doc,
        "Đề tài của nhóm tập trung vào Speech Emotion Recognition, tức là nhận diện cảm xúc từ giọng nói. Về dài hạn, hệ thống không chỉ dừng ở phân loại cảm xúc, mà hướng tới presentation feedback system: giúp người nói biết bài thuyết trình có đang quá căng thẳng, thiếu năng lượng, đơn điệu hoặc thay đổi cảm xúc bất thường hay không. Tuy nhiên, ở giai đoạn hiện tại, bài toán lõi cần giải quyết là phân loại cảm xúc từ audio."
    )
    add_p(
        doc,
        "Giải pháp được chọn là 06D Emotion2Vec Co-Attention SER. Đây là mô hình multi-representation vì một audio được nhìn qua bốn dạng biểu diễn khác nhau: temporal acoustic sequence, spectrogram image, pretrained emotion embedding và statistical vector. Cách này phù hợp hơn baseline đơn nhánh vì cảm xúc trong giọng nói không nằm trọn trong một feature duy nhất."
    )
    add_table(
        doc,
        ["Lựa chọn", "Vai trò", "Lý do chọn", "Giá trị cho project"],
        [
            ["Temporal acoustic branch", "Học chuỗi MFCC/delta/RMS/ZCR/spectral theo thời gian.", "Cảm xúc là dynamics, không chỉ một frame tĩnh.", "Bắt thay đổi tone, energy, timbre trong câu nói."],
            ["Spectrogram branch", "Học ảnh log-Mel 3 kênh bằng Residual CNN + SE2D.", "Log-Mel giữ time-frequency structure rõ.", "Bắt vùng phổ/formant/texture theo thời gian."],
            ["Emotion2vec branch", "Dùng frozen pretrained emotion representation.", "Tận dụng model học trước nhưng không tốn GPU fine-tune lớn.", "Tăng khả năng biểu diễn cảm xúc trừu tượng."],
            ["Statistical branch", "Tóm tắt feature thành vector cố định cho MLP/SVM.", "Feature engineering cổ điển vẫn hữu ích trong SER.", "Bổ sung global descriptors và probability cho stacking."],
            ["Co-attention + stacking", "Cho embedding tương tác và kết hợp xác suất.", "Fusion tốt hơn nối vector thô hoặc vote cứng.", "Tạo p_final[6] và confidence cho demo feedback."],
        ],
        widths=[1.45, 2.05, 2.0, 1.9],
        font_size=8,
    )
    add_p(
        doc,
        "Giải pháp này cũng phù hợp điều kiện tài nguyên. Nhóm không fine-tune toàn bộ emotion2vec mà dùng frozen embedding, chỉ train adapter/classifier nhỏ. Như vậy mô hình vẫn có điểm học thuật nhờ pretrained representation và co-attention, nhưng vẫn có thể chạy trên Kaggle/Colab."
    )
    add_reference_trace_section(doc)
    add_15_method_selection_deep_justification(doc)

    add_heading(doc, "3.1. Present the architecture or workflow of the proposed solution", 2)
    add_p(
        doc,
        "Workflow của 06D là: audio 16 kHz -> preprocess -> split protocol -> extract acoustic/pretrained/statistical features -> four branches -> co-attention fusion -> classifier -> stacking -> output. Split phải được thực hiện trước augmentation/model selection để tránh data leakage. Validation dùng để chọn best epoch và fit stacking, test chỉ dùng báo cáo cuối."
    )
    add_code(
        doc,
        "Input audio 16 kHz\n"
        "  -> Branch A: X_temporal [B, 132, T] -> TemporalBranch -> z_temporal [B, 192]\n"
        "  -> Branch B: X_spectral [B, 3, 96, T] -> SpectralBranch -> z_spectral [B, 192]\n"
        "  -> Branch C: X_e2v [B, D_e2v] -> MLPBranch -> z_e2v [B, 192]\n"
        "  -> Branch D: X_stats [B, 796] -> MLPBranch -> z_stats [B, 128]\n"
        "  -> CoAttention(z_e2v queries z_temporal/z_spectral) -> z_context [B, 192]\n"
        "  -> concat([z_t, z_s, z_e, z_context, z_stats]) [B, 896]\n"
        "  -> Fusion MLP -> classifier -> p_deep[6]\n"
        "  -> stacking with p_stats / p_e2v -> p_final[6]"
    )

    add_feature_section(doc)
    add_feature_extraction_pipeline_section(doc)
    add_31_block_code_deep_explanation(doc)

    add_heading(doc, "3.1.4. Branch A - TemporalBranch: 1D-CNN + BiLSTM + AttentionPooling", 3)
    add_p(
        doc,
        "Branch A xử lý chuỗi đặc trưng theo thời gian. Trong 06D, temporal input có 132 channels: 40 MFCC, 40 delta MFCC, 40 delta-delta MFCC, RMS, ZCR, centroid, bandwidth, rolloff và 7 spectral contrast. Mỗi frame vì vậy không chỉ có màu giọng mà còn có thông tin dynamics, năng lượng và hình dạng phổ."
    )
    add_code(
        doc,
        "mfcc = librosa.feature.mfcc(..., n_mfcc=40)\n"
        "delta = librosa.feature.delta(mfcc)\n"
        "delta2 = librosa.feature.delta(mfcc, order=2)\n"
        "rms = librosa.feature.rms(...)\n"
        "zcr = librosa.feature.zero_crossing_rate(...)\n"
        "centroid = librosa.feature.spectral_centroid(...)\n"
        "bandwidth = librosa.feature.spectral_bandwidth(...)\n"
        "rolloff = librosa.feature.spectral_rolloff(...)\n"
        "contrast = librosa.feature.spectral_contrast(...)  # 7 bands\n"
        "temporal = np.vstack([mfcc, delta, delta2, rms, zcr, centroid, bandwidth, rolloff, contrast])\n"
        "# temporal shape: [132, T]"
    )
    add_table(
        doc,
        ["Block trong 06D", "Cấu trúc/code-level", "Input -> Output", "Hoạt động và lý do dùng"],
        [
            ["Conv1D block 1", "Conv1d(132 -> 128, kernel=7, padding=3) + BatchNorm1d + GELU + Dropout", "[B,132,T] -> [B,128,T]", "Kernel 1D trượt theo thời gian để học pattern cục bộ: energy tăng nhanh, MFCC đổi mạnh, spectral cue xuất hiện trong vài frame."],
            ["Conv1D block 2", "Conv1d(128 -> 160, kernel=5, padding=2) + BatchNorm1d + GELU + MaxPool1d(2)", "[B,128,T] -> [B,160,T/2]", "Tăng số channel feature map và giảm chiều thời gian để giữ pattern quan trọng, giảm nhiễu và giảm chi phí tính toán."],
            ["BiLSTM", "LSTM(input_size=160, hidden_size=96, bidirectional=True)", "[B,T/2,160] -> [B,T/2,192]", "Forward context đọc quá khứ -> hiện tại; backward context đọc tương lai -> hiện tại. Cảm xúc thường cần ngữ cảnh hai chiều."],
            ["AttentionPooling", "Linear(192 -> 96) + Tanh + Linear(96 -> 1) + softmax over time", "[B,T/2,192] -> [B,192]", "Không phải frame nào cũng giàu cảm xúc. Attention học trọng số frame và pool có trọng số."],
            ["Projection", "LayerNorm(192) + Linear(192 -> 192) + GELU + Dropout", "[B,192] -> z_temporal [B,192]", "Chuẩn hóa và đưa temporal representation về cùng kích thước với các embedding khác để fusion."],
        ],
        widths=[1.25, 2.05, 1.25, 2.65],
        font_size=7.4,
    )
    add_p(
        doc,
        "Có thể hiểu output của Conv1D là nhiều channel feature map. Một channel có thể nhạy với đoạn tăng năng lượng, channel khác nhạy với vùng phổ sáng, channel khác nhạy với MFCC thay đổi nhanh. BiLSTM sau đó đặt các pattern này vào ngữ cảnh trước/sau, còn attention chọn frame quan trọng nhất cho quyết định cảm xúc."
    )

    add_heading(doc, "3.1.5. Branch B - SpectralBranch: ResidualSEBlock + SE2D", 3)
    add_p(
        doc,
        "Branch B khai thác log-Mel spectrogram 3 kênh: log-Mel, delta log-Mel và delta-delta log-Mel. Đây là biểu diễn dạng ảnh time-frequency. So với MFCC, log-Mel giữ rõ hơn vị trí năng lượng theo thời gian và dải tần Mel, nên phù hợp với 2D-CNN."
    )
    add_code(
        doc,
        "mel = librosa.feature.melspectrogram(..., n_mels=96)\n"
        "logmel = librosa.power_to_db(mel, ref=np.max)\n"
        "d_logmel = librosa.feature.delta(logmel)\n"
        "d2_logmel = librosa.feature.delta(logmel, order=2)\n"
        "spectral = np.stack([logmel, d_logmel, d2_logmel], axis=0)\n"
        "# spectral shape: [3, 96, T]"
    )
    add_table(
        doc,
        ["Block", "Cấu trúc/code-level", "Input -> Output", "Ý nghĩa"],
        [
            ["ResidualSEBlock 1", "Conv2d(3 -> 32) + BN + GELU + Conv2d + BN + SE2D + shortcut", "[B,3,96,T] -> [B,32,96,T]", "Học pattern time-frequency cục bộ từ 3 kênh log-Mel động."],
            ["MaxPool2d", "MaxPool2d(2)", "Giảm time/frequency resolution", "Giữ pattern lớn hơn, giảm chi phí."],
            ["ResidualSEBlock 2", "32 -> 64 channels", "maps -> richer maps", "Học pattern phổ phức tạp hơn."],
            ["ResidualSEBlock 3", "64 -> 128 channels", "maps -> [B,128,...]", "Tạo feature map sâu hơn cho cảm xúc."],
            ["SE2D", "AdaptiveAvgPool2d(1) -> Linear -> GELU -> Linear -> Sigmoid", "feature maps -> channel-reweighted maps", "Tính trọng số cho từng channel, nhấn mạnh map chứa cue cảm xúc."],
            ["Projection", "AdaptiveAvgPool2d((1,1)) + Flatten + LayerNorm(128) + Linear(128 -> 192)", "maps -> z_spectral [B,192]", "Tạo embedding phổ cố định cho fusion."],
        ],
        widths=[1.3, 2.25, 1.35, 2.35],
        font_size=7.4,
    )
    add_p(
        doc,
        "Trong text cũ có nhắc AST/pretrained spectrogram. Tuy nhiên với 06D hiện tại, Branch B không dùng AST; nó dùng CNN-SE tự train trên log-Mel 3 kênh. Pretrained branch của 06D là emotion2vec ở Branch C."
    )

    add_heading(doc, "3.1.6. Branch C - Emotion2Vec branch: frozen pretrained embedding + adapter MLP", 3)
    add_p(
        doc,
        "Branch C dùng waveform/raw speech để lấy emotion2vec embedding. emotion2vec là pretrained speech emotion representation. Trong 06D, phần pretrained được giữ frozen; mô hình chỉ train MLPBranch phía sau để ánh xạ embedding này vào không gian 192-D phù hợp với fusion."
    )
    add_code(
        doc,
        "waveform 16 kHz -> emotion2vec_base (frozen)\n"
        "hidden states / embedding -> X_e2v [B, D_e2v]\n"
        "MLPBranch(e2v_dim, out_dim=192, hidden=384):\n"
        "    Linear(D_e2v -> 384) -> LayerNorm -> GELU -> Dropout\n"
        "    Linear(384 -> 192) -> GELU -> Dropout\n"
        "output: z_e2v [B, 192]"
    )
    add_table(
        doc,
        ["Block", "Input -> Output", "Hoạt động", "Lý do dùng"],
        [
            ["Frozen emotion2vec", "waveform -> X_e2v", "Trích xuất embedding cảm xúc từ model học trước.", "Tận dụng tri thức pretrained mà không cần fine-tune nặng."],
            ["Adapter MLP", "X_e2v -> z_e2v [B,192]", "Nén/biến đổi embedding sang không gian chung.", "Cho phép 06D train nhẹ trên dataset hiện tại."],
            ["Role in co-attention", "z_e2v -> query", "z_e2v hỏi temporal/spectral xem cue nào quan trọng.", "Embedding pretrained đóng vai trò emotion-guided signal."],
        ],
        widths=[1.35, 1.4, 2.35, 2.2],
        font_size=8,
    )
    add_p(
        doc,
        "Lý do không fine-tune toàn bộ emotion2vec là vì dataset SER hiện tại không lớn và project hướng tới demo thực tế. Fine-tune quá mạnh trên bốn dataset nhỏ có thể làm model học dataset/speaker style thay vì emotion signal. Frozen embedding giúp mô hình ổn định hơn và tiết kiệm GPU."
    )

    add_heading(doc, "3.1.7. Branch D - Statistical branch: X_stats, MLPBranch và RBF-SVM", 3)
    add_p(
        doc,
        "Branch D là nhánh thống kê. Trong 06D, stats vector được tạo từ 132 temporal channels. Với mỗi channel, notebook lấy mean, std, min, max, median và IQR, sau đó cộng thêm 4 thống kê waveform-level: energy, max absolute amplitude, std waveform và mean absolute amplitude. Vì vậy kích thước stats thực tế xấp xỉ 132 x 6 + 4 = 796 chiều."
    )
    add_code(
        doc,
        "base = temporal  # [132, T]\n"
        "q25, q75 = percentile(base, 25/75, axis=time)\n"
        "stats = concat([mean, std, min, max, median, iqr, energy, peak, waveform_std, mean_abs])\n"
        "# X_stats shape: [B, 796]\n\n"
        "Stats MLP: MLPBranch(stats_dim, out_dim=128, hidden=256) -> z_stats [B,128]\n"
        "RBF-SVM: StandardScaler + PCA/VarianceThreshold if configured + SVC(kernel='rbf', probability=True)"
    )
    add_table(
        doc,
        ["Khối", "Input -> Output", "Hoạt động", "Vì sao dùng"],
        [
            ["Statistical extraction", "[132,T] -> [796]", "Tóm tắt mỗi channel theo mean/std/min/max/median/IQR.", "SVM/MLP cần vector cố định; thống kê giảm phụ thuộc độ dài audio."],
            ["StandardScaler", "X_stats -> scaled X_stats", "Chuẩn hóa scale từng feature.", "RBF-SVM dựa trên khoảng cách nên rất nhạy với scale."],
            ["Stats MLP", "X_stats -> z_stats [B,128]", "Học nonlinear projection từ stats.", "Cho stats branch tham gia deep fusion."],
            ["RBF-SVM", "scaled stats -> p_stats [B,6]", "Kernel RBF đo tương đồng phi tuyến.", "Bổ sung model cổ điển mạnh cho stacking."],
            ["Probability output", "decision -> probability", "SVC(probability=True) tạo xác suất từng lớp.", "Stacking cần xác suất chứ không chỉ nhãn cứng."],
        ],
        widths=[1.35, 1.45, 2.35, 2.15],
        font_size=8,
    )

    add_heading(doc, "3.1.8. EmotionGuidedCoAttention và classifier", 3)
    add_p(
        doc,
        "Fusion của 06D không chỉ concatenate mọi vector ngay từ đầu. Mô hình dùng EmotionGuidedCoAttention: z_e2v đóng vai trò query, còn z_temporal và z_spectral đóng vai trò key/value. Ý nghĩa trực quan: embedding emotion2vec hỏi hai nhánh acoustic xem thông tin nào phù hợp nhất với cảm xúc đang xét."
    )
    add_code(
        doc,
        "q = z_e2v.unsqueeze(1)                         # [B,1,192]\n"
        "kv = torch.stack([z_temporal, z_spectral], 1)  # [B,2,192]\n"
        "context, weights = MultiheadAttention(q, kv, kv)\n"
        "z_context = LayerNorm(z_e2v + context.squeeze(1))\n\n"
        "z = concat([z_t, z_s, z_e, z_context, z_stats])\n"
        "# 192 + 192 + 192 + 192 + 128 = 896\n"
        "fusion MLP: Linear(896 -> 384) -> LayerNorm -> GELU -> Dropout -> Linear(384 -> 192)\n"
        "classifier: Linear(192 -> 6) -> logits -> softmax -> p_deep"
    )
    add_table(
        doc,
        ["Thành phần", "Input -> Output", "Ý nghĩa"],
        [
            ["Query", "z_e2v [B,192] -> [B,1,192]", "Câu hỏi từ pretrained emotion representation."],
            ["Key/Value", "z_temporal + z_spectral -> [B,2,192]", "Hai nguồn acoustic để attention chọn thông tin."],
            ["Multi-head attention", "q,kv -> context + weights", "Nhiều head học nhiều kiểu quan hệ emotion-acoustic."],
            ["Fusion MLP", "[B,896] -> [B,192]", "Kết hợp 5 vector: temporal, spectral, e2v, context, stats."],
            ["Classifier", "[B,192] -> [B,6]", "Tạo logits cho 6 emotion."],
        ],
        widths=[1.4, 1.7, 3.8],
        font_size=8,
    )

    add_heading(doc, "3.2. Implement the selected speech processing model, algorithm, or method", 1)
    add_p(
        doc,
        "Notebook 06D triển khai đầy đủ pipeline trên ser_processed. Dữ liệu đầu vào gồm metadata.csv và audio_16k. Metadata chứa sample_id, dataset, speaker_id, emotion, gender, duration và split gốc. Notebook tạo lại label_id, dataset_id và speaker_global để phục vụ split/evaluation."
    )
    add_implementation_architecture_walkthrough_v2(doc)
    add_implementation_shape_section(doc)
    add_augmentation_detail_section(doc)
    add_reference_for_extraction_augmentation_section(doc)
    add_table(
        doc,
        ["Bước", "Trong notebook 06D", "Mục đích"],
        [
            ["Load metadata", "Giữ 6 emotion chung: neutral, happy, sad, angry, fear, disgust.", "Đồng nhất label space giữa RAVDESS, CREMA-D, TESS, SAVEE."],
            ["Preprocess audio", "load_audio_fixed: load 16 kHz, pad/trim theo TARGET_DURATION=3s.", "Tạo input cùng độ dài để batch training ổn định."],
            ["Feature extraction", "extract_acoustic_features: tạo temporal, spectral, stats.", "Cung cấp input cho Branch A/B/D."],
            ["Emotion2vec extraction", "Emotion2VecExtractor.extract_one.", "Cung cấp X_e2v cho Branch C."],
            ["Feature cache", "build_or_load_feature_cache.", "Tránh trích xuất lại feature khi chạy lại notebook."],
            ["Data split", "combined_random, combined_strict_no_tess.", "Đánh giá cả split dễ và generalization khó hơn."],
            ["Training", "train_deep_model + run_epoch.", "Train deep co-attention model, chọn best epoch bằng val macro-F1."],
            ["Classical models", "stats_rbf_svm, emotion2vec_logreg, emotion2vec_mlp.", "Tạo model phụ để stacking."],
            ["Stacking", "validation predictions -> meta/weighted ensemble.", "Kết hợp p_deep, p_stats và p_e2v mà không dùng test để chọn."],
            ["Reports", "metrics.csv, per_dataset.csv, figures, predictions.", "Báo cáo accuracy, macro-F1, UAR, confusion matrix."],
        ],
        widths=[1.3, 3.0, 2.8],
        font_size=8,
    )
    add_p(
        doc,
        "Điểm kiểm soát quan trọng là validation/test không được dùng để train scaler, chọn stacking weight hoặc augment dữ liệu. Điều này giúp kết quả test phản ánh khả năng tổng quát thật hơn. 06D dùng macro-F1 làm tiêu chí chọn best epoch vì macro-F1 đối xử tương đối công bằng hơn giữa các lớp cảm xúc."
    )

    add_heading(doc, "3.3. Provide preliminary experimental results or initial training results to demonstrate the feasibility of the proposed solution", 1)
    add_p(
        doc,
        "Kết quả dưới đây lấy từ output thật của notebook 06D và notebook 06D single-dataset. Báo cáo tách ba nhóm evaluation: combined_random, combined_strict_no_tess và single-dataset. Cách tách này quan trọng vì cùng một mô hình có thể đạt điểm cao khi split random, nhưng giảm mạnh khi gặp speaker/domain khác hoặc khi test trên dataset nhiều speaker như CREMA-D."
    )

    add_heading(doc, "3.3.1. Kết quả combined random và combined strict", 2)
    metric_rows = []
    for r in metrics:
        if r.get("split") == "test":
            metric_rows.append([r["protocol"], r["model"], r["n_samples"], pct(r["accuracy"]), pct(r["macro_f1"]), pct(r["uar"]), r.get("best_epoch") or "-"])
    add_table(doc, ["Protocol", "Model", "n", "Accuracy", "Macro-F1", "UAR", "Best epoch"], metric_rows, widths=[1.25, 1.65, 0.55, 0.85, 0.85, 0.85, 0.8], font_size=7.2)

    add_p(
        doc,
        "Kết quả chính: stacking_full đạt 80.39% accuracy và 80.51% macro-F1 trên combined_random. Khi chuyển sang combined_strict_no_tess, stacking_full giảm còn 69.59% accuracy và 69.60% macro-F1. Khoảng giảm này là bằng chứng rằng mô hình học được emotion signal ở mức khả thi, nhưng vẫn còn phụ thuộc đáng kể vào speaker/domain style. Nói cách khác, 06D đã tốt hơn baseline đơn giản nhưng chưa đạt mức domain-robust."
    )
    add_image(doc, LEADERBOARD_IMG, "Hình 7. Macro-F1 leaderboard của các nhánh/mô hình trong 06D.", width=6.5)
    add_image(doc, CURVES_IMG, "Hình 8. Training curves của 06D. Train F1 tăng cao hơn validation F1, cho thấy nguy cơ overfit.", width=6.5)

    add_p(
        doc,
        "Hình training curves được đưa vào để chứng minh feasibility và đồng thời chỉ ra giới hạn hiện tại. Nếu train macro-F1 tăng nhanh nhưng validation macro-F1 đứng lại, mô hình đang học tốt trên train nhưng không cải thiện tương ứng trên validation. Trong SER, nguyên nhân thường là model học speaker identity, dataset recording condition hoặc lexical/acoustic style thay vì chỉ học emotion."
    )

    add_heading(doc, "3.3.2. Per-dataset trong combined protocols", 2)
    per_rows = []
    for r in per_dataset:
        if r.get("model") == "stacking_full":
            per_rows.append([r["protocol"], r["dataset"], r["n_samples"], pct(r["accuracy"]), pct(r["macro_f1"]), pct(r["uar"])])
    add_table(doc, ["Protocol", "Dataset", "n", "Accuracy", "Macro-F1", "UAR"], per_rows, widths=[1.45, 1.0, 0.55, 1.0, 1.0, 1.0], font_size=8)
    add_p(
        doc,
        "Per-dataset result cho thấy TESS rất dễ trong combined_random, trong khi CREMA-D khó hơn do nhiều speaker và recording variation hơn. Nếu một dataset có giọng, cách đọc hoặc môi trường ghi âm đồng nhất, model dễ học shortcut domain. CREMA-D đa dạng hơn nên phản ánh rõ hơn khả năng tổng quát. Vì vậy, khi báo cáo không nên chỉ lấy TESS hoặc accuracy trung bình để kết luận mô hình đã mạnh."
    )

    add_heading(doc, "3.3.3. Kết quả single-dataset experiments", 2)
    add_p(
        doc,
        "Single-dataset experiments train/test riêng trên từng dataset theo kiểu 80/20 và lấy validation từ train. Protocol này gần với nhiều paper single-dataset hơn strict speaker-aware, nên dùng để so sánh tương đối với các paper báo cáo RAVDESS, CREMA-D, SAVEE, TESS riêng lẻ. Tuy nhiên nó vẫn không chứng minh generalization sang speaker/domain mới mạnh như strict protocol."
    )
    split_rows = []
    for r in single_split:
        if r.get("split") in {"train", "validation", "test"}:
            split_rows.append([r["protocol"], r["split"], r["n_samples"], r["speaker_count"], r["split_mode"]])
    add_table(doc, ["Protocol", "Split", "n", "Speakers", "Split mode"], split_rows, widths=[1.4, 0.9, 0.7, 0.8, 2.25], font_size=7.4)

    single_rows = []
    for r in single_metrics:
        if r.get("split") == "test" and r.get("model") == "stacking_full":
            single_rows.append([r["protocol"], r["n_samples"], pct(r["accuracy"]), pct(r["macro_f1"]), pct(r["uar"])])
    add_table(doc, ["Single-dataset protocol", "Test n", "Accuracy", "Macro-F1", "UAR"], single_rows, widths=[1.7, 0.8, 1.0, 1.0, 1.0], font_size=8)
    add_p(
        doc,
        "Kết quả single-dataset cho thấy 06D mạnh hơn rõ trên RAVDESS, SAVEE và đặc biệt TESS, nhưng vẫn thấp trên CREMA-D. Cụ thể, stacking_full đạt khoảng 71.19% trên CREMA-D, 86.32% trên RAVDESS, 83.33% trên SAVEE và 100% trên TESS. Điều này khớp với nhận xét rằng khó khăn chính không phải mô hình hoàn toàn không học được cảm xúc, mà là CREMA-D và combined strict chứa nhiều biến thiên speaker/domain hơn."
    )

    add_heading(doc, "3.3.4. Confusion matrix: vì sao cần và cách đọc", 2)
    add_p(
        doc,
        "Confusion matrix được đưa vào vì accuracy chỉ nói tổng số mẫu đúng, còn không cho biết mô hình nhầm cảm xúc nào với cảm xúc nào. Trong matrix, hàng là nhãn thật, cột là nhãn dự đoán. Các ô trên đường chéo là dự đoán đúng; ô ngoài đường chéo là lỗi nhầm lớp. Nếu sad thường bị nhầm neutral, hoặc fear bị nhầm happy/disgust, ta biết mô hình chưa tách tốt cue acoustic giữa các nhóm cảm xúc đó."
    )
    add_image(doc, CM_RANDOM_IMG, "Hình 9. Confusion matrix của stacking_full trên combined_random.", width=5.8)
    add_image(doc, CM_STRICT_IMG, "Hình 10. Confusion matrix của stacking_full trên combined_strict_no_tess.", width=5.8)
    add_p(
        doc,
        "Hai confusion matrix combined cho thấy random split dễ hơn strict split. Khi chuyển sang strict, các lỗi ngoài đường chéo tăng vì speaker/domain mới làm phân bố feature lệch đi. Đây là lý do nhóm báo cáo macro-F1/UAR song song với accuracy: macro-F1 và UAR giúp nhìn công bằng hơn khi một số emotion hoặc dataset có số lượng mẫu khác nhau."
    )
    add_image(doc, CM_SINGLE_CREMA_IMG, "Hình 11. Confusion matrix single CREMA-D. CREMA-D là single-dataset khó nhất trong kết quả 06D.", width=5.8)
    add_image(doc, CM_SINGLE_RAVDESS_IMG, "Hình 12. Confusion matrix single RAVDESS.", width=5.8)
    add_image(doc, CM_SINGLE_SAVEE_IMG, "Hình 13. Confusion matrix single SAVEE.", width=5.8)
    add_image(doc, CM_SINGLE_TESS_IMG, "Hình 14. Confusion matrix single TESS. Kết quả rất cao nhưng cần cẩn thận vì TESS chỉ có 2 speaker nữ và recording style khá đồng nhất.", width=5.8)

    add_heading(doc, "3.3.5. So sánh references và phân tích vì sao 06D chưa đạt trên 90%", 2)
    ref_rows = [[r.get("model", ""), r.get("protocol", ""), r.get("reported_accuracy_text", ""), r.get("main_idea", "")] for r in refs]
    add_table(doc, ["Reference", "Protocol / split", "Reported result", "Main idea"], ref_rows, widths=[1.6, 1.7, 1.4, 2.3], font_size=7.2)
    add_p(
        doc,
        "Các con số trên 90% trong reference không nên so trực tiếp một dòng với combined_strict_no_tess của nhóm. Một số paper dùng single-dataset hoặc random split, thường dễ hơn vì train/test có thể cùng speaker style hoặc cùng điều kiện ghi âm. Một số paper cũng không công bố đầy đủ code, seed, danh sách file train/test hoặc speaker split, nên khó tái lập chính xác. Vì vậy báo cáo nên ghi rõ: 06D so sánh trực tiếp nhất với các setup có code/split rõ; còn paper không rõ split được dùng làm tham khảo ý tưởng kiến trúc."
    )
    add_table(
        doc,
        ["Nhóm/reference", "Có code/split rõ không?", "Họ làm gì tốt hơn?", "Tại sao có thể cao hơn 06D?"],
        [
            [
                "Emonity repo - https://github.com/sv6095/Emonity",
                "Có repo tham khảo, nhưng khi đưa về dataset/split của nhóm vẫn cần kiểm soát lại preprocessing và split.",
                "Pipeline của họ thiên về practical ensemble: nhiều feature acoustic, augmentation, nhiều backbone như 1D-CNN/2D-CNN/CNN-BiLSTM và weighted ensemble.",
                "Ensemble nhiều model độc lập thường giảm lỗi của từng nhánh. Nếu split random/single dễ hơn strict, điểm có thể cao hơn rõ."
            ],
            [
                "Ahmed et al. weighted ensemble - https://arxiv.org/abs/2112.05666",
                "Có mô tả rõ hướng 1D-CNN, CNN-LSTM, CNN-GRU và weighted ensemble; split không phải strict speaker-aware như protocol khó của nhóm.",
                "Họ dùng augmentation và ensemble các recurrent CNN variants. LSTM/GRU học temporal context, còn weighted ensemble chọn trọng số theo validation.",
                "Các model chuyên biệt cho từng dataset/split có thể fit distribution tốt hơn; augmentation và ensemble làm random/single-dataset accuracy tăng."
            ],
            [
                "Chowdhury/DCRF-BiLSTM style feature engineering",
                "Không phải tất cả chi tiết code/split đều có thể tái lập từ tài liệu local; dùng làm reference về feature engineering và BiLSTM/CRF.",
                "Họ nhấn mạnh handcrafted feature engineering sâu hơn, có thể gồm nhiều thống kê/selection/chuẩn hóa trước khi đưa vào recurrent/CRF.",
                "06D hiện có stats vector nhưng RBF-SVM còn yếu; nếu feature thủ công chưa đủ tinh, nhánh classical không kéo được overall lên như paper feature-engineering mạnh."
            ],
            [
                "Ullah et al. 1D-CNN feature fusion",
                "Báo cáo combined 4 dataset nhưng split chi tiết không đủ để tái lập hoàn toàn từ local materials.",
                "Dùng ZCR, energy, entropy of energy, RMS, MFCC rồi đưa vào 1D-CNN.",
                "Nếu split là random và preprocessing/feature selection được tối ưu, model có thể học dataset style tốt hơn, dẫn tới accuracy cao."
            ],
        ],
        widths=[1.45, 1.35, 2.15, 2.15],
        font_size=6.8,
    )
    add_p(
        doc,
        "Vì sao 06D còn yếu: thứ nhất, stats_rbf_svm là nhánh thấp nhất trong cả combined và single CREMA-D, nên stacking nhận một nguồn xác suất chưa mạnh. Thứ hai, emotion2vec đang frozen; frozen giúp tiết kiệm tài nguyên nhưng không tối ưu sâu cho 6-label setup của nhóm. Thứ ba, 3 giây pad/trim có thể làm mất cue ở audio dài hoặc thêm silence ở audio ngắn. Thứ tư, 4 dataset có style khác nhau: TESS chỉ 2 speaker nữ và rất sạch, SAVEE chỉ 4 speaker nam, RAVDESS diễn xuất có cấu trúc, CREMA-D nhiều speaker hơn và khó hơn. Khi gộp lại, mô hình phải học emotion signal đồng thời bỏ qua speaker/domain signal, đây là bài toán khó hơn nhiều so với single-dataset random."
    )
    add_p(
        doc,
        "Tuy vậy, kết quả không phải thất bại. 06D đạt trên 80% random, gần 70% strict và đạt tốt trên RAVDESS/SAVEE/TESS single-dataset. Điều này đủ chứng minh feasibility cho midterm: pipeline chạy được, có mô hình multi-branch, có pretrained representation, có co-attention, có stacking và có đánh giá theo nhiều protocol. Điểm cần phát triển tiếp là domain-aware adapter, feature engineering sâu hơn cho stats branch, segment/dynamic duration, calibration và thử lại ensemble chuyên biệt nhưng vẫn giữ split sạch."
    )
    add_p(
        doc,
        "Kết luận: 06D là proposed solution hợp lý cho midterm vì đã có pipeline hoàn chỉnh từ audio -> feature extraction -> multi-branch model -> co-attention -> stacking -> evaluation. Nó cũng mở đường cho hệ thống feedback giọng nói: p_final[6] và confidence có thể được tính theo từng segment để tạo emotion timeline, phát hiện đoạn nói quá căng, thiếu năng lượng hoặc đơn điệu."
    )

    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except PermissionError:
        try:
            doc.save(ALT_DOCX_PATH)
            return ALT_DOCX_PATH
        except PermissionError:
            try:
                doc.save(ALT_DOCX_PATH_V2)
                return ALT_DOCX_PATH_V2
            except PermissionError:
                doc.save(ALT_DOCX_PATH_V3)
                return ALT_DOCX_PATH_V3


def make_zip(docx):
    files = [
        docx,
        ARCH_IMG,
        EVIDENCE_WAVE_RMS,
        EVIDENCE_MFCC,
        EVIDENCE_ZCR,
        EVIDENCE_SPECTRAL,
        EVIDENCE_LOGMEL,
        BILSTM_IMG,
        DOMAIN_DATASET_IMG,
        CURVES_IMG,
        LEADERBOARD_IMG,
        CM_RANDOM_IMG,
        CM_STRICT_IMG,
        CM_SINGLE_CREMA_IMG,
        CM_SINGLE_RAVDESS_IMG,
        CM_SINGLE_SAVEE_IMG,
        CM_SINGLE_TESS_IMG,
    ]
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in files:
            if Path(p).exists():
                zf.write(p, Path(p).relative_to(ROOT))
    return ZIP_PATH


def main():
    docx = build_doc()
    zipped = make_zip(docx)
    print("DOCX", docx)
    print("ZIP", zipped)


if __name__ == "__main__":
    main()
