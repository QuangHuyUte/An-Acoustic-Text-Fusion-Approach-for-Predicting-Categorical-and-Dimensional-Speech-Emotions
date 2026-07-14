from pathlib import Path
import csv
import zipfile

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "Midterm_Param" / "06D_midterm_assets"
FEATURE_DIR = ROOT / "06_Advanced_Model" / "architecture_diagram" / "assets"
REPORT_DIR = ROOT / "06_Advanced_Model" / "outputs" / "06D_Emotion2Vec_CoAttention_outputs_package" / "reports"
ASSET_DIR.mkdir(parents=True, exist_ok=True)


def font(size, bold=False):
    files = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for f in files:
        if Path(f).exists():
            return ImageFont.truetype(f, size)
    return ImageFont.load_default()


F = {
    "title": font(62, True),
    "sub": font(32),
    "h": font(40, True),
    "h2": font(32, True),
    "body": font(27),
    "small": font(23),
    "tiny": font(20),
    "mini": font(18),
}

C = {
    "bg": "#F5F7FB",
    "ink": "#102A43",
    "muted": "#52606D",
    "line": "#D6E0EA",
    "panel": "#FFFFFF",
    "slate": "#334155",
    "blue": "#2563EB",
    "blue_bg": "#ECF3FF",
    "cyan": "#0891B2",
    "cyan_bg": "#E9FBFF",
    "purple": "#7C3AED",
    "purple_bg": "#F4F0FF",
    "green": "#059669",
    "green_bg": "#EAFBF3",
    "orange": "#D97706",
    "orange_bg": "#FFF5E8",
}


def rr(draw, box, fill, outline, radius=22, width=4):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def wrap(draw, text, fnt, max_w):
    lines = []
    for para in str(text).split("\n"):
        words = para.split()
        if not words:
            lines.append("")
            continue
        cur = ""
        for word in words:
            trial = f"{cur} {word}".strip()
            if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_w or not cur:
                cur = trial
            else:
                lines.append(cur)
                cur = word
        lines.append(cur)
    return lines


def text(draw, xy, s, fnt, fill=None, max_w=None, gap=6):
    fill = fill or C["ink"]
    x, y = xy
    lines = wrap(draw, s, fnt, max_w) if max_w else str(s).split("\n")
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += draw.textbbox((0, 0), line or "Ag", font=fnt)[3] + gap
    return y


def arrow(draw, p1, p2, color=C["slate"], width=5):
    x1, y1 = p1
    x2, y2 = p2
    draw.line((x1, y1, x2, y2), fill=color, width=width)
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 18, y2 - 9), (x2 - 18, y2 + 9)]
    else:
        pts = [(x2, y2), (x2 + 18, y2 - 9), (x2 + 18, y2 + 9)]
    draw.polygon(pts, fill=color)


def v_arrow(draw, p1, p2, color=C["slate"], width=5):
    x1, y1 = p1
    x2, y2 = p2
    draw.line((x1, y1, x2, y2), fill=color, width=width)
    pts = [(x2, y2), (x2 - 9, y2 - 18), (x2 + 9, y2 - 18)]
    draw.polygon(pts, fill=color)


def paste_fit(base, src, box):
    x1, y1, x2, y2 = box
    img = ImageOps.fit(src.convert("RGB"), (x2 - x1, y2 - y1), method=Image.Resampling.LANCZOS)
    mask = Image.new("L", img.size, 0)
    md = ImageDraw.Draw(mask)
    md.rounded_rectangle((0, 0, img.size[0], img.size[1]), radius=18, fill=255)
    base.paste(img, (x1, y1), mask)


def feature_crop(name):
    src = Image.open(FEATURE_DIR / "savee_feature_example.png")
    w, h = src.size
    boxes = {
        "wave": (60, 70, w - 170, 320),
        "spec": (60, 350, w - 170, 660),
        "mfcc": (60, 705, w - 170, 1005),
        "logmel": (60, 1060, w - 170, 1370),
    }
    return src.crop(boxes[name])


def prep_card(draw, x, y, title, body, color, fill, h=150):
    rr(draw, (x, y, x + 425, y + h), fill, color, 22, 4)
    text(draw, (x + 24, y + 18), title, F["h2"], color)
    text(draw, (x + 24, y + 58), body, F["small"], C["ink"], 365, 4)


def layer(draw, x, y, w, title, body, color):
    rr(draw, (x, y, x + w, y + 132), "#FFFFFF", color, 16, 3)
    text(draw, (x + 18, y + 14), title, F["h2"], color, w - 36, 4)
    text(draw, (x + 18, y + 60), body, F["tiny"], C["muted"], w - 36, 4)


def branch(img, draw, y, tag, title, color, fill, feat_name, input_text, cue_text, layers, out_text):
    x0, x1, h = 620, 2765, 350
    rr(draw, (x0, y, x1, y + h), fill, color, 24, 5)
    rr(draw, (x0 + 25, y + 24, x0 + 83, y + 82), color, color, 14, 2)
    text(draw, (x0 + 45, y + 36), tag, F["h2"], "#FFFFFF")
    text(draw, (x0 + 105, y + 28), title, F["h"], color)

    rr(draw, (x0 + 35, y + 105, x0 + 330, y + 278), "#FFFFFF", C["line"], 14, 2)
    paste_fit(img, feature_crop(feat_name), (x0 + 48, y + 118, x0 + 317, y + 265))

    rr(draw, (x0 + 360, y + 105, x0 + 805, y + 278), "#FFFFFF", color, 16, 3)
    text(draw, (x0 + 385, y + 122), "Feature input", F["h2"], color)
    text(draw, (x0 + 385, y + 166), input_text, F["small"], C["ink"], 390, 4)
    rr(draw, (x0 + 360, y + 292, x0 + 805, y + 332), "#FFFFFF", color, 12, 2)
    text(draw, (x0 + 382, y + 302), cue_text, F["mini"], color, 400, 2)

    lx = x0 + 835
    layer_w = 285
    for i, (lt, lb) in enumerate(layers):
        px = lx + i * 330
        layer(draw, px, y + 128, layer_w, lt, lb, color)
        if i < len(layers) - 1:
            arrow(draw, (px + layer_w, y + 194), (px + 314, y + 194), color, 4)

    rr(draw, (x1 - 345, y + 118, x1 - 45, y + 252), "#FFFFFF", color, 15, 3)
    text(draw, (x1 - 318, y + 142), "Branch output", F["small"], color, 245, 4)
    text(draw, (x1 - 318, y + 178), out_text, F["body"], C["ink"], 245, 5)
    return (x1, y + h // 2)


def draw_landscape():
    W, H = 3840, 2160
    img = Image.new("RGB", (W, H), C["bg"])
    d = ImageDraw.Draw(img)

    # Header.
    rr(d, (60, 45, W - 60, 190), "#FFFFFF", C["line"], 26, 2)
    text(d, (98, 70), "06D Emotion2Vec Co-Attention SER - Landscape Architecture", F["title"], C["ink"])
    text(d, (102, 138), "A single multi-representation SER model: audio -> features -> four branches -> co-attention + stacking -> 6 emotion probabilities", F["sub"], C["muted"])

    # Left pipeline panel.
    rr(d, (70, 250, 560, 2040), "#FFFFFF", "#94A3B8", 26, 4)
    text(d, (110, 282), "Input pipeline", F["h"], C["slate"])
    prep_card(d, 105, 345, "1. Raw audio", "single utterance / segment\n16 kHz waveform", C["blue"], C["blue_bg"], 145)
    prep_card(d, 105, 535, "2. Preprocess", "resample -> mono\nnormalize -> crop/pad 3s", C["cyan"], C["cyan_bg"], 145)
    prep_card(d, 105, 725, "3. Split first", "train / validation / test\nno augmentation before split", C["purple"], C["purple_bg"], 145)
    prep_card(d, 105, 915, "4. Feature router", "same audio -> temporal,\nlog-Mel, emotion2vec, stats", C["orange"], C["orange_bg"], 160)
    v_arrow(d, (318, 490), (318, 535), C["slate"], 4)
    v_arrow(d, (318, 680), (318, 725), C["slate"], 4)
    v_arrow(d, (318, 870), (318, 915), C["slate"], 4)

    # Mini real feature panel.
    rr(d, (105, 1165, 525, 1945), "#F8FAFC", C["line"], 22, 3)
    text(d, (132, 1195), "Feature examples", F["h2"], C["slate"])
    for i, (name, label) in enumerate([("wave", "waveform"), ("mfcc", "MFCC"), ("logmel", "log-Mel")]):
        y = 1270 + i * 195
        rr(d, (132, y, 355, y + 145), "#FFFFFF", C["line"], 12, 2)
        paste_fit(img, feature_crop(name), (142, y + 12, 345, y + 133))
        text(d, (375, y + 58), label, F["small"], C["ink"])

    # Main model.
    rr(d, (600, 250, 2790, 2040), "#FFFFFF", "#94A3B8", 26, 4)
    text(d, (640, 280), "Four representation branches inside one model", F["h"], C["ink"])
    text(d, (642, 324), "Each branch captures a different emotion cue. The outputs are embeddings/probabilities, not final separate decisions.", F["body"], C["muted"])

    branch_outputs = []
    branch_outputs.append(branch(
        img, d, 390, "A", "Temporal acoustic branch", C["blue"], C["blue_bg"], "mfcc",
        "40 MFCC + delta + delta2\nRMS/ZCR/spectral sequence\nshape: [B, 132, T]",
        "Captures: timbre + dynamics",
        [("1D-CNN", "Conv1D + BN/ReLU\nlocal prosody filters"), ("BiLSTM/GRU", "forward/backward\nlong-range context"), ("Attention", "learn frame weights\npool emotional frames")],
        "z_temporal\n192-D embedding",
    ))
    branch_outputs.append(branch(
        img, d, 785, "B", "Spectrogram branch", C["cyan"], C["cyan_bg"], "logmel",
        "log-Mel + delta + delta2\n96 mel bins over time\nshape: [B, 3, 96, T]",
        "Captures: time-frequency patterns",
        [("2D-CNN", "time-frequency kernels\nlocal spectral shapes"), ("Residual CNN", "skip connection\nstable deep filters"), ("SE attention", "channel re-weighting\nkeep useful maps")],
        "z_spectral\n192-D embedding",
    ))
    branch_outputs.append(branch(
        img, d, 1180, "C", "Emotion pretrained branch", C["purple"], C["purple_bg"], "wave",
        "raw waveform 16 kHz\nfrozen emotion2vec hidden states\nadapter only is trained",
        "Captures: pretrained emotion cues",
        [("emotion2vec", "self-supervised encoder\nfrozen weights"), ("Mean pool", "utterance-level vector\naverage hidden states"), ("Adapter MLP", "Linear + GELU\nproject to label space")],
        "z_e2v\np_e2v optional",
    ))
    branch_outputs.append(branch(
        img, d, 1575, "D", "Statistical branch", C["green"], C["green_bg"], "spec",
        "clip-level handcrafted vector\nMFCC/chroma/spectral/energy\nmean/std/min/max/percentile",
        "Captures: global stable descriptors",
        [("Scaler", "standardize features\nfit on train only"), ("Stats MLP", "nonlinear stats view\ncompact descriptor"), ("RBF-SVM", "classical margin model\ncalibrated probability")],
        "p_stats\nSVM probability",
    ))

    # Right fusion and output.
    rr(d, (2840, 250, 3770, 1645), "#FFFFFF", "#64748B", 26, 4)
    text(d, (2885, 282), "Fusion + final decision", F["h"], C["slate"])
    text(d, (2888, 326), "The final output is produced after combining all four views.", F["body"], C["muted"], 770)
    rr(d, (2888, 375, 3715, 435), "#F8FAFC", C["line"], 14, 2)
    text(d, (2915, 394), "embedding path -> co-attention -> p_deep | stats probability path -> stacking", F["small"], C["slate"], 760)

    # Embedding fusion lane: A/B/C embeddings go through co-attention.
    rr(d, (2888, 470, 3715, 1045), "#F8FAFC", "#CBD5E1", 18, 3)
    text(d, (2918, 500), "Embedding fusion lane", F["h2"], C["slate"])
    rr(d, (2920, 575, 3205, 900), "#FFFFFF", "#94A3B8", 18, 3)
    text(d, (2950, 605), "A + B + C embeddings", F["h2"], C["slate"])
    for i, (name, color) in enumerate([("z_temporal", C["blue"]), ("z_spectral", C["cyan"]), ("z_e2v", C["purple"])]):
        yy = 670 + i * 62
        rr(d, (2950, yy, 3175, yy + 40), C["panel"], color, 10, 2)
        text(d, (2970, yy + 10), name, F["mini"], color)
    rr(d, (3290, 555, 3690, 765), C["orange_bg"], C["orange"], 20, 4)
    text(d, (3320, 585), "Co-attention", F["h2"], C["orange"])
    text(d, (3320, 628), "emotion2vec attends to\ntemporal + spectral cues", F["small"], C["ink"], 330, 4)
    rr(d, (3290, 845, 3690, 990), C["blue_bg"], C["blue"], 20, 4)
    text(d, (3320, 878), "Deep classifier", F["h2"], C["blue"])
    text(d, (3320, 920), "MLP -> softmax -> p_deep", F["small"], C["ink"])
    arrow(d, (3205, 735), (3290, 660), C["slate"], 5)
    v_arrow(d, (3490, 765), (3490, 845), C["orange"], 5)

    # Probability ensemble lane: Branch D probability bypasses co-attention.
    rr(d, (2888, 1110, 3715, 1640), "#FFFFFF", "#86EFAC", 18, 3)
    text(d, (2918, 1140), "Probability ensemble lane", F["h2"], C["green"])
    rr(d, (2920, 1438, 3205, 1565), "#FFFFFF", C["green"], 18, 3)
    text(d, (2950, 1465), "D probability", F["h2"], C["green"])
    text(d, (2950, 1510), "p_stats from SVM", F["body"], C["ink"], 235)
    rr(d, (3290, 1405, 3690, 1588), C["green_bg"], C["green"], 20, 4)
    text(d, (3320, 1436), "Weighted stacking", F["h2"], C["green"])
    text(d, (3320, 1483), "p_final = combine\np_deep + p_stats", F["small"], C["ink"], 330, 5)
    arrow(d, (3205, 1500), (3290, 1500), C["green"], 4)
    v_arrow(d, (3490, 990), (3490, 1405), C["blue"], 5)

    rr(d, (2840, 1700, 3770, 2040), C["green_bg"], C["green"], 24, 5)
    text(d, (2885, 1738), "Output", F["h"], C["green"])
    # Probability vector area.
    rr(d, (2888, 1810, 3195, 1995), "#FFFFFF", "#86EFAC", 16, 2)
    text(d, (2920, 1837), "p_final[6]", F["h2"], C["green"])
    bar_x, bar_y = 2935, 1877
    for j, val in enumerate([0.22, 0.47, 0.18, 0.82, 0.35, 0.28]):
        bh = int(65 * val)
        x = bar_x + j * 32
        d.rounded_rectangle((x, bar_y + 70 - bh, x + 22, bar_y + 70), radius=6, fill=C["green"])
    # Feedback value area.
    rr(d, (3225, 1810, 3420, 1995), "#FFFFFF", "#86EFAC", 16, 2)
    text(d, (3250, 1837), "returns", F["h2"], C["green"])
    text(d, (3250, 1883), "label\nconfidence\ntrend", F["small"], C["ink"], 142, 5)
    # Label grid area.
    rr(d, (3450, 1810, 3730, 1995), "#FFFFFF", "#86EFAC", 16, 2)
    text(d, (3475, 1837), "6 labels", F["h2"], C["green"])
    labels = ["neutral", "happy", "sad", "angry", "fear", "disgust"]
    for i, lab in enumerate(labels):
        x = 3475 + (i % 2) * 122
        y = 1880 + (i // 2) * 34
        rr(d, (x, y, x + 108, y + 28), C["green_bg"], "#86EFAC", 8, 2)
        text(d, (x + 13, y + 6), lab, F["mini"], C["green"])

    # Connect branch outputs. Embeddings go to co-attention; probabilities bypass to stacking.
    embedding_bus_x = 2868
    a_y, b_y, c_y = branch_outputs[0][1], branch_outputs[1][1], branch_outputs[2][1]
    for (x, y), color in [(branch_outputs[0], C["blue"]), (branch_outputs[1], C["cyan"]), (branch_outputs[2], C["purple"])]:
        d.line((x, y, embedding_bus_x, y), fill=color, width=4)
    d.line((embedding_bus_x, a_y, embedding_bus_x, c_y), fill="#64748B", width=6)
    arrow(d, (embedding_bus_x, 735), (2920, 735), "#64748B", 5)
    # Auxiliary probabilities bypass co-attention and enter stacking.
    x, y = branch_outputs[3]
    # Route the statistical probability around the standalone Output block.
    d.line((x, y, 2816, y, 2816, 1500, 2920, 1500), fill=C["green"], width=4, joint="curve")
    d.polygon([(2920, 1500), (2920 - 16, 1492), (2920 - 16, 1508)], fill=C["green"])
    v_arrow(d, (3490, 1588), (3490, 1700), C["green"], 5)

    out = ASSET_DIR / "06D_architecture_landscape_final.png"
    img.save(out)
    return out


def make_html():
    path = ASSET_DIR / "06D_architecture_landscape_final_source.html"
    path.write_text(
        """<!doctype html>
<html lang="vi"><head><meta charset="utf-8"><title>06D landscape architecture</title>
<style>
body{margin:0;background:#f5f7fb;font-family:Segoe UI,Arial,sans-serif;color:#102a43}
.canvas{width:3840px;height:2160px;box-sizing:border-box;padding:60px}
.card{background:white;border:4px solid #d6e0ea;border-radius:26px;padding:26px}
.branch{display:grid;grid-template-columns:70px 270px 430px 1fr 300px;gap:24px;align-items:center;border-radius:24px;margin:28px 0;padding:24px}
.blue{border:5px solid #2563eb;background:#ecf3ff}.cyan{border:5px solid #0891b2;background:#e9fbff}
.purple{border:5px solid #7c3aed;background:#f4f0ff}.green{border:5px solid #059669;background:#eafbf3}
.layers{display:flex;gap:28px}.layer{background:white;border:3px solid currentColor;border-radius:16px;padding:14px 18px;min-width:210px}
</style></head><body><div class="canvas">
<div class="card"><h1>06D Emotion2Vec Co-Attention SER - Landscape Architecture</h1>
<p>A single multi-representation SER model: audio -> features -> four branches -> co-attention + stacking -> six emotion probabilities.</p></div>
</div></body></html>""",
        encoding="utf-8",
    )
    return path


def read_metrics():
    path = REPORT_DIR / "06D_emotion2vec_coattention_metrics.csv"
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def make_doc(image, html):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    p = doc.add_paragraph()
    r = p.add_run("Kien truc ngang 06D - Emotion2Vec Co-Attention SER")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(16, 42, 67)
    doc.add_paragraph("Ban landscape final: tranh nhieu khoang trong chet, chu lon hon, moi branch the hien input -> layers -> output.")
    doc.add_picture(str(image), width=Inches(7.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Ket qua 06D dung trong bao cao", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    headers = ["Protocol", "Model", "Accuracy", "Macro-F1", "UAR"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in read_metrics():
        if row.get("model") == "stacking_full":
            cells = t.add_row().cells
            vals = [
                row.get("protocol", ""),
                row.get("model", ""),
                f"{float(row.get('accuracy', 0)):.4f}",
                f"{float(row.get('macro_f1', 0)):.4f}",
                f"{float(row.get('uar', 0)):.4f}",
            ]
            for i, v in enumerate(vals):
                cells[i].text = v
    doc.add_paragraph(f"HTML/CSS source: {html}")
    out = ROOT / "Midterm_Param" / "Midterm_Param_06D_Landscape_Architecture_Final.docx"
    doc.save(out)
    return out


def make_zip(files):
    out = ASSET_DIR / "06D_landscape_architecture_final_package.zip"
    if out.exists():
        out.unlink()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for f in files:
            if f.exists():
                z.write(f, f.relative_to(ROOT))
    return out


def main():
    image = draw_landscape()
    html = make_html()
    doc = make_doc(image, html)
    z = make_zip([image, html, doc])
    print("IMAGE", image)
    print("HTML", html)
    print("DOCX", doc)
    print("ZIP", z)


if __name__ == "__main__":
    main()
