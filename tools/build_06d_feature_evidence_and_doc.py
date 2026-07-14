from pathlib import Path
import csv
import math
import wave
import zipfile

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SER_DIR = ROOT / "01&02_Data_and_DataProcessing" / "ser_processed"
META_PATH = SER_DIR / "metadata.csv"
AUDIO_DIR = SER_DIR / "audio_16k"
OUT_DIR = ROOT / "Midterm_Param"
ASSET_DIR = OUT_DIR / "06D_midterm_assets"
EVIDENCE_DIR = ASSET_DIR / "feature_evidence_6_emotions"
REPORT_DIR = ROOT / "06_Advanced_Model" / "outputs" / "06D_Emotion2Vec_CoAttention_outputs_package" / "reports"
EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "Midterm_Param_Sections_1.5_3.1_3.2_3.3_06D_Detailed_With_Feature_Evidence.docx"
ZIP_PATH = ASSET_DIR / "Midterm_Param_06D_Detailed_With_Feature_Evidence_Package.zip"

ARCH_IMG = ASSET_DIR / "06D_architecture_landscape_final.png"
CURVES_IMG = ASSET_DIR / "06D_training_curves.png"
LEADERBOARD_IMG = ASSET_DIR / "06D_macro_f1_leaderboard.png"
CM_RANDOM_IMG = ASSET_DIR / "06D_confusion_combined_random_stacking_full.png"
CM_STRICT_IMG = ASSET_DIR / "06D_confusion_combined_strict_no_tess_stacking_full.png"

COMMON_EMOTIONS = ["neutral", "happy", "sad", "angry", "fear", "disgust"]
SR = 16000
N_FFT = 512
WIN = 400
HOP = 160
N_MELS = 96
N_MFCC = 40

COLORS = {
    "ink": "#102A43",
    "muted": "#52606D",
    "blue": "#2563EB",
    "cyan": "#0891B2",
    "purple": "#7C3AED",
    "green": "#059669",
    "orange": "#D97706",
    "red": "#DC2626",
    "bg": "#F6F8FC",
    "panel": "#FFFFFF",
    "grid": "#E2E8F0",
}


def font(size, bold=False):
    files = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for f in files:
        if Path(f).exists():
            return ImageFont.truetype(f, size)
    return ImageFont.load_default()


F = {
    "title": font(34, True),
    "h": font(25, True),
    "body": font(20),
    "small": font(16),
    "tiny": font(13),
}


def rr(draw, box, fill, outline="#CBD5E1", radius=14, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def read_wav(path, target_len=SR * 3):
    with wave.open(str(path), "rb") as w:
        raw = w.readframes(w.getnframes())
        sw = w.getsampwidth()
        ch = w.getnchannels()
        sr = w.getframerate()
    if sw == 2:
        x = np.frombuffer(raw, dtype=np.int16).astype(np.float32) / 32768.0
    elif sw == 4:
        x = np.frombuffer(raw, dtype=np.int32).astype(np.float32) / 2147483648.0
    else:
        x = np.frombuffer(raw, dtype=np.uint8).astype(np.float32)
        x = (x - 128.0) / 128.0
    if ch > 1:
        x = x.reshape(-1, ch).mean(axis=1)
    if sr != SR:
        # Simple linear resample fallback, normally unused because audio_16k is already 16 kHz.
        old = np.linspace(0, 1, len(x), endpoint=False)
        new_len = int(len(x) * SR / sr)
        new = np.linspace(0, 1, new_len, endpoint=False)
        x = np.interp(new, old, x).astype(np.float32)
    if len(x) < target_len:
        x = np.pad(x, (0, target_len - len(x)))
    else:
        x = x[:target_len]
    peak = np.max(np.abs(x)) + 1e-9
    x = x / max(peak, 0.05)
    return x.astype(np.float32)


def frames(x):
    if len(x) < WIN:
        x = np.pad(x, (0, WIN - len(x)))
    n = 1 + (len(x) - WIN) // HOP
    idx = np.arange(WIN)[None, :] + HOP * np.arange(n)[:, None]
    return x[idx] * np.hamming(WIN)[None, :]


def stft_power(x):
    fr = frames(x)
    if WIN < N_FFT:
        fr = np.pad(fr, ((0, 0), (0, N_FFT - WIN)))
    spec = np.fft.rfft(fr, n=N_FFT, axis=1)
    power = (np.abs(spec) ** 2).T
    mag = np.sqrt(power + 1e-10)
    return power, mag


def hz_to_mel(hz):
    return 2595.0 * np.log10(1.0 + hz / 700.0)


def mel_to_hz(mel):
    return 700.0 * (10 ** (mel / 2595.0) - 1.0)


def mel_filterbank(sr=SR, n_fft=N_FFT, n_mels=N_MELS):
    fmin, fmax = 0.0, sr / 2
    mel_pts = np.linspace(hz_to_mel(fmin), hz_to_mel(fmax), n_mels + 2)
    hz_pts = mel_to_hz(mel_pts)
    bins = np.floor((n_fft + 1) * hz_pts / sr).astype(int)
    fb = np.zeros((n_mels, n_fft // 2 + 1), dtype=np.float32)
    for m in range(1, n_mels + 1):
        left, center, right = bins[m - 1], bins[m], bins[m + 1]
        center = max(center, left + 1)
        right = max(right, center + 1)
        for k in range(left, center):
            if 0 <= k < fb.shape[1]:
                fb[m - 1, k] = (k - left) / (center - left)
        for k in range(center, right):
            if 0 <= k < fb.shape[1]:
                fb[m - 1, k] = (right - k) / (right - center)
    return fb


MEL_FB = mel_filterbank()
FREQS = np.fft.rfftfreq(N_FFT, d=1 / SR)


def dct_type_2(x, n=N_MFCC):
    m = x.shape[0]
    k = np.arange(n)[:, None]
    i = np.arange(m)[None, :]
    basis = np.cos(np.pi / m * (i + 0.5) * k)
    return basis @ x


def delta(mat):
    return np.gradient(mat, axis=1)


def compute_features(x):
    power, mag = stft_power(x)
    mel = MEL_FB @ power
    logmel = 10.0 * np.log10(mel + 1e-8)
    mfcc = dct_type_2(logmel, N_MFCC)
    d_mfcc = delta(mfcc)
    d2_mfcc = delta(d_mfcc)

    fr = frames(x)
    rms = np.sqrt(np.mean(fr ** 2, axis=1) + 1e-10)
    zcr = np.mean(np.abs(np.diff(np.signbit(fr), axis=1)), axis=1)
    denom = np.sum(mag, axis=0) + 1e-10
    centroid = np.sum(FREQS[:, None] * mag, axis=0) / denom
    bandwidth = np.sqrt(np.sum(((FREQS[:, None] - centroid[None, :]) ** 2) * mag, axis=0) / denom)
    cumsum = np.cumsum(mag, axis=0)
    threshold = 0.85 * cumsum[-1, :]
    rolloff_idx = np.argmax(cumsum >= threshold[None, :], axis=0)
    rolloff = FREQS[rolloff_idx]

    db = 20.0 * np.log10(mag + 1e-8)
    bands = np.array_split(np.arange(db.shape[0]), 7)
    contrast = []
    for b in bands:
        vals = db[b, :]
        hi = np.percentile(vals, 90, axis=0)
        lo = np.percentile(vals, 10, axis=0)
        contrast.append(hi - lo)
    contrast = np.vstack(contrast)

    return {
        "waveform": x,
        "rms": rms,
        "zcr": zcr,
        "centroid": centroid,
        "bandwidth": bandwidth,
        "rolloff": rolloff,
        "logmel": logmel,
        "mfcc": mfcc,
        "delta": d_mfcc,
        "delta2": d2_mfcc,
        "contrast": contrast,
    }


def select_samples():
    rows = []
    with META_PATH.open("r", encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            sid = row["sample_id"]
            row["local_audio"] = AUDIO_DIR / f"{sid}.wav"
            if row["emotion"] in COMMON_EMOTIONS and row["local_audio"].exists():
                rows.append(row)
    selected = {}
    # Prefer SAVEE for consistent speaker/domain examples, then fallback to any dataset.
    for emotion in COMMON_EMOTIONS:
        candidates = [r for r in rows if r["emotion"] == emotion and r["dataset"] == "SAVEE"]
        if not candidates:
            candidates = [r for r in rows if r["emotion"] == emotion]
        selected[emotion] = candidates[len(candidates) // 2]
    return selected


def color_interp(c1, c2, t):
    return tuple(int(a + (b - a) * t) for a, b in zip(c1, c2))


def heat_color(t):
    t = float(np.clip(t, 0, 1))
    stops = [
        (0.0, (25, 20, 72)),
        (0.35, (91, 36, 130)),
        (0.65, (220, 83, 96)),
        (1.0, (255, 230, 120)),
    ]
    for (a, ca), (b, cb) in zip(stops[:-1], stops[1:]):
        if a <= t <= b:
            return color_interp(ca, cb, (t - a) / (b - a))
    return stops[-1][1]


def draw_heat(draw, mat, box, robust=True):
    x1, y1, x2, y2 = box
    arr = np.array(mat, dtype=np.float32)
    if robust:
        lo, hi = np.percentile(arr, 5), np.percentile(arr, 95)
    else:
        lo, hi = arr.min(), arr.max()
    arr = (arr - lo) / (hi - lo + 1e-8)
    h, w = arr.shape
    img = Image.new("RGB", (w, h))
    px = img.load()
    for yy in range(h):
        for xx in range(w):
            px[xx, yy] = heat_color(arr[yy, xx])
    img = img.resize((x2 - x1, y2 - y1), Image.Resampling.BILINEAR)
    draw._image.paste(img, (x1, y1))
    draw.rectangle((x1, y1, x2, y2), outline="#CBD5E1", width=1)


def draw_line(draw, values, box, color="#2563EB", y_range=None, width=3):
    x1, y1, x2, y2 = box
    arr = np.asarray(values, dtype=np.float32)
    if y_range is None:
        lo, hi = np.percentile(arr, 2), np.percentile(arr, 98)
    else:
        lo, hi = y_range
    hi = hi if hi > lo else lo + 1e-6
    pts = []
    for i, v in enumerate(arr):
        x = x1 + (x2 - x1) * i / max(len(arr) - 1, 1)
        y = y2 - (np.clip(v, lo, hi) - lo) / (hi - lo) * (y2 - y1)
        pts.append((x, y))
    draw.rectangle((x1, y1, x2, y2), fill="#FFFFFF", outline="#CBD5E1", width=1)
    for frac in [0.25, 0.5, 0.75]:
        yy = y1 + (y2 - y1) * frac
        draw.line((x1, yy, x2, yy), fill="#EEF2F7", width=1)
    if len(pts) > 1:
        draw.line(pts, fill=color, width=width)


def add_label(draw, box, text, fill="#102A43", font_obj=None):
    x1, y1, x2, y2 = box
    draw.text((x1, y1), text, fill=fill, font=font_obj or F["small"])


def make_canvas(title, subtitle, rows=6, cols=1, cell_w=340, cell_h=120, left=175, top=120):
    w = max(left + cols * cell_w + 45, 1260)
    h = top + rows * cell_h + 70
    img = Image.new("RGB", (w, h), COLORS["bg"])
    d = ImageDraw.Draw(img)
    d.text((28, 25), title, fill=COLORS["ink"], font=F["title"])
    d.text((30, 68), subtitle, fill=COLORS["muted"], font=F["body"])
    return img, d, left, top, cell_w, cell_h


def build_wave_rms(samples):
    img, d, left, top, cw, ch = make_canvas(
        "Evidence A: Waveform and RMS energy across 6 emotions",
        "RMS follows the short-term loudness/intensity of the waveform.",
        cols=2, cell_w=470, cell_h=135,
    )
    d.text((left, top - 32), "Waveform amplitude", fill=COLORS["blue"], font=F["h"])
    d.text((left + cw, top - 32), "Frame-level RMS energy", fill=COLORS["orange"], font=F["h"])
    for r, emotion in enumerate(COMMON_EMOTIONS):
        y = top + r * ch
        d.text((28, y + 42), emotion, fill=COLORS["ink"], font=F["h"])
        f = samples[emotion]["features"]
        draw_line(d, f["waveform"][::40], (left, y + 10, left + cw - 30, y + ch - 20), COLORS["blue"], y_range=(-1, 1), width=2)
        draw_line(d, f["rms"], (left + cw, y + 10, left + 2 * cw - 30, y + ch - 20), COLORS["orange"], y_range=(0, max(0.02, np.percentile(f["rms"], 98))), width=3)
    out = EVIDENCE_DIR / "feature_evidence_waveform_rms_6emotions.png"
    img.save(out)
    return out


def build_mfcc_delta(samples):
    img, d, left, top, cw, ch = make_canvas(
        "Evidence B: MFCC, delta MFCC and delta-delta MFCC across 6 emotions",
        "MFCC shows spectral envelope/timbre; delta and delta-delta show speed and acceleration of change.",
        cols=3, cell_w=330, cell_h=132,
    )
    for i, h in enumerate(["MFCC", "Delta MFCC", "Delta-delta MFCC"]):
        d.text((left + i * cw, top - 32), h, fill=[COLORS["blue"], COLORS["purple"], COLORS["red"]][i], font=F["h"])
    for r, emotion in enumerate(COMMON_EMOTIONS):
        y = top + r * ch
        d.text((28, y + 42), emotion, fill=COLORS["ink"], font=F["h"])
        f = samples[emotion]["features"]
        for c, key in enumerate(["mfcc", "delta", "delta2"]):
            draw_heat(d, f[key], (left + c * cw, y + 8, left + (c + 1) * cw - 20, y + ch - 18))
    out = EVIDENCE_DIR / "feature_evidence_mfcc_delta_6emotions.png"
    img.save(out)
    return out


def build_zcr(samples):
    img, d, left, top, cw, ch = make_canvas(
        "Evidence C: Zero-Crossing Rate across 6 emotions",
        "ZCR is higher when the signal changes sign more often, often linked to sharp/noisy/high-frequency texture.",
        cols=1, cell_w=950, cell_h=125,
    )
    d.text((left, top - 32), "ZCR curve over time", fill=COLORS["green"], font=F["h"])
    for r, emotion in enumerate(COMMON_EMOTIONS):
        y = top + r * ch
        d.text((28, y + 38), emotion, fill=COLORS["ink"], font=F["h"])
        f = samples[emotion]["features"]
        draw_line(d, f["zcr"], (left, y + 8, left + cw - 30, y + ch - 18), COLORS["green"], y_range=(0, max(0.05, np.percentile(f["zcr"], 98))), width=3)
    out = EVIDENCE_DIR / "feature_evidence_zcr_6emotions.png"
    img.save(out)
    return out


def build_spectral_shape(samples):
    img, d, left, top, cw, ch = make_canvas(
        "Evidence D: Spectral centroid, bandwidth and rolloff across 6 emotions",
        "Centroid describes brightness; bandwidth describes spread; rolloff describes high-frequency energy boundary.",
        cols=3, cell_w=330, cell_h=132,
    )
    headers = [("Centroid / brightness", "blue"), ("Bandwidth / spread", "purple"), ("Rolloff / high freq.", "orange")]
    for i, (h, c) in enumerate(headers):
        d.text((left + i * cw, top - 32), h, fill=COLORS[c], font=F["h"])
    for r, emotion in enumerate(COMMON_EMOTIONS):
        y = top + r * ch
        d.text((28, y + 42), emotion, fill=COLORS["ink"], font=F["h"])
        f = samples[emotion]["features"]
        draw_line(d, f["centroid"], (left, y + 8, left + cw - 20, y + ch - 18), COLORS["blue"], y_range=(0, 6000), width=3)
        draw_line(d, f["bandwidth"], (left + cw, y + 8, left + 2 * cw - 20, y + ch - 18), COLORS["purple"], y_range=(0, 5000), width=3)
        draw_line(d, f["rolloff"], (left + 2 * cw, y + 8, left + 3 * cw - 20, y + ch - 18), COLORS["orange"], y_range=(0, 8000), width=3)
    out = EVIDENCE_DIR / "feature_evidence_spectral_shape_6emotions.png"
    img.save(out)
    return out


def build_logmel_contrast(samples):
    img, d, left, top, cw, ch = make_canvas(
        "Evidence E: Log-Mel spectrogram and spectral contrast across 6 emotions",
        "Log-Mel keeps time-frequency energy; contrast shows peak-valley structure in frequency bands.",
        cols=2, cell_w=470, cell_h=135,
    )
    d.text((left, top - 32), "Log-Mel spectrogram", fill=COLORS["cyan"], font=F["h"])
    d.text((left + cw, top - 32), "7-band spectral contrast", fill=COLORS["red"], font=F["h"])
    for r, emotion in enumerate(COMMON_EMOTIONS):
        y = top + r * ch
        d.text((28, y + 42), emotion, fill=COLORS["ink"], font=F["h"])
        f = samples[emotion]["features"]
        draw_heat(d, f["logmel"], (left, y + 8, left + cw - 30, y + ch - 18))
        draw_heat(d, f["contrast"], (left + cw, y + 8, left + 2 * cw - 30, y + ch - 18))
    out = EVIDENCE_DIR / "feature_evidence_logmel_contrast_6emotions.png"
    img.save(out)
    return out


def build_feature_evidence():
    selected = select_samples()
    samples = {}
    for emotion, row in selected.items():
        x = read_wav(row["local_audio"])
        samples[emotion] = {"row": row, "features": compute_features(x)}
    paths = [
        build_wave_rms(samples),
        build_mfcc_delta(samples),
        build_zcr(samples),
        build_spectral_shape(samples),
        build_logmel_contrast(samples),
    ]
    with (EVIDENCE_DIR / "selected_samples.csv").open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["emotion", "sample_id", "dataset", "speaker_id", "source_filename"])
        for emotion in COMMON_EMOTIONS:
            row = samples[emotion]["row"]
            writer.writerow([emotion, row["sample_id"], row["dataset"], row["speaker_id"], row["source_filename"]])
    return paths


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, RGBColor(16, 42, 67), font_size)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.size = Pt(10.5)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    return p


def add_image(doc, path, caption, width=6.5):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(82, 96, 109)


def read_csv(path):
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def pct(x):
    try:
        return f"{float(x) * 100:.2f}%"
    except Exception:
        return str(x)


def build_doc(evidence_paths):
    metrics = read_csv(REPORT_DIR / "06D_emotion2vec_coattention_metrics.csv")
    per_dataset = read_csv(REPORT_DIR / "06D_emotion2vec_coattention_per_dataset.csv")
    refs = read_csv(REPORT_DIR / "06D_reference_model_comparison.csv")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Phân tích chi tiết kiến trúc 06D kèm dẫn chứng đặc trưng trên 6 cảm xúc")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(16, 42, 67)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Các mục Midterm Param: 1.5, 3.1, 3.2, 3.3")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(82, 96, 109)

    add_p(doc, "Tài liệu này viết lại phần proposed solution theo mô hình 06D Emotion2Vec Co-Attention SER. Điểm bổ sung chính so với bản trước là mỗi lập luận về đặc trưng âm thanh đều có hình dẫn chứng từ audio thật trên 6 cảm xúc: neutral, happy, sad, angry, fear và disgust. Các hình này không phải hình minh họa chung; chúng được tạo trực tiếp từ các file WAV trong ser_processed/audio_16k.")
    add_image(doc, ARCH_IMG, "Hình 1. Kiến trúc 06D Emotion2Vec Co-Attention SER.", width=7.0)

    add_heading(doc, "1.5. Select the proposed solution for the project and justify the selection", 1)
    add_p(doc, "Giải pháp được chọn là 06D Emotion2Vec Co-Attention SER, một mô hình Speech Emotion Recognition nhiều biểu diễn. Mô hình nhận audio 16 kHz và dự đoán xác suất cho 6 cảm xúc chung: neutral, happy, sad, angry, fear và disgust.")
    add_p(doc, "Lý do chọn hướng nhiều biểu diễn là vì cảm xúc trong giọng nói không chỉ nằm ở một cue duy nhất. Cùng một nhãn angry chẳng hạn có thể biểu hiện qua biên độ mạnh hơn, RMS cao hơn, phổ sáng hơn, tốc độ thay đổi MFCC mạnh hơn, hoặc pattern log-Mel khác với neutral/sad. Ngược lại, sad có thể có năng lượng thấp hơn, nhịp nói chậm hơn hoặc vùng phổ mờ hơn. Vì vậy, mô hình cần nhìn audio qua nhiều góc: waveform/pretrained speech, temporal MFCC sequence, time-frequency log-Mel và handcrafted statistics.")
    add_table(
        doc,
        ["Nhánh", "Biểu diễn chính", "Cue cảm xúc được khai thác", "Lý do đưa vào 06D"],
        [
            ["A. Temporal acoustic", "MFCC, delta, delta-delta, RMS, ZCR, spectral sequence", "Màu giọng và sự thay đổi theo thời gian.", "Cảm xúc thường là dynamics của giọng nói, nên cần mô hình chuỗi."],
            ["B. Spectrogram", "Log-Mel, delta log-Mel, delta-delta log-Mel", "Vùng năng lượng theo thời gian - tần số.", "CNN học tốt pattern dạng ảnh trên spectrogram."],
            ["C. Emotion pretrained", "Raw waveform -> frozen emotion2vec", "Cue cảm xúc đã được pretrained model học trước.", "Tận dụng tri thức từ model lớn nhưng vẫn nhẹ vì frozen."],
            ["D. Statistical", "Vector thống kê thủ công", "Đặc trưng ổn định toàn clip.", "Bổ sung hướng machine learning cổ điển và dùng được với RBF-SVM."],
        ],
        widths=[1.35, 2.0, 2.0, 2.0],
        font_size=8,
    )
    add_p(doc, "Do tài nguyên hạn chế trên Kaggle/Colab, 06D không fine-tune toàn bộ pretrained model. Thay vào đó, emotion2vec được dùng ở chế độ frozen, chỉ train adapter MLP và các nhánh nhỏ hơn. Đây là lựa chọn cân bằng giữa tính học thuật, khả năng chạy thực tế và khả năng phát triển về sau thành hệ thống feedback giọng nói.")

    add_heading(doc, "3.1. Present the architecture or workflow of the proposed solution", 1)
    add_p(doc, "Workflow tổng thể gồm: input audio 16 kHz -> preprocess -> split train/validation/test -> feature extraction -> bốn branch -> co-attention -> deep classifier -> validation-weighted stacking -> output xác suất 6 cảm xúc. Một nguyên tắc quan trọng là split phải xảy ra trước augmentation hoặc model selection để tránh data leakage.")
    add_heading(doc, "3.1.1. Dẫn chứng trực quan cho các đặc trưng âm thanh", 2)
    add_p(doc, "Phần này bổ sung bằng chứng trực quan cho các lập luận về feature. Mỗi hình có đủ 6 emotion để có thể đối chiếu pattern giữa các cảm xúc. Hình được lấy từ audio thật trong ser_processed, ưu tiên cùng dataset SAVEE để giảm nhiễu do khác domain khi minh họa.")

    add_table(
        doc,
        ["Đặc trưng", "Nó thể hiện điều gì?", "Dẫn chứng hình ảnh", "Liên hệ với branch"],
        [
            ["Waveform + RMS", "Waveform cho thấy biên độ thô; RMS thể hiện độ lớn/năng lượng ngắn hạn của giọng.", "Hình 2", "Branch A dùng RMS theo frame; Branch C dùng waveform cho emotion2vec."],
            ["MFCC + delta + delta-delta", "MFCC mô tả timbre/spectral envelope; delta mô tả tốc độ thay đổi; delta-delta mô tả gia tốc thay đổi.", "Hình 3", "Branch A dùng chuỗi MFCC động; Branch D lấy thống kê MFCC."],
            ["ZCR", "Tần suất đổi dấu của waveform, liên quan voiced/unvoiced, độ sắc/nhiễu/high-frequency texture.", "Hình 4", "Branch A dùng ZCR theo frame; Branch D lấy thống kê ZCR."],
            ["Centroid, bandwidth, rolloff", "Centroid mô tả độ sáng; bandwidth mô tả độ rộng phổ; rolloff mô tả biên năng lượng tần số cao.", "Hình 5", "Branch A và Branch D dùng spectral shape cues."],
            ["Log-Mel + spectral contrast", "Log-Mel giữ bản đồ năng lượng time-frequency; contrast mô tả độ chênh peak-valley theo dải tần.", "Hình 6", "Branch B dùng log-Mel; Branch A/D dùng spectral contrast."],
        ],
        widths=[1.35, 2.55, 0.85, 2.25],
        font_size=8,
    )

    captions = [
        "Hình 2. Waveform và RMS trên 6 emotion. RMS là bằng chứng trực quan cho cường độ/loudness theo thời gian.",
        "Hình 3. MFCC, delta MFCC và delta-delta MFCC trên 6 emotion. Delta/delta-delta cho thấy tốc độ và gia tốc thay đổi của timbre/spectral envelope.",
        "Hình 4. ZCR trên 6 emotion. ZCR cao hơn thường phản ánh tín hiệu đổi dấu nhanh hơn, liên quan texture sắc/nhiễu hoặc high-frequency.",
        "Hình 5. Spectral centroid, bandwidth và rolloff trên 6 emotion. Đây là bằng chứng cho độ sáng, độ rộng phổ và vùng năng lượng tần số cao.",
        "Hình 6. Log-Mel spectrogram và spectral contrast trên 6 emotion. Log-Mel cho thấy cấu trúc time-frequency, contrast cho thấy độ rõ peak-valley theo dải tần.",
    ]
    for path, cap in zip(evidence_paths, captions):
        add_image(doc, path, cap, width=7.0)

    add_heading(doc, "3.1.2. Branch A - Temporal acoustic branch", 2)
    add_p(doc, "Branch A nhận tensor [B, 132, T]. 132 feature gồm 40 MFCC, 40 delta MFCC, 40 delta-delta MFCC, RMS, ZCR, centroid, bandwidth, rolloff và 7 spectral contrast. Hình 2-5 cho thấy vì sao các feature này có ý nghĩa: RMS cho intensity, ZCR cho texture, MFCC/delta cho timbre dynamics, còn spectral features cho brightness và spread của phổ.")
    add_code(doc, "temporal = concatenate([mfcc, delta, delta2, rms, zcr, centroid, bandwidth, rolloff, contrast], axis=0)\n# temporal shape: [132, T]\n# model: 1D-CNN -> BiLSTM/BiGRU -> Attention pooling -> z_temporal")
    add_table(
        doc,
        ["Khối", "Input -> Output", "Hoạt động", "Lý do dùng"],
        [
            ["1D-CNN", "[B,132,T] -> feature maps", "Kernel 1D quét theo thời gian để phát hiện pattern cục bộ.", "Nhẹ và phù hợp chuỗi acoustic frame."],
            ["BiLSTM/BiGRU", "feature maps -> contextual sequence", "Đọc chuỗi hai hướng để lấy ngữ cảnh trước/sau.", "Cảm xúc kéo dài qua nhiều frame."],
            ["Attention pooling", "sequence -> z_temporal", "Học trọng số cho frame quan trọng.", "Không phải mọi frame đều giàu cảm xúc."],
        ],
        widths=[1.2, 1.6, 2.3, 2.2],
        font_size=8,
    )

    add_heading(doc, "3.1.3. Branch B - Spectrogram branch", 2)
    add_p(doc, "Branch B nhận [B, 3, 96, T], gồm log-Mel, delta log-Mel và delta-delta log-Mel. Log-Mel giữ rõ cấu trúc thời gian - tần số hơn MFCC, vì trục ngang là thời gian, trục dọc là Mel bins và màu là năng lượng. Hình 6 cho thấy mỗi emotion có pattern năng lượng khác nhau theo thời gian và dải tần.")
    add_code(doc, "spectral = stack([logmel, delta_logmel, delta2_logmel], axis=0)\n# spectral shape: [3, 96, T]\n# model: 2D-CNN -> Residual CNN -> SE/channel attention -> z_spectral")
    add_table(
        doc,
        ["Khối", "Input -> Output", "Hoạt động", "Lý do dùng"],
        [
            ["2D-CNN", "[B,3,96,T] -> maps", "Học pattern cục bộ trên mặt phẳng time-frequency.", "Spectrogram giống ảnh nên CNN phù hợp."],
            ["Residual CNN", "maps -> deeper maps", "Skip connection giúp train ổn định hơn.", "Giảm mất tín hiệu khi mạng sâu hơn."],
            ["SE attention", "maps -> reweighted maps", "Tính trọng số cho từng channel.", "Nhấn mạnh channel chứa cue cảm xúc."],
        ],
        widths=[1.2, 1.6, 2.3, 2.2],
        font_size=8,
    )

    add_heading(doc, "3.1.4. Branch C - Emotion pretrained branch", 2)
    add_p(doc, "Branch C dùng waveform thô 16 kHz đưa vào frozen emotion2vec. Waveform ở Hình 2 cho thấy tín hiệu thô vẫn chứa biên độ, nhịp và khoảng nghỉ. emotion2vec học representation từ raw speech nên có thể giữ các cue không được mô tả hết bằng MFCC hoặc log-Mel thủ công.")
    add_code(doc, "waveform -> frozen emotion2vec_base -> hidden states -> mean pooling -> Adapter MLP -> z_e2v / p_e2v")
    add_p(doc, "Trong 06D, emotion2vec không fine-tune toàn bộ. Chỉ adapter MLP được train. Cách này phù hợp tài nguyên Kaggle và giảm nguy cơ overfit vào bộ dataset nhỏ.")

    add_heading(doc, "3.1.5. Branch D - Statistical RBF-SVM branch", 2)
    add_p(doc, "Branch D lấy thống kê trên các feature như MFCC, delta, delta-delta, chroma, spectral contrast, RMS, ZCR, energy và entropy. Lý do dùng thống kê là SVM cần vector cố định cho mỗi audio. Hình 2-6 cho thấy các feature gốc thay đổi theo thời gian; Branch D tóm tắt chúng bằng mean, std, min, max, percentile hoặc entropy.")
    add_code(doc, "stat_vec = concatenate([stats(mfcc), stats(delta), stats(delta2), stats(chroma), stats(spectral), stats(rms), stats(zcr), stats(energy), stats(entropy)])\np_stats = StandardScaler() + RBF_SVM(probability=True, class_weight='balanced')")

    add_heading(doc, "3.1.6. Fusion and output", 2)
    add_p(doc, "Ba embedding z_temporal, z_spectral và z_e2v đi qua co-attention để tạo p_deep. Nhánh statistical tạo p_stats riêng. Cuối cùng validation-weighted stacking kết hợp p_deep, p_stats và optional p_e2v để tạo p_final[6]. Output gồm vector xác suất 6 cảm xúc, nhãn dự đoán và confidence.")

    add_heading(doc, "3.2. Implement the selected speech processing model, algorithm, or method", 1)
    add_p(doc, "Notebook 06D triển khai pipeline từ ser_processed. Dữ liệu metadata chứa sample_id, filepath, dataset, speaker_id, emotion và split. Audio đã chuẩn hóa nằm trong audio_16k. Feature extraction tạo bốn nhóm input: X_temporal, X_spectral, X_e2v và X_stats.")
    add_table(
        doc,
        ["Bước", "Triển khai", "Kiểm soát chất lượng"],
        [
            ["Load data", "Đọc metadata.csv và audio_16k.", "Giữ dataset/speaker_id để đánh giá random và strict."],
            ["Feature extraction", "Tạo temporal, spectral, emotion2vec embedding, statistical vector.", "Các feature có hình evidence từ 6 emotion trong tài liệu này."],
            ["Split", "combined_random và combined_strict_no_tess.", "Không augment hoặc chọn tham số bằng test set."],
            ["Deep training", "Train temporal + spectrogram + emotion2vec adapter + co-attention.", "Chọn best epoch bằng validation macro-F1."],
            ["SVM training", "Train StandardScaler + RBF-SVM trên X_stats.", "Fit scaler/SVM trên train split בלבד."],
            ["Stacking", "Kết hợp xác suất theo validation.", "Test chỉ dùng báo cáo cuối."],
        ],
        widths=[1.2, 3.0, 3.0],
        font_size=8,
    )
    add_p(doc, "Các metric chính gồm accuracy, macro-F1, weighted-F1 và UAR. Macro-F1 được ưu tiên vì bài toán có nhiều lớp cảm xúc; metric này tránh việc lớp nhiều mẫu làm kết quả nhìn tốt giả tạo.")

    add_heading(doc, "3.3. Provide preliminary experimental results or initial training results to demonstrate the feasibility of the proposed solution", 1)
    metrics = [r for r in metrics if r.get("split") == "test"]
    add_table(
        doc,
        ["Protocol", "Model", "n", "Accuracy", "Macro-F1", "UAR"],
        [[r["protocol"], r["model"], r["n_samples"], pct(r["accuracy"]), pct(r["macro_f1"]), pct(r["uar"])] for r in metrics],
        widths=[1.35, 1.9, 0.55, 0.9, 0.9, 0.9],
        font_size=7.5,
    )
    add_p(doc, "Kết quả chính: stacking_full đạt khoảng 80.39% accuracy và 80.51% macro-F1 trên combined_random. Trên combined_strict_no_tess, stacking_full đạt khoảng 69.59% accuracy và 69.60% macro-F1. Điều này chứng minh mô hình khả thi, nhưng cũng cho thấy bài toán generalization còn khó.")
    add_image(doc, LEADERBOARD_IMG, "Hình 7. Macro-F1 leaderboard từ output 06D.", width=6.5)
    add_image(doc, CURVES_IMG, "Hình 8. Training curves từ output 06D.", width=6.5)

    add_heading(doc, "3.3.1. Kết quả theo từng dataset", 2)
    per_rows = [[r["protocol"], r["dataset"], r["n_samples"], pct(r["accuracy"]), pct(r["macro_f1"]), pct(r["uar"])] for r in per_dataset if r.get("model") == "stacking_full"]
    add_table(doc, ["Protocol", "Dataset", "n", "Accuracy", "Macro-F1", "UAR"], per_rows, widths=[1.5, 1.0, 0.55, 1.0, 1.0, 1.0], font_size=8)
    add_p(doc, "Per-dataset result cho thấy TESS dễ hơn trong combined_random, còn CREMA-D khó hơn do speaker/domain đa dạng. Strict protocol làm kết quả giảm vì mô hình phải tổng quát tốt hơn sang speaker/domain ít trùng với train.")
    add_image(doc, CM_RANDOM_IMG, "Hình 9. Confusion matrix combined_random.", width=5.8)
    add_image(doc, CM_STRICT_IMG, "Hình 10. Confusion matrix combined_strict_no_tess.", width=5.8)

    add_heading(doc, "3.3.2. So sánh với references và nhận xét", 2)
    add_table(
        doc,
        ["Reference", "Protocol / split", "Reported result", "Main idea"],
        [[r["model"], r["protocol"], r["reported_accuracy_text"], r["main_idea"]] for r in refs],
        widths=[1.6, 1.7, 1.4, 2.3],
        font_size=7.5,
    )
    add_bullets(doc, [
        "Các bài báo có accuracy trên 90% thường dùng single-dataset hoặc random split; không phải lúc nào cũng strict speaker-aware như bài của mình.",
        "06D mạnh hơn baseline ở khả năng kết hợp nhiều biểu diễn, nhưng statistical branch hiện còn yếu, đặc biệt ở strict split.",
        "Khoảng cách train/validation trong curves cho thấy mô hình vẫn có dấu hiệu overfit speaker/domain.",
        "Future work hợp lý là tăng chất lượng feature engineering, domain-aware adapter, calibration và đánh giá single-dataset để so sánh công bằng hơn với references.",
    ])

    doc.save(DOCX_PATH)
    return DOCX_PATH


def make_zip(docx, evidence_paths):
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    files = [docx, ARCH_IMG, CURVES_IMG, LEADERBOARD_IMG, CM_RANDOM_IMG, CM_STRICT_IMG, EVIDENCE_DIR / "selected_samples.csv", *evidence_paths]
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as z:
        for p in files:
            if Path(p).exists():
                z.write(p, Path(p).relative_to(ROOT))
    return ZIP_PATH


def main():
    evidence = build_feature_evidence()
    docx = build_doc(evidence)
    zipped = make_zip(docx, evidence)
    print("DOCX", docx)
    print("ZIP", zipped)
    for p in evidence:
        print("EVIDENCE", p)


if __name__ == "__main__":
    main()
